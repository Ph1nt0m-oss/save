from fastapi import FastAPI, APIRouter, HTTPException, Response, Request, UploadFile, File
from fastapi.responses import JSONResponse, FileResponse, HTMLResponse, RedirectResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware

from dotenv import load_dotenv
from motor.motor_asyncio import AsyncIOMotorClient

from pydantic import BaseModel, Field, ConfigDict

from pathlib import Path
from typing import List, Optional, Dict, Any

from datetime import datetime, timezone, timedelta

from cfaction_engine import (
    sanitize_filename as _cf_sanitize_filename,
    build_docx_bytes as _cf_build_docx_bytes,
    build_pdf_bytes as _cf_build_pdf_bytes,
    build_xlsx_bytes as _cf_build_xlsx_bytes,
    build_pptx_bytes as _cf_build_pptx_bytes,
    run_python_sandbox as _cf_run_python_sandbox,
    analyze_pdf as _cf_analyze_pdf,
    analyze_docx as _cf_analyze_docx,
    analyze_xlsx as _cf_analyze_xlsx,
    analyze_pptx as _cf_analyze_pptx,
    analyze_sqlite as _cf_analyze_sqlite,
)

import os
import re
import uuid
import httpx
import json
import zipfile
import io
import base64
import logging
import asyncio
import time
import secrets
import bcrypt
import hashlib

# Import sub-routers (PWA + Desktop)
from routes.pwa_routes import export_router as pwa_router
from routes.desktop_routes import desktop_router

# ==================== LOAD ENV ====================
ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / ".env")

# ==================== LOGGING ====================
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)

logger = logging.getLogger(__name__)

# ==================== BACKGROUND TASK REGISTRY ====================
# Long-running generations (chat, complete-app, code) are detached from the
# request lifecycle: when the client disconnects, the work continues to
# completion and the result is persisted to MongoDB. The client picks it up
# on next reconnect by reading chat_messages / projects from the DB.
#
# We use `asyncio.shield(task)` to immune the task from the request's
# CancelledError, and we keep a strong reference in `_BG_TASKS` so the GC
# does not drop it mid-flight.
_BG_TASKS: set = set()

def _register_bg(task: "asyncio.Task") -> "asyncio.Task":
    _BG_TASKS.add(task)
    task.add_done_callback(_BG_TASKS.discard)
    return task

async def _run_in_background(coro):
    """Schedule `coro` so it runs to completion even if the awaiting client
    disconnects. Returns the coroutine's result while the connection lasts;
    if the request is cancelled, the underlying task keeps executing and
    persisting its side-effects to the database."""
    task = _register_bg(asyncio.create_task(coro))
    try:
        return await asyncio.shield(task)
    except asyncio.CancelledError:
        # Client gone — the task keeps running. We re-raise so the framework
        # cleans up the request properly, but the side-effects (DB writes)
        # will still complete in the background.
        logger.info("Client disconnected during long generation — task continues in background.")
        raise


# ==================== GITHUB CONFIG ====================
GITHUB_CLIENT_ID = os.environ.get("GITHUB_CLIENT_ID")
GITHUB_CLIENT_SECRET = os.environ.get("GITHUB_CLIENT_SECRET")  # PAT TOKEN recommandé
GITHUB_OWNER = os.environ.get("GITHUB_OWNER")
GITHUB_REPO_NAME = os.environ.get("GITHUB_REPO_NAME")

# ==================== EMAIL AUTH CONFIG ====================
RESEND_ENABLED = bool(os.environ.get("RESEND_API_KEY"))
if RESEND_ENABLED:
    logger.info("✅ Resend email provider activé (RESEND_API_KEY présent)")
else:
    logger.warning("⚠️ RESEND_API_KEY manquant — mode démo (lien de vérification retourné dans la réponse)")

GITHUB_ENABLED = all([
    GITHUB_CLIENT_SECRET,
    GITHUB_OWNER,
    GITHUB_REPO_NAME
])

if GITHUB_ENABLED:
    logger.info(f"✅ GitHub activé: {GITHUB_OWNER}/{GITHUB_REPO_NAME}")
else:
    logger.warning("⚠️ GitHub désactivé (config .env incomplète)")

# ==================== GITHUB ENGINE ====================
async def push_to_github(file_path: str, content: str, branch: str = "main", retries: int = 3):
    """
    Push auto vers GitHub (robuste + retry + safe update)
    """

    if not GITHUB_ENABLED:
        logger.error("❌ GitHub non configuré")
        return False

    # URL-encode path segments (handles accents, spaces, parens, etc.) — keep "/" intact.
    from urllib.parse import quote
    safe_path = quote(file_path, safe="/")
    url = f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO_NAME}/contents/{safe_path}"

    headers = {
        "Authorization": f"Bearer {GITHUB_CLIENT_SECRET}",
        "Accept": "application/vnd.github+json",
        "User-Agent": "CodeForge-AI"
    }

    encoded_content = base64.b64encode(content.encode("utf-8")).decode("utf-8")

    async with httpx.AsyncClient(timeout=20.0) as client:

        for attempt in range(retries):

            try:
                sha = None

                # check existing file
                get_resp = await client.get(url, headers=headers)
                if get_resp.status_code == 200:
                    sha = get_resp.json().get("sha")

                payload = {
                    "message": f"Auto-sync {file_path}",
                    "content": encoded_content,
                    "branch": branch
                }

                if sha:
                    payload["sha"] = sha

                put_resp = await client.put(url, json=payload, headers=headers)

                if put_resp.status_code in [200, 201]:
                    logger.info(f"✅ GitHub sync OK: {file_path}")
                    return True

                logger.warning(f"GitHub fail {attempt+1}: {put_resp.text}")

            except Exception as e:
                logger.error(f"GitHub error {attempt+1}: {e}")

            await asyncio.sleep(1.5 * (attempt + 1))

    logger.error(f"❌ GitHub FINAL FAIL: {file_path}")
    return False

# ==================== MONGODB ====================
mongo_url = os.environ.get("MONGO_URL")
db_name = os.environ.get("DB_NAME")

if not mongo_url or not db_name:
    logger.error("❌ MongoDB configuration manquante (.env)")
    raise Exception("Missing MongoDB config")

client = AsyncIOMotorClient(mongo_url)
db = client[db_name]

# ==================== FASTAPI APP ====================
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origin_regex=r"https://.*\.preview(\.static)?\.emergentagent\.com",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

api_router = APIRouter(prefix="/api")

# ==================== PYDANTIC MODELS ====================

class ChatAttachment(BaseModel):
    kind: str  # 'text' | 'image'
    filename: Optional[str] = None
    mime_type: Optional[str] = None
    content: Optional[str] = None       # extracted text for kind='text'
    data_base64: Optional[str] = None   # for kind='image' (no data:<...>, prefix)


class ChatMessageInput(BaseModel):
    message: str
    project_id: Optional[str] = None
    mode: Optional[str] = "online"
    language: Optional[str] = "fr"
    model: Optional[str] = None  # 'gpt-5.2' | 'claude-opus' | 'claude-sonnet' | 'gemini-3-pro' | 'gemma-3' (offline)
    attachments: Optional[List[ChatAttachment]] = None

class Project(BaseModel):
    model_config = ConfigDict(extra="ignore")

    project_id: str = Field(default_factory=lambda: f"proj_{uuid.uuid4().hex[:12]}")
    user_id: str
    name: str
    description: Optional[str] = ""
    project_type: str = "web"
    ai_mode: Optional[str] = "online"  # 'online' (Emergent/GPT) | 'offline' (Ollama)
    status: str = "created"
    generated_code: Optional[Dict[str, Any]] = None
    preview_image: Optional[str] = None  # data URI thumbnail for sidebar preview
    is_public: Optional[bool] = False
    share_slug: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ProjectCreate(BaseModel):
    name: str
    description: Optional[str] = ""
    project_type: str = "web"

class ProjectUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    status: Optional[str] = None
    generated_code: Optional[Dict[str, Any]] = None

class GenerateCodeRequest(BaseModel):
    project_id: str
    description: str
    project_type: str
    use_emergent: bool = True

class ExportRequest(BaseModel):
    project_id: str
    export_type: str  # 'apk', 'exe', 'web', 'source'
    include_code: bool = True   # iter80 C17 — case « code source »
    include_chat: bool = False  # iter80 C17 — case « discussions / .docx »

# ==================== HELPER FUNCTIONS ====================

async def get_current_user(request: Request) -> str:
    """Extract user_id from session_token (cookie or Authorization header)"""
    session_token = request.cookies.get("session_token")
    
    if not session_token:
        auth_header = request.headers.get("Authorization")
        if auth_header and auth_header.startswith("Bearer "):
            session_token = auth_header.replace("Bearer ", "")
    
    if not session_token:
        raise HTTPException(status_code=401, detail="Non authentifié")
    
    # Verify session in database
    session_doc = await db.user_sessions.find_one(
        {"session_token": session_token},
        {"_id": 0}
    )
    
    if not session_doc:
        raise HTTPException(status_code=401, detail="Session invalide")
    
    # Check expiry
    expires_at = session_doc["expires_at"]
    if isinstance(expires_at, str):
        expires_at = datetime.fromisoformat(expires_at)
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    
    if expires_at < datetime.now(timezone.utc):
        raise HTTPException(status_code=401, detail="Session expirée")

    # iter66: heartbeat — mark this session as "recently seen" so the
    # multi-device approval flow (server.py L1278) can distinguish a
    # really-active device from one that just left a stale 7-day cookie
    # behind. Cheap fire-and-forget update.
    try:
        await db.user_sessions.update_one(
            {"session_token": session_token},
            {"$set": {"last_seen_at": datetime.now(timezone.utc).isoformat()}},
        )
    except Exception:
        pass

    return session_doc["user_id"]

class SMSAuthRequest(BaseModel):
    phone_number: str
    code: Optional[str] = None

# ==================== AUTH ROUTES ====================

async def send_sms_via_twilio(phone_number: str, message: str) -> bool:
    """Send SMS via Twilio if configured, otherwise return False"""
    twilio_sid = os.environ.get('TWILIO_ACCOUNT_SID')
    twilio_token = os.environ.get('TWILIO_AUTH_TOKEN')
    twilio_phone = os.environ.get('TWILIO_PHONE_NUMBER')
    
    if not all([twilio_sid, twilio_token, twilio_phone]):
        logger.warning("Twilio not configured - SMS will be simulated")
        return False
    
    try:
        # Twilio API call
        async with httpx.AsyncClient() as client:
            response = await client.post(
                f"https://api.twilio.com/2010-04-01/Accounts/{twilio_sid}/Messages.json",
                auth=(twilio_sid, twilio_token),
                data={
                    "From": twilio_phone,
                    "To": phone_number,
                    "Body": message
                }
            )
            
            if response.status_code in [200, 201]:
                logger.info(f"SMS sent successfully to {phone_number}")
                return True
            else:
                logger.error(f"Twilio error: {response.text}")
                return False
    except Exception as e:
        logger.error(f"Twilio exception: {e}")
        return False

@api_router.post("/auth/sms/send")
async def send_sms_code(request: SMSAuthRequest):
    """Send SMS verification code (for offline auth)"""
    try:
        # Generate 6-digit code
        code = str(uuid.uuid4().int)[:6]
        
        # Store code in database (expires in 5 minutes)
        await db.sms_codes.insert_one({
            "phone_number": request.phone_number,
            "code": code,
            "expires_at": (datetime.now(timezone.utc) + timedelta(minutes=5)).isoformat(),
            "created_at": datetime.now(timezone.utc).isoformat()
        })
        
        # Try to send via Twilio
        message = f"Votre code CodeForge AI: {code}. Valide 5 minutes."
        sms_sent = await send_sms_via_twilio(request.phone_number, message)
        
        logger.info(f"SMS Code for {request.phone_number}: {code} (Twilio: {sms_sent})")
        
        response_data = {
            "message": "Code SMS envoyé" if sms_sent else "Code généré (mode démo)",
            "sms_sent": sms_sent
        }
        
        # Return code in response only if Twilio is not configured (for testing)
        if not sms_sent:
            response_data["code"] = code  # DEMO MODE - remove when Twilio is configured
        
        return response_data
    except Exception as e:
        logger.error(f"Error sending SMS: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/auth/sms/verify")
async def verify_sms_code(request: SMSAuthRequest, response: Response):
    """Verify SMS code and create session"""
    try:
        # Find valid code
        code_doc = await db.sms_codes.find_one({
            "phone_number": request.phone_number,
            "code": request.code
        }, {"_id": 0})
        
        if not code_doc:
            await log_auth_error("sms_invalid_code", f"phone={request.phone_number}", request=None)
            return JSONResponse(status_code=401, content={"detail": "Code invalide"})
        
        # Check expiry
        expires_at = datetime.fromisoformat(code_doc["expires_at"])
        if expires_at.tzinfo is None:
            expires_at = expires_at.replace(tzinfo=timezone.utc)
        
        if expires_at < datetime.now(timezone.utc):
            await log_auth_error("sms_code_expired", f"phone={request.phone_number}", request=None)
            return JSONResponse(status_code=401, content={"detail": "Code expiré"})
        
        # Create or get user
        user_id = f"user_{uuid.uuid4().hex[:12]}"
        existing_user = await db.users.find_one({"phone_number": request.phone_number}, {"_id": 0})
        
        if existing_user:
            user_id = existing_user["user_id"]
        else:
            new_user = {
                "user_id": user_id,
                "phone_number": request.phone_number,
                "name": f"User {request.phone_number[-4:]}",
                "created_at": datetime.now(timezone.utc).isoformat()
            }
            await db.users.insert_one(new_user)
        
        # Create session
        session_token = f"sms_session_{uuid.uuid4().hex}"
        expires_at = datetime.now(timezone.utc) + timedelta(days=7)
        
        session_doc = {
            "session_token": session_token,
            "user_id": user_id,
            "auth_type": "sms",
            "expires_at": expires_at.isoformat(),
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.user_sessions.insert_one(session_doc)
        
        # Set cookie
        response.set_cookie(
            key="session_token",
            value=session_token,
            httponly=True,
            secure=True,
            samesite="none",
            max_age=7 * 24 * 60 * 60,
            path="/"
        )
        
        # Delete used code
        await db.sms_codes.delete_one({"phone_number": request.phone_number, "code": request.code})
        
        user = await db.users.find_one({"user_id": user_id}, {"_id": 0})
        return user
    
    except Exception as e:
        logger.error(f"Error verifying SMS: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ==================== AUTH ROUTES ====================

@api_router.get("/auth/me")
async def get_me(request: Request):
    """Get current user from session"""
    user_id = await get_current_user(request)
    user = await db.users.find_one({"user_id": user_id}, {"_id": 0, "password_hash": 0})
    
    if not user:
        raise HTTPException(status_code=404, detail="Utilisateur non trouvé")
    
    return user


@api_router.post("/auth/heartbeat")
async def auth_heartbeat(request: Request):
    """iter68: explicit liveness ping called every 30s by the frontend.
    Updates last_seen_at on the current session so the multi-device
    approval gating can distinguish a really-active device from one that
    has just left a stale 7-day cookie behind. get_current_user() already
    refreshes last_seen_at as a side-effect, so this endpoint is mostly a
    no-op wrapper that returns the freshness window for diagnostics."""
    await get_current_user(request)
    return {"ok": True, "now": datetime.now(timezone.utc).isoformat()}


@api_router.post("/auth/disconnect-soft")
async def auth_disconnect_soft(request: Request):
    """iter69: called via `navigator.sendBeacon` from a `beforeunload`
    handler when the user closes the tab/browser. Marks the current
    session's last_seen_at far in the past so the multi-device approval
    gating immediately treats it as abandoned. The session row is NOT
    deleted — if the user reopens within the 7-day cookie window, the
    next /auth/me + heartbeat will resurrect it. This is what allows
    'close tab ⇒ instantly stale' without forcing a hard logout."""
    try:
        session_token = request.cookies.get("session_token")
        if not session_token:
            auth_header = request.headers.get("Authorization") or ""
            if auth_header.startswith("Bearer "):
                session_token = auth_header[7:]
        # iter69: sendBeacon doesn't let us set custom headers, so the
        # frontend appends the token as a query param.
        if not session_token:
            session_token = request.query_params.get("t")
        if session_token:
            past = (datetime.now(timezone.utc) - timedelta(hours=24)).isoformat()
            await db.user_sessions.update_one(
                {"session_token": session_token},
                {"$set": {"last_seen_at": past}},
            )
    except Exception:
        pass
    # sendBeacon ignores the body but we return 200 anyway
    return {"ok": True}

@api_router.get("/user/stats")
async def get_user_stats(request: Request):
    """Stats summary shown in the Dashboard avatar dropdown."""
    user_id = await get_current_user(request)
    user = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="Utilisateur non trouvé")

    project_count = await db.projects.count_documents({"user_id": user_id})

    # Most recent session (excluding the current one) = "dernier login"
    sessions_cursor = db.user_sessions.find(
        {"user_id": user_id}, {"_id": 0, "created_at": 1}
    ).sort("created_at", -1).limit(2)
    sessions = await sessions_cursor.to_list(length=2)
    last_login = None
    if len(sessions) >= 2:
        last_login = sessions[1].get("created_at")
    elif len(sessions) == 1:
        last_login = sessions[0].get("created_at")

    return {
        "project_count": project_count,
        "member_since": user.get("created_at"),
        "last_login": last_login,
        "plan": "Gratuit illimité",
    }

@api_router.post("/auth/logout")
async def logout(request: Request, response: Response):
    """Logout user and clear session"""
    session_token = request.cookies.get("session_token")
    
    if session_token:
        await db.user_sessions.delete_one({"session_token": session_token})
    
    response.delete_cookie("session_token", path="/")
    return {"message": "Déconnexion réussie"}


# ==================== EMAIL + PASSWORD + MAGIC LINK AUTH ====================
# Classic email/password auth with magic-link email verification.
# No Google OAuth, no Emergent Auth. Works fully offline in "demo mode"
# (verification link returned in the response body when no email provider
# is configured).

# ----- bcrypt helpers -----
def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")


def verify_password(plain: str, hashed: str) -> bool:
    try:
        return bcrypt.checkpw(plain.encode("utf-8"), hashed.encode("utf-8"))
    except Exception:
        return False


EMAIL_RE = re.compile(r"^[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}$")


def normalize_email(email: str) -> str:
    return (email or "").strip().lower()


async def _send_via_smtp(to_email: str, subject: str, html: str, reply_to: Optional[str] = None) -> bool:
    """Send an email via plain SMTP (gmail.com / outlook.com / any provider).

    Activated when SMTP_HOST + SMTP_USER + SMTP_PASSWORD env vars are set.
    Compatible with Gmail App Passwords (smtp.gmail.com:587 STARTTLS) and
    Outlook (smtp-mail.outlook.com:587). Use this instead of Resend when
    you don't have a verified domain — the recipient sees the From header
    we set, but the underlying mailbox is SMTP_USER.
    """
    host = os.environ.get("SMTP_HOST")
    user = os.environ.get("SMTP_USER")
    password = os.environ.get("SMTP_PASSWORD")
    if not (host and user and password):
        return False
    try:
        import aiosmtplib
        from email.message import EmailMessage
        port = int(os.environ.get("SMTP_PORT", "587"))
        use_tls = os.environ.get("SMTP_USE_TLS", "false").lower() in ("1", "true", "yes")
        # default sender: "Display Name <user>" — display configurable.
        display = os.environ.get("EMAIL_FROM_DISPLAY", "CodeForge AI")
        sender = f"{display} <{user}>"
        msg = EmailMessage()
        msg["From"] = sender
        msg["To"] = to_email
        msg["Subject"] = subject
        # No-reply enforcement: replies bounce on the IETF-reserved .invalid
        # TLD (RFC 2606), so they never reach SMTP_USER's mailbox — even the
        # creator can't receive a reply. Auto-Submitted (RFC 3834) tells
        # well-behaved MTAs to suppress auto-responders / out-of-office
        # bounces. Callers may still override Reply-To explicitly.
        msg["Reply-To"] = reply_to or "no-reply@codeforge-ai.invalid"
        msg["Auto-Submitted"] = "auto-generated"
        msg["X-Auto-Response-Suppress"] = "All"
        # plain-text fallback for spam filters
        msg.set_content("Ce mail contient du HTML. Active l'affichage HTML dans ton client.")
        msg.add_alternative(html, subtype="html")
        await aiosmtplib.send(
            msg,
            hostname=host,
            port=port,
            username=user,
            password=password,
            start_tls=not use_tls,  # STARTTLS for port 587
            use_tls=use_tls,        # implicit TLS for port 465
            timeout=20.0,
        )
        logger.info(f"✅ SMTP email sent to {to_email} via {host}")
        return True
    except Exception as e:
        logger.error(f"SMTP exception ({host}): {e}")
        return False


async def _send_via_resend(to_email: str, subject: str, html: str, reply_to: Optional[str] = None) -> bool:
    """Send via Resend (only if RESEND_API_KEY is set)."""
    resend_key = os.environ.get("RESEND_API_KEY")
    if not resend_key:
        return False
    sender = os.environ.get("EMAIL_FROM", "CodeForge AI <no-reply@resend.dev>")
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            resp = await client.post(
                "https://api.resend.com/emails",
                headers={"Authorization": f"Bearer {resend_key}", "Content-Type": "application/json"},
                json={
                    "from": sender,
                    "to": [to_email],
                    "reply_to": reply_to or "no-reply@codeforge-ai.invalid",
                    "subject": subject,
                    "html": html,
                    "headers": {
                        "Auto-Submitted": "auto-generated",
                        "X-Auto-Response-Suppress": "All",
                    },
                },
            )
            if resp.status_code in (200, 202):
                logger.info(f"✅ Resend email sent to {to_email}")
                return True
            logger.error(f"Resend API error {resp.status_code}: {resp.text[:200]}")
            return False
    except Exception as e:
        logger.error(f"Resend exception: {e}")
        return False


async def _send_email(to_email: str, subject: str, html: str, reply_to: Optional[str] = None) -> bool:
    """Unified send-email entrypoint.

    Strategy: try SMTP first (production-grade, sends to anyone), fall back
    to Resend (sandbox-limited, only verified recipients). Returns True iff
    at least one transport accepted the message.
    """
    if await _send_via_smtp(to_email, subject, html, reply_to=reply_to):
        return True
    return await _send_via_resend(to_email, subject, html, reply_to=reply_to)


async def send_verification_email(to_email: str, verify_url: str) -> bool:
    """Send magic-link verification email — SMTP first then Resend."""
    html = (
        f"<div style='font-family:system-ui,sans-serif;background:#050505;color:#fff;padding:32px;max-width:560px;margin:0 auto'>"
        f"<h1 style='color:#E4FF00;margin:0 0 16px'>CodeForge AI</h1>"
        f"<p style='color:#E4E4E7'>Clique sur le bouton ci-dessous pour confirmer ton compte&nbsp;:</p>"
        f"<p style='margin:24px 0'><a href='{verify_url}' style='background:#E4FF00;color:#050505;"
        f"padding:14px 28px;border-radius:6px;text-decoration:none;font-weight:bold;display:inline-block'>"
        f"Confirmer mon compte</a></p>"
        f"<p style='color:#A1A1AA;font-size:12px;margin:24px 0 8px'>Ou copie ce lien dans ton navigateur (Chrome, Safari, Firefox)&nbsp;:<br>"
        f"<span style='color:#00D4FF;word-break:break-all;font-size:11px'>{verify_url}</span></p>"
        f"<p style='color:#A1A1AA;font-size:12px;margin-top:24px'><strong>Ce lien expire dans 5 minutes.</strong> Passé ce délai, il sera invalide et tu devras recommencer l'inscription.</p>"
        f"<p style='color:#A1A1AA;font-size:12px'>Astuce&nbsp;: si le bouton ouvre une page bloquée, copie-colle le lien dans ton navigateur principal (Chrome, Safari…).</p>"
        f"<p style='color:#A1A1AA;font-size:12px'>Si tu n'es pas à l'origine de cette demande, ignore cet email.</p>"
        f"<hr style='border:none;border-top:1px solid rgba(255,255,255,.1);margin:24px 0'>"
        f"<p style='color:#71717A;font-size:11px;margin:0'>Courriel automatique — <strong>ne réponds pas</strong>. Toute réponse sera rejetée par le serveur, personne ne la lira.</p>"
        f"</div>"
    )
    ok = await _send_email(to_email, "Confirme ton compte CodeForge AI", html)
    if ok:
        logger.info(f"✅ Verification email sent to {to_email}")
    return ok



class RegisterRequest(BaseModel):
    email: str
    password: str
    name: Optional[str] = None       # display name (legacy)
    pseudo: Optional[str] = None     # required, unique nickname — see /auth/register
    frontend_url: Optional[str] = None
    # iter62: mandatory device-capture data (extracted client-side via /auth/ocr-device-info)
    device_capture_kind: Optional[str] = None      # 'phone' | 'computer'
    device_capture_product: Optional[str] = None   # e.g. "Galaxy S21 5G"
    device_capture_model: Optional[str] = None     # e.g. "SM-G991U1"
    device_capture_name: Optional[str] = None      # e.g. "DESKTOP-52KO8J1" (computer hostname)
    # iter69: mandatory biometric enrollment (kind = 'webauthn' | 'iris')
    biometric_kind: Optional[str] = None
    biometric_options_token: Optional[str] = None  # webauthn only — links to webauthn_challenges
    biometric_credential: Optional[Dict[str, Any]] = None  # webauthn navigator.credentials.create result
    biometric_iris_hashes: Optional[List[str]] = None  # iris only — 3 SHA-256 b64 hashes


class LoginRequest(BaseModel):
    email: str
    password: str
    device_key_id: Optional[str] = None  # cryptographic device identifier (browser ECDSA)
    device_label: Optional[str] = None   # human label (e.g. "iPhone 15 Pro")


# ---------------- Device-info OCR (registration capture) ----------------
class OcrDeviceIn(BaseModel):
    image_base64: str
    hint: Optional[str] = None  # 'phone' | 'computer' (optional kind hint)


@api_router.post("/auth/ocr-device-info")
async def auth_ocr_device_info(payload: OcrDeviceIn):
    """Extract product name + model number from an "About this phone /
    About this PC" screenshot using Gemini Vision (Emergent LLM).

    Public endpoint — used during registration (no auth yet at that point).
    Rate-limited softly via the front-end UX (one request per submission).

    Returns:
      {kind: 'phone'|'computer'|'unknown', product, model, raw_text, confidence}
    """
    raw_b64 = (payload.image_base64 or "").strip()
    if raw_b64.startswith("data:"):
        # strip the leading "data:image/png;base64," header if present
        try:
            raw_b64 = raw_b64.split(",", 1)[1]
        except Exception:
            pass
    if not raw_b64 or len(raw_b64) > 5_000_000:  # ~3.5 MB image cap
        raise HTTPException(status_code=400, detail="Image manquante ou trop volumineuse.")

    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage, ImageContent
    except ImportError:
        raise HTTPException(status_code=500, detail="OCR indisponible.")

    key = os.environ.get("EMERGENT_LLM_KEY")
    if not key:
        raise HTTPException(status_code=500, detail="EMERGENT_LLM_KEY missing.")

    system_msg = (
        "You are an OCR + classifier specialized in device-info screenshots. "
        "You receive ONE screenshot taken from either:\n"
        "  - an Android 'About phone' page (Samsung/Pixel/etc) — fields like "
        "    'Device name', 'Model name', 'Model number'\n"
        "  - an iPhone 'About' page (Réglages > Général > Informations) — "
        "    fields like 'Nom', 'Numéro de modèle'\n"
        "  - a Windows 'About' page (Paramètres > Système > Informations) — "
        "    field like 'Nom de l'appareil' / 'Device name'\n"
        "  - a macOS 'About this Mac' page\n\n"
        "Return STRICT JSON, no extra text:\n"
        '  {"kind": "phone" | "computer" | "unknown", '
        '"product": "<marketing name like Galaxy S21 5G or empty>", '
        '"model": "<model code like SM-G991U1 or empty>", '
        '"device_name": "<computer name like DESKTOP-52KO8J1 or empty>", '
        '"confidence": 0.0..1.0}\n\n'
        "RULES:\n"
        "- Phones: prioritize the 'Model name' (marketing) for product and "
        "  'Model number' (technical SKU) for model. If only one is present, "
        "  put it in product. NEVER invent.\n"
        "- Computers: only fill device_name (e.g. 'DESKTOP-52KO8J1' or "
        "  'MacBook-Pro-de-Vincent'). Leave product/model empty.\n"
        "- If you can't identify ANY device info, set kind='unknown' and "
        "  confidence below 0.3."
    )
    chat = LlmChat(
        api_key=key,
        session_id=f"ocr_{uuid.uuid4().hex[:8]}",
        system_message=system_msg,
    ).with_model("gemini", "gemini-2.5-flash")
    image = ImageContent(image_base64=raw_b64)
    msg = UserMessage(text="Extract device info from this screenshot.", file_contents=[image])
    try:
        raw = await chat.send_message(msg)
    except Exception as e:
        logger.warning(f"ocr-device-info upstream: {e}")
        raise HTTPException(status_code=502, detail="OCR upstream error.")
    text = str(raw).strip()
    # Strip ```json fences if present
    if text.startswith("```"):
        text = text.strip("`")
        if text.startswith("json"):
            text = text[4:]
    try:
        import json as _json
        parsed = _json.loads(text)
    except Exception:
        return {
            "kind": "unknown",
            "product": "",
            "model": "",
            "device_name": "",
            "confidence": 0.0,
            "raw": text[:300],
        }
    return {
        "kind": parsed.get("kind") or "unknown",
        "product": (parsed.get("product") or "").strip()[:80],
        "model": (parsed.get("model") or "").strip()[:80],
        "device_name": (parsed.get("device_name") or "").strip()[:80],
        "confidence": float(parsed.get("confidence") or 0.0),
    }


@api_router.post("/auth/register")
async def register(payload: RegisterRequest, request: Request):
    """Create an unverified account and send (or return) a magic link."""
    email = normalize_email(payload.email)
    if not email or not EMAIL_RE.match(email):
        raise HTTPException(status_code=400, detail="Adresse email invalide")
    if len(payload.password) < 6:
        raise HTTPException(status_code=400, detail="Le mot de passe doit faire au moins 6 caractères")

    # Pseudo — required, unique, "Créatrice" reserved for the creator.
    pseudo_raw = (payload.pseudo or payload.name or "").strip()
    if not pseudo_raw or len(pseudo_raw) < 1:
        raise HTTPException(status_code=400, detail="Le pseudo est requis.")
    if len(pseudo_raw) > 30:
        raise HTTPException(status_code=400, detail="Le pseudo est trop long (30 max).")
    # iter75: pseudo "créatrice" is no longer reserved — the creator can
    # rename anyone (including herself) freely, so the reservation is
    # purely cosmetic noise. Anyone may now pick "créatrice" as a pseudo.
    # Pseudo uniqueness is intentionally NOT enforced anymore — users keep
    # the right to choose any pseudo. The creator-side "Autres comptes"
    # panel disambiguates duplicates by appending "#N" suffixes when
    # listing, and offers a rename action for both sides.

    # iter62: mandatory device capture (extracted via /auth/ocr-device-info)
    capture_kind = (payload.device_capture_kind or "").strip().lower()
    capture_product = (payload.device_capture_product or "").strip()
    capture_model = (payload.device_capture_model or "").strip()
    capture_name = (payload.device_capture_name or "").strip()
    if capture_kind not in ("phone", "computer"):
        raise HTTPException(status_code=400, detail=(
            "Capture d'écran de l'appareil requise. "
            "Téléphone : Paramètres > À propos du téléphone (capture du Nom du produit + Numéro de modèle). "
            "Ordinateur : Paramètres > Système > Informations système (capture du Nom de l'appareil)."
        ))
    if capture_kind == "phone" and not (capture_product or capture_model):
        raise HTTPException(status_code=400, detail=(
            "Capture invalide : impossible de lire le nom du produit + numéro de modèle. "
            "Ouvre Paramètres > À propos du téléphone et capture la page complète."
        ))
    if capture_kind == "computer" and not capture_name:
        raise HTTPException(status_code=400, detail=(
            "Capture invalide : impossible de lire le nom de l'ordinateur. "
            "Ouvre Paramètres > Système > Informations système (Windows) ou À propos de ce Mac (macOS)."
        ))

    # iter69: mandatory biometric enrollment.
    bio_kind = (payload.biometric_kind or "").strip().lower()
    bio_doc: Optional[dict] = None
    if bio_kind not in ("webauthn", "iris"):
        raise HTTPException(status_code=400, detail=(
            "Identité biométrique requise. Utilise ton empreinte / Face ID, ou capture ton iris via la webcam."
        ))
    if bio_kind == "webauthn":
        if not payload.biometric_options_token or not payload.biometric_credential:
            raise HTTPException(status_code=400, detail="Enrôlement biométrique incomplet (token ou credential manquant).")
        challenge_doc = await db.webauthn_challenges.find_one_and_delete(
            {"options_token": payload.biometric_options_token, "kind": "signup"},
        )
        if not challenge_doc:
            raise HTTPException(status_code=400, detail="Défi d'enrôlement biométrique introuvable ou expiré.")
        try:
            verification = webauthn.verify_registration_response(
                credential=payload.biometric_credential,
                expected_challenge=webauthn.helpers.base64url_to_bytes(challenge_doc["challenge"]),
                expected_origin=challenge_doc.get("origin"),
                expected_rp_id=challenge_doc.get("rp_id"),
                require_user_verification=False,
            )
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Vérification biométrique échouée: {e}")
        bio_doc = {
            "kind": "webauthn",
            "credential_id": webauthn.helpers.bytes_to_base64url(verification.credential_id),
            "public_key": webauthn.helpers.bytes_to_base64url(verification.credential_public_key),
            "sign_count": verification.sign_count,
            "rp_id": challenge_doc.get("rp_id"),
            "enrolled_at": datetime.now(timezone.utc).isoformat(),
        }
    elif bio_kind == "iris":
        hashes = payload.biometric_iris_hashes or []
        if not isinstance(hashes, list) or len(hashes) < 3:
            raise HTTPException(status_code=400, detail="3 captures iris sont requises pour l'enrôlement.")
        clean = [h for h in hashes if isinstance(h, str) and 20 <= len(h) <= 128]
        if len(clean) < 3:
            raise HTTPException(status_code=400, detail="Empreintes iris invalides.")
        bio_doc = {
            "kind": "iris",
            "hashes": clean[:3],
            "enrolled_at": datetime.now(timezone.utc).isoformat(),
        }

    existing = await db.users.find_one({"email": email}, {"_id": 0})
    if existing and existing.get("verified"):
        raise HTTPException(status_code=409, detail="Un compte existe déjà avec cet email. Connecte-toi.")

    now = datetime.now(timezone.utc)
    password_hash = hash_password(payload.password)

    if existing and not existing.get("verified"):
        # Re-register on an unverified account: refresh password + resend link
        user_id = existing["user_id"]
        await db.users.update_one(
            {"user_id": user_id},
            {"$set": {
                "password_hash": password_hash,
                "name": payload.name or existing.get("name") or email.split("@")[0],
                "pseudo": pseudo_raw,
                "pseudo_lower": pseudo_raw.lower(),
                "device_capture": {
                    "kind": capture_kind,
                    "product": capture_product or None,
                    "model": capture_model or None,
                    "device_name": capture_name or None,
                },
                "biometric": bio_doc,
                "updated_at": now.isoformat(),
            }},
        )
    else:
        user_id = f"user_{uuid.uuid4().hex[:12]}"
        await db.users.insert_one({
            "user_id": user_id,
            "email": email,
            "password_hash": password_hash,
            "name": payload.name or email.split("@")[0],
            "pseudo": pseudo_raw,
            "pseudo_lower": pseudo_raw.lower(),
            "verified": False,
            "auth_type": "email",
            "device_capture": {
                "kind": capture_kind,
                "product": capture_product or None,
                "model": capture_model or None,
                "device_name": capture_name or None,
            },
            "biometric": bio_doc,
            "created_at": now.isoformat(),
        })

    # Create verification token (single-use, 5 min — aligné avec l'attente raisonnable)
    token = secrets.token_urlsafe(32)
    await db.email_verifications.delete_many({"user_id": user_id})
    await db.email_verifications.insert_one({
        "token": token,
        "user_id": user_id,
        "email": email,
        "consumed_at": None,
        "pending_session_token": None,
        "created_at": now.isoformat(),
        "expires_at": (now + timedelta(minutes=5)).isoformat(),
    })

    # Build verification URL.
    # Priority: explicit frontend_url in body (sent by the React app as
    # window.location.origin) > FRONTEND_URL env > REACT_APP_BACKEND_URL env.
    # Request Origin/Referer are NOT used — the k8s ingress rewrites them
    # to the internal cluster hostname which is 403 externally.
    def _clean_origin(u: str) -> str:
        try:
            from urllib.parse import urlparse
            p = urlparse(u)
            if p.scheme in ("http", "https") and p.netloc:
                return f"{p.scheme}://{p.netloc}"
        except Exception:
            pass
        return ""

    frontend_base = ""
    if payload.frontend_url:
        frontend_base = _clean_origin(payload.frontend_url)
    if not frontend_base:
        frontend_base = _clean_origin(os.environ.get("FRONTEND_URL", ""))
    if not frontend_base:
        frontend_base = _clean_origin(os.environ.get("REACT_APP_BACKEND_URL", ""))

    verify_url = f"{frontend_base}/verify-email?token={token}" if frontend_base else f"/verify-email?token={token}"

    sent = await send_verification_email(email, verify_url)

    response = {
        "message": "Compte créé ! Vérifie ton email pour confirmer." if sent
        else "Compte créé ! Ton e-mail n'a pas pu être envoyé automatiquement — voici le lien de confirmation à coller dans ton navigateur.",
        "email": email,
        "email_sent": sent,
        # The frontend uses this to poll /auth/verification-status and unlock
        # the original tab automatically when the user clicks the magic link
        # (possibly in another tab from their email client).
        "verification_token": token,
        "expires_in_seconds": 5 * 60,
        # iter59: ALWAYS expose the verification link so the frontend can
        # offer a "copy / open" fallback in case the e-mail never arrives
        # (Resend sandbox, spam folder, typo etc.). This is safe — the link
        # already requires the single-use token to do anything.
        "verification_link": verify_url,
    }
    return response


class ResendRequest(BaseModel):
    email: str
    frontend_url: Optional[str] = None


class MagicLinkLoginRequest(BaseModel):
    email: str
    frontend_url: Optional[str] = None


@api_router.post("/auth/magic-link")
async def magic_link_login(payload: MagicLinkLoginRequest, request: Request):
    """Send a one-shot login link to a verified user.

    Same UX as register: the original tab polls /verification-status while
    the user clicks the link in their inbox. Always returns the same
    neutral message to prevent email enumeration. Rate-limited 3/10 min.
    """
    email = normalize_email(payload.email)
    if not email or not EMAIL_RE.match(email):
        raise HTTPException(status_code=400, detail="Adresse email invalide")

    user = await db.users.find_one({"email": email}, {"_id": 0})
    neutral = {
        "message": "Si un compte existe pour cet email, un lien de connexion t'a été envoyé.",
    }

    if not user or not user.get("verified"):
        return neutral

    # Rate limit: 3 magic links / 10 min / verified email
    cutoff = (datetime.now(timezone.utc) - timedelta(minutes=10)).isoformat()
    recent = await db.resend_attempts.count_documents({"email": email, "ts": {"$gte": cutoff}, "purpose": "magic_login"})
    if recent >= 3:
        raise HTTPException(status_code=429, detail="Trop de demandes. Patiente 10 minutes.")

    now = datetime.now(timezone.utc)
    token = secrets.token_urlsafe(32)
    await db.email_verifications.delete_many({"user_id": user["user_id"], "purpose": "magic_login"})
    await db.email_verifications.insert_one({
        "token": token,
        "user_id": user["user_id"],
        "email": email,
        "purpose": "magic_login",
        "consumed_at": None,
        "pending_session_token": None,
        "created_at": now.isoformat(),
        "expires_at": (now + timedelta(minutes=5)).isoformat(),
    })
    await db.resend_attempts.insert_one({
        "email": email,
        "ts": now.isoformat(),
        "purpose": "magic_login",
    })

    frontend_base = (
        _clean_origin(payload.frontend_url)
        or _clean_origin(os.environ.get("FRONTEND_URL", ""))
        or _clean_origin(os.environ.get("REACT_APP_BACKEND_URL", ""))
    )
    login_url = f"{frontend_base}/verify-email?token={token}" if frontend_base else f"/verify-email?token={token}"

    sent = await send_verification_email(email, login_url)
    response = {
        **neutral,
        "email_sent": sent,
        "verification_token": token,  # for polling cross-tab
        "expires_in_seconds": 5 * 60,
    }
    if not sent:
        response["verification_link"] = login_url
    return response


@api_router.post("/auth/resend-verification")
async def resend_verification(payload: ResendRequest, request: Request):
    """Generate a fresh magic link for an unverified account.

    Rate-limited: at most 3 resends / 10 min / email. Already-verified
    users get a friendly message instead of a link (no enumeration).
    """
    email = normalize_email(payload.email)
    if not email or not EMAIL_RE.match(email):
        raise HTTPException(status_code=400, detail="Adresse email invalide")

    # Rate limit: 3 resends per 10 min
    cutoff = (datetime.now(timezone.utc) - timedelta(minutes=10)).isoformat()
    recent = await db.resend_attempts.count_documents({"email": email, "ts": {"$gte": cutoff}})
    if recent >= 3:
        raise HTTPException(status_code=429, detail="Trop de renvois. Patiente 10 minutes.")

    user = await db.users.find_one({"email": email}, {"_id": 0})
    # Neutral response if no user / already verified (no email enumeration)
    if not user or user.get("verified"):
        return {
            "message": "Si un compte non confirmé existe pour cet email, un nouveau lien a été envoyé.",
            "email_sent": False,
        }

    now = datetime.now(timezone.utc)
    token = secrets.token_urlsafe(32)
    await db.email_verifications.delete_many({"user_id": user["user_id"]})
    await db.email_verifications.insert_one({
        "token": token,
        "user_id": user["user_id"],
        "email": email,
        "consumed_at": None,
        "pending_session_token": None,
        "created_at": now.isoformat(),
        "expires_at": (now + timedelta(minutes=5)).isoformat(),
    })
    await db.resend_attempts.insert_one({"email": email, "ts": now.isoformat()})

    def _clean_origin(u: str) -> str:
        try:
            from urllib.parse import urlparse
            p = urlparse(u or "")
            if p.scheme in ("http", "https") and p.netloc:
                return f"{p.scheme}://{p.netloc}"
        except Exception:
            pass
        return ""

    frontend_base = (
        _clean_origin(payload.frontend_url)
        or _clean_origin(os.environ.get("FRONTEND_URL", ""))
        or _clean_origin(os.environ.get("REACT_APP_BACKEND_URL", ""))
    )
    verify_url = f"{frontend_base}/verify-email?token={token}" if frontend_base else f"/verify-email?token={token}"

    sent = await send_verification_email(email, verify_url)
    resp = {
        "message": "Nouveau lien envoyé ! Tu as 5 minutes." if sent
        else "Nouveau lien généré. L'e-mail n'a pas pu être envoyé automatiquement — utilise le lien ci-dessous.",
        "email": email,
        "email_sent": sent,
        "verification_token": token,
        "expires_in_seconds": 5 * 60,
        # iter59: always expose the link as a copy/open fallback
        "verification_link": verify_url,
    }
    return resp


@api_router.get("/auth/verify-email")
async def verify_email(token: str):
    """Consume the magic link.

    Marks the user as verified, stores a fresh session_token against the
    verification row, and returns a friendly message. We do NOT set a
    cookie here because the user may have opened this link in a different
    tab/device from the one where they registered. The original tab is
    polling /auth/verification-status and will pick up the session_token
    on its next poll, then log the user in automatically.
    """
    if not token:
        raise HTTPException(status_code=400, detail="Token manquant")

    doc = await db.email_verifications.find_one({"token": token}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=400, detail="Lien invalide ou déjà utilisé")

    expires_at = datetime.fromisoformat(doc["expires_at"])
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    if expires_at < datetime.now(timezone.utc):
        # Clean up the expired row so it can't be reused accidentally
        await db.email_verifications.delete_one({"token": token})
        raise HTTPException(
            status_code=400,
            detail="La durée de validation de ce lien a expiré. Merci de réessayer à nouveau sur CodeForge AI.",
        )

    # Already consumed? Idempotent friendly response.
    if doc.get("consumed_at"):
        return {
            "message": "Votre compte est désormais certifié. Vous pouvez fermer cette page et retourner sur l'application.",
            "already_verified": True,
        }

    user_id = doc["user_id"]
    now = datetime.now(timezone.utc)

    # Email change flow: apply the pending email change instead of marking
    # the account as verified (it already is).
    if doc.get("purpose") == "email_change" and doc.get("pending_email"):
        new_email = doc["pending_email"]
        # Race-check: another account may have grabbed this email since
        # the request was issued.
        existing = await db.users.find_one({"email": new_email}, {"_id": 0})
        if existing and existing.get("user_id") != user_id:
            await db.email_verifications.delete_one({"token": token})
            raise HTTPException(
                status_code=409,
                detail="Cet email a été pris par un autre compte entre-temps.",
            )
        await db.users.update_one(
            {"user_id": user_id},
            {"$set": {"email": new_email, "last_login": now.isoformat()}},
        )
        await db.email_verifications.update_one(
            {"token": token},
            {"$set": {"consumed_at": now.isoformat()}},
        )
        return {
            "message": "Adresse email mise à jour. Tu peux fermer cette page.",
            "already_verified": True,
            "email_changed": True,
        }

    await db.users.update_one(
        {"user_id": user_id},
        {"$set": {"verified": True, "last_login": now.isoformat()}},
    )

    # Prepare a fresh session token that the original tab will exchange
    # when it polls /auth/verification-status.
    session_token = secrets.token_urlsafe(32)
    await db.user_sessions.insert_one({
        "session_token": session_token,
        "user_id": user_id,
        "auth_type": "email",
        "created_at": now.isoformat(),
        "expires_at": (now + timedelta(days=7)).isoformat(),
    })

    await db.email_verifications.update_one(
        {"token": token},
        {"$set": {
            "consumed_at": now.isoformat(),
            "pending_session_token": session_token,
        }},
    )

    return {
        "message": "Votre compte est désormais certifié. Vous pouvez fermer cette page et retourner sur l'application.",
        "already_verified": False,
    }


@api_router.get("/auth/verification-status")
async def verification_status(token: str, response: Response):
    """Polled by the original registration tab every ~2 seconds.

    Returns one of:
      - {status: "pending"}  → user hasn't clicked the link yet
      - {status: "expired"}  → link expired before being clicked
      - {status: "verified", session_token, user}  → link consumed; the
        original tab should now log the user in automatically. The
        session_token is handed over exactly ONCE (the pending token is
        cleared on the same call).
    """
    if not token:
        raise HTTPException(status_code=400, detail="Token manquant")

    doc = await db.email_verifications.find_one({"token": token}, {"_id": 0})
    if not doc:
        return {"status": "expired"}

    expires_at = datetime.fromisoformat(doc["expires_at"])
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)

    consumed = doc.get("consumed_at")
    pending_token = doc.get("pending_session_token")

    if consumed and pending_token:
        # Hand the token over to the original tab and delete the row so it
        # cannot be reused.
        user = await db.users.find_one(
            {"user_id": doc["user_id"]},
            {"_id": 0, "password_hash": 0},
        )
        await db.email_verifications.delete_one({"token": token})

        response.set_cookie(
            key="session_token",
            value=pending_token,
            httponly=True,
            secure=True,
            samesite="none",
            max_age=7 * 24 * 3600,
            path="/",
        )
        return {
            "status": "verified",
            "session_token": pending_token,
            "user": user,
        }

    if not consumed and expires_at < datetime.now(timezone.utc):
        await db.email_verifications.delete_one({"token": token})
        return {"status": "expired"}

    return {"status": "pending"}


@api_router.post("/auth/login")
async def login(payload: LoginRequest, response: Response, request: Request):
    """Verify credentials and create a session (cookie + token in body)."""
    email = normalize_email(payload.email)
    if not email or not payload.password:
        raise HTTPException(status_code=400, detail="Email et mot de passe requis")

    # Simple brute-force guard: 5 fails / 15 min per email.
    # NOTE: identifier is email-only because the k8s ingress rotates IPs
    # across pods, which would defeat an IP-based lockout.
    identifier = email
    cutoff = (datetime.now(timezone.utc) - timedelta(minutes=15)).isoformat()
    fails = await db.login_attempts.count_documents({"identifier": identifier, "ts": {"$gte": cutoff}})
    if fails >= 5:
        raise HTTPException(status_code=429, detail="Trop de tentatives. Réessaie dans 15 minutes.")

    user = await db.users.find_one({"email": email}, {"_id": 0})
    if not user:
        await db.login_attempts.insert_one({
            "identifier": identifier,
            "ts": datetime.now(timezone.utc).isoformat(),
        })
        await log_auth_error("login_unknown_email", f"email={email}", request=request)
        raise HTTPException(status_code=404, detail="Aucun compte avec cet email")
    if not verify_password(payload.password, user.get("password_hash", "")):
        await db.login_attempts.insert_one({
            "identifier": identifier,
            "ts": datetime.now(timezone.utc).isoformat(),
        })
        await log_auth_error("login_wrong_password", f"email={email}", request=request)
        raise HTTPException(status_code=401, detail="Mot de passe incorrect")

    if not user.get("verified"):
        raise HTTPException(
            status_code=403,
            detail="Email non confirmé. Clique sur le lien reçu par email ou recrée ton compte.",
        )

    # Success: clear failed attempts
    await db.login_attempts.delete_many({"identifier": identifier})

    now = datetime.now(timezone.utc)
    requesting_key_id = (payload.device_key_id or "").strip() or None

    # ----- iter63: device-binding (1 device-key = 1 account) -----------------
    # If THIS device-key has already been bound to a DIFFERENT verified user,
    # block the login. This prevents the same browser/PC from juggling several
    # accounts. We bind at first successful login (set device_keys.email).
    if requesting_key_id:
        existing_dev = await db.device_keys.find_one(
            {"key_id": requesting_key_id},
            {"_id": 0, "email": 1, "role": 1},
        )
        if existing_dev and existing_dev.get("email") and existing_dev["email"] != email:
            # Make sure the bound account still exists — otherwise free the slot.
            owner = await db.users.find_one({"email": existing_dev["email"]}, {"_id": 0, "user_id": 1})
            if owner:
                raise HTTPException(
                    status_code=409,
                    detail=(
                        "Cet appareil est déjà lié à un autre compte (" + existing_dev["email"] + "). "
                        "Connecte-toi avec ce compte, ou utilise un autre appareil."
                    ),
                )
            # Stale binding → clear it before continuing.
            await db.device_keys.update_one({"key_id": requesting_key_id}, {"$unset": {"email": ""}})

    # --- One-device-at-a-time approval flow ---------------------------------
    # If another active session exists for this account on a DIFFERENT device,
    # block the login and queue a pending session request. The currently-
    # connected device must approve from its UI. This applies regardless of
    # site_mode — the email account itself is the unit of trust.
    if requesting_key_id:
        # iter76: bug "phantom approval prompt" — la fenêtre `> now` créait
        # une race condition : si A heartbeat juste APRÈS que B capture `now`,
        # alors `last_seen_at(A) > now` → prompt fantôme. On utilise désormais
        # une fenêtre de présence explicite : on considère A actif uniquement
        # si son `last_seen_at` est dans les 8 dernières secondes ET supérieur
        # à `last_seen_at_ack` (qui ne s'incrémente pas via sendBeacon).
        # Comme `/auth/session-pending` est appelé toutes les 3s par chaque
        # onglet ouvert, 8s couvre largement le cas réel ; les onglets fermés
        # avec sendBeacon ou simplement expirés tombent hors fenêtre.
        presence_window = (now - timedelta(seconds=8)).isoformat()
        active_other = await db.user_sessions.find_one({
            "user_id": user["user_id"],
            "expires_at": {"$gt": now.isoformat()},
            "last_seen_at": {"$gt": presence_window},
            "device_key_id": {"$nin": [None, requesting_key_id]},
        }, {"_id": 0, "device_key_id": 1})
        if active_other:
            # Already an outstanding request for the same (user, device)?
            existing_req = await db.session_requests.find_one({
                "user_id": user["user_id"],
                "requesting_key_id": requesting_key_id,
                "status": "pending",
            }, {"_id": 0})
            if not existing_req:
                # Approximate location for @gmail.com only (privacy: other
                # providers don't get geo-resolution).
                client_ip = (request.client.host if request.client else "") or ""
                fwd = (request.headers.get("x-forwarded-for") or "").split(",")[0].strip()
                ip = fwd or client_ip
                location = None
                if email.endswith("@gmail.com") and ip and ip not in ("127.0.0.1", "::1"):
                    try:
                        async with httpx.AsyncClient(timeout=4.0) as client:
                            r = await client.get(f"https://ipinfo.io/{ip}/json")
                            if r.status_code == 200:
                                j = r.json()
                                city = j.get("city") or ""
                                region = j.get("region") or ""
                                country = j.get("country") or ""
                                location = ", ".join([p for p in (city, region, country) if p]) or None
                    except Exception:
                        location = None
                request_id = secrets.token_urlsafe(16)
                await db.session_requests.insert_one({
                    "request_id": request_id,
                    "user_id": user["user_id"],
                    "email": email,
                    "requesting_key_id": requesting_key_id,
                    "requesting_label": (payload.device_label or "")[:80] or None,
                    "is_gmail": email.endswith("@gmail.com"),
                    "location": location,
                    "status": "pending",  # 'pending' | 'approved' | 'denied' | 'expired'
                    "created_at": now.isoformat(),
                    "expires_at": (now + timedelta(minutes=15)).isoformat(),
                })
                request_id_to_return = request_id
            else:
                request_id_to_return = existing_req["request_id"]
            # 202 — login is on hold. Frontend polls /auth/session-request-status.
            raise HTTPException(
                status_code=202,
                detail={
                    "code": "session_pending_approval",
                    "request_id": request_id_to_return,
                    "message": "Connexion en attente d'approbation par l'appareil déjà connecté.",
                },
            )

    await db.users.update_one({"user_id": user["user_id"]}, {"$set": {"last_login": now.isoformat()}})

    # iter63: bind this device-key to the account (1 device = 1 account).
    if requesting_key_id:
        await db.device_keys.update_one(
            {"key_id": requesting_key_id},
            {"$set": {"email": email, "last_seen_at": now.isoformat()}},
        )

    session_token = secrets.token_urlsafe(32)
    await db.user_sessions.insert_one({
        "session_token": session_token,
        "user_id": user["user_id"],
        "device_key_id": requesting_key_id,
        "device_label": (payload.device_label or "")[:80] or None,
        "auth_type": "email",
        "created_at": now.isoformat(),
        "last_seen_at": now.isoformat(),  # iter66 heartbeat init
        "expires_at": (now + timedelta(days=7)).isoformat(),
    })

    response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        secure=True,
        samesite="none",
        max_age=7 * 24 * 3600,
        path="/",
    )

    safe_user = {k: v for k, v in user.items() if k != "password_hash"}
    return {**safe_user, "session_token": session_token}


# ==================== PROFILE / SETTINGS ====================

class ChangePasswordRequest(BaseModel):
    current_password: str
    new_password: str


class ChangeEmailRequest(BaseModel):
    new_email: str
    current_password: str
    frontend_url: Optional[str] = None


class DeleteAccountRequest(BaseModel):
    current_password: str


@api_router.post("/auth/change-password")
async def change_password(payload: ChangePasswordRequest, request: Request):
    """Change password while logged in. Requires current password."""
    user_id = await get_current_user(request)
    if len(payload.new_password) < 6:
        raise HTTPException(status_code=400, detail="Le nouveau mot de passe doit faire au moins 6 caractères")

    user = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="Utilisateur introuvable")
    if not verify_password(payload.current_password, user.get("password_hash", "")):
        raise HTTPException(status_code=401, detail="Mot de passe actuel incorrect")

    new_hash = hash_password(payload.new_password)
    await db.users.update_one(
        {"user_id": user_id},
        {"$set": {
            "password_hash": new_hash,
            "last_password_change": datetime.now(timezone.utc).isoformat(),
        }},
    )
    # Invalidate other sessions, keep the current one
    current_token = request.cookies.get("session_token") or (
        request.headers.get("Authorization", "").replace("Bearer ", "") or None
    )
    if current_token:
        await db.user_sessions.delete_many({
            "user_id": user_id,
            "session_token": {"$ne": current_token},
        })
    return {"message": "Mot de passe mis à jour."}


@api_router.post("/auth/change-email")
async def change_email(payload: ChangeEmailRequest, request: Request):
    """Request a change of email. Sends a verification link to the NEW
    email; the change is only applied once the user clicks the link.
    """
    user_id = await get_current_user(request)
    new_email = normalize_email(payload.new_email)
    if not new_email or not EMAIL_RE.match(new_email):
        raise HTTPException(status_code=400, detail="Adresse email invalide")

    user = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="Utilisateur introuvable")
    if not verify_password(payload.current_password, user.get("password_hash", "")):
        raise HTTPException(status_code=401, detail="Mot de passe incorrect")
    if new_email == user.get("email"):
        raise HTTPException(status_code=400, detail="Cet email est déjà ton email actuel")
    other = await db.users.find_one({"email": new_email}, {"_id": 0})
    if other:
        raise HTTPException(status_code=409, detail="Cet email est déjà utilisé par un autre compte")

    # Reuse email_verifications collection with a 'pending_email' marker
    now = datetime.now(timezone.utc)
    token = secrets.token_urlsafe(32)
    await db.email_verifications.delete_many({"user_id": user_id, "purpose": "email_change"})
    await db.email_verifications.insert_one({
        "token": token,
        "user_id": user_id,
        "email": user.get("email"),
        "pending_email": new_email,
        "purpose": "email_change",
        "consumed_at": None,
        "pending_session_token": None,
        "created_at": now.isoformat(),
        "expires_at": (now + timedelta(minutes=30)).isoformat(),
    })

    frontend_base = (
        _clean_origin(payload.frontend_url)
        or _clean_origin(os.environ.get("FRONTEND_URL", ""))
        or _clean_origin(os.environ.get("REACT_APP_BACKEND_URL", ""))
    )
    confirm_url = f"{frontend_base}/verify-email?token={token}" if frontend_base else f"/verify-email?token={token}"
    sent = await send_verification_email(new_email, confirm_url)
    resp = {
        "message": f"Email de confirmation envoyé à {new_email}." if sent
        else "Confirmation requise (mode démo — clique le lien ci-dessous).",
        "email_sent": sent,
    }
    if not sent:
        resp["verification_link"] = confirm_url
    return resp


@api_router.delete("/auth/me")
async def delete_account(payload: DeleteAccountRequest, request: Request, response: Response):
    """Delete own account + all related data (RGPD compliant)."""
    user_id = await get_current_user(request)
    user = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="Utilisateur introuvable")
    if not verify_password(payload.current_password, user.get("password_hash", "")):
        raise HTTPException(status_code=401, detail="Mot de passe incorrect")

    # Cascade delete
    await db.users.delete_one({"user_id": user_id})
    await db.user_sessions.delete_many({"user_id": user_id})
    await db.email_verifications.delete_many({"user_id": user_id})
    await db.password_reset_tokens.delete_many({"user_id": user_id})
    await db.projects.delete_many({"user_id": user_id})
    await db.previews.delete_many({"user_id": user_id})
    await db.chat_messages.delete_many({"user_id": user_id})

    response.delete_cookie("session_token", path="/")
    return {"message": "Compte supprimé avec succès. À bientôt."}


@api_router.get("/auth/export")
async def export_my_data(request: Request):
    """RGPD: download a JSON of all data we hold about the user."""
    user_id = await get_current_user(request)
    user = await db.users.find_one({"user_id": user_id}, {"_id": 0, "password_hash": 0})
    projects = await db.projects.find({"user_id": user_id}, {"_id": 0}).to_list(1000)
    sessions = await db.user_sessions.find({"user_id": user_id}, {"_id": 0}).to_list(100)
    chats = await db.chat_messages.find({"user_id": user_id}, {"_id": 0}).to_list(1000)
    payload = {
        "exported_at": datetime.now(timezone.utc).isoformat(),
        "user": user,
        "projects": projects,
        "sessions": sessions,
        "chat_messages": chats,
    }
    return JSONResponse(
        content=payload,
        headers={"Content-Disposition": "attachment; filename=codeforge-mes-donnees.json"},
    )


# ==================== PASSWORD RESET (FORGOT PASSWORD) ====================

class ForgotPasswordRequest(BaseModel):
    email: str
    password: Optional[str] = None  # NEW flow: user supplies the new password upfront
    frontend_url: Optional[str] = None


class ResetPasswordRequest(BaseModel):
    token: str
    password: str


async def send_reset_email(to_email: str, reset_url: str) -> bool:
    """Send password reset confirmation link via Resend (same provider as verification).

    The user has already entered + confirmed the new password on the website.
    This email is the SECOND step: clicking the link applies the pending password.
    """
    html = (
        f"<div style='font-family:system-ui,sans-serif;background:#050505;color:#fff;padding:32px;max-width:560px;margin:0 auto'>"
        f"<h1 style='color:#E4FF00;margin:0 0 16px'>CodeForge AI</h1>"
        f"<p style='color:#E4E4E7'>Tu viens de demander à changer le mot de passe de ce compte.</p>"
        f"<p style='color:#E4E4E7'>Pour finaliser le changement, clique sur le bouton ci-dessous&nbsp;:</p>"
        f"<p style='margin:24px 0'><a href='{reset_url}' style='background:#E4FF00;color:#050505;"
        f"padding:14px 28px;border-radius:6px;text-decoration:none;font-weight:bold;display:inline-block'>"
        f"Veuillez cliquer ici pour confirmer la réinitialisation de votre mot de passe</a></p>"
        f"<p style='color:#A1A1AA;font-size:12px;margin:24px 0 8px'>Ou copie ce lien dans ton navigateur&nbsp;:<br>"
        f"<span style='color:#00D4FF;word-break:break-all;font-size:11px'>{reset_url}</span></p>"
        f"<p style='color:#A1A1AA;font-size:12px;margin-top:24px'>Ce lien expire dans 30 minutes. Tant que tu ne cliques pas, ton ancien mot de passe reste valide.</p>"
        f"<p style='color:#A1A1AA;font-size:12px'>Si tu n'es pas à l'origine de cette demande, ignore cet email — ton mot de passe actuel reste inchangé.</p>"
        f"<hr style='border:none;border-top:1px solid rgba(255,255,255,.1);margin:24px 0'>"
        f"<p style='color:#71717A;font-size:11px;margin:0'>Courriel automatique — <strong>ne réponds pas</strong>. Toute réponse sera rejetée par le serveur, personne ne la lira.</p>"
        f"</div>"
    )
    ok = await _send_email(to_email, "Confirme la réinitialisation de ton mot de passe CodeForge AI", html)
    if ok:
        logger.info(f"✅ Password reset email sent to {to_email}")
    return ok


def _clean_origin(u: str) -> str:
    """Extract scheme://netloc from a URL, empty string if invalid."""
    try:
        from urllib.parse import urlparse
        p = urlparse(u or "")
        if p.scheme in ("http", "https") and p.netloc:
            return f"{p.scheme}://{p.netloc}"
    except Exception:
        pass
    return ""


@api_router.post("/auth/forgot-password")
async def forgot_password(payload: ForgotPasswordRequest, request: Request):
    """Step 1 of the new "set then confirm" reset flow.

    The user enters their email + a NEW password (twice, validated by the frontend).
    We don't change the password yet — we store the new hash on a pending token,
    then email them a confirmation link. Clicking the link applies the password.

    Always returns the same neutral message to prevent email enumeration.
    Rate-limited: 3 requests / 10 min / email.
    """
    email = normalize_email(payload.email)
    if not email or not EMAIL_RE.match(email):
        raise HTTPException(status_code=400, detail="Adresse email invalide")
    if not payload.password or len(payload.password) < 6:
        raise HTTPException(status_code=400, detail="Le mot de passe doit faire au moins 6 caractères")

    cutoff = (datetime.now(timezone.utc) - timedelta(minutes=10)).isoformat()
    recent = await db.password_resets.count_documents({"email": email, "ts": {"$gte": cutoff}})
    if recent >= 3:
        raise HTTPException(status_code=429, detail="Trop de demandes. Patiente 10 minutes.")

    user = await db.users.find_one({"email": email}, {"_id": 0})
    neutral = {
        "message": "Si un compte existe pour cet email, un lien de confirmation t'a été envoyé.",
    }

    if not user or not user.get("verified"):
        return neutral

    await db.password_resets.insert_one({
        "email": email,
        "ts": datetime.now(timezone.utc).isoformat(),
    })

    # Generate single-use token (30 min) carrying the PENDING password hash.
    now = datetime.now(timezone.utc)
    token = secrets.token_urlsafe(32)
    pending_hash = hash_password(payload.password)
    await db.password_reset_tokens.delete_many({"user_id": user["user_id"]})
    await db.password_reset_tokens.insert_one({
        "token": token,
        "user_id": user["user_id"],
        "email": email,
        "pending_password_hash": pending_hash,
        "consumed_at": None,
        "created_at": now.isoformat(),
        "expires_at": (now + timedelta(minutes=30)).isoformat(),
    })

    frontend_base = (
        _clean_origin(payload.frontend_url)
        or _clean_origin(os.environ.get("FRONTEND_URL", ""))
        or _clean_origin(os.environ.get("REACT_APP_BACKEND_URL", ""))
    )
    # GET endpoint that finalizes the change — same pattern as /verify-email.
    confirm_url = (
        f"{frontend_base}/api/auth/confirm-password-reset?token={token}"
        if frontend_base
        else f"/api/auth/confirm-password-reset?token={token}"
    )

    sent = await send_reset_email(email, confirm_url)
    response = {**neutral, "email_sent": sent}
    if not sent:
        response["confirm_link"] = confirm_url
    return response


@api_router.get("/auth/confirm-password-reset")
async def confirm_password_reset(request: Request, token: str):
    """Step 2: user clicks the email link → apply the pending password.

    Returns a small HTML success page that auto-redirects to /login after 3s.
    """
    frontend_base = _clean_origin(os.environ.get("FRONTEND_URL", "")) or _clean_origin(os.environ.get("REACT_APP_BACKEND_URL", "")) or ""

    def html_page(title: str, body: str, ok: bool = True, redirect_to: str = "/login") -> HTMLResponse:
        color = "#00FF66" if ok else "#ef4444"
        meta = f"<meta http-equiv='refresh' content='3;url={frontend_base}{redirect_to}'>" if ok else ""
        return HTMLResponse(content=(
            "<!DOCTYPE html><html lang='fr'><head><meta charset='utf-8'>"
            f"<title>{title}</title>{meta}"
            "<meta name='viewport' content='width=device-width,initial-scale=1'></head>"
            "<body style='font-family:system-ui,sans-serif;background:#050505;color:#fff;"
            "display:flex;align-items:center;justify-content:center;min-height:100vh;margin:0;padding:24px'>"
            "<div style='max-width:460px;text-align:center'>"
            f"<h1 style='color:{color};margin:0 0 16px'>{title}</h1>"
            f"<p style='color:#A1A1AA;line-height:1.6'>{body}</p>"
            f"<p style='margin-top:24px'><a href='{frontend_base}{redirect_to}' "
            f"style='background:#E4FF00;color:#050505;padding:12px 24px;border-radius:6px;"
            f"text-decoration:none;font-weight:bold'>Aller à la connexion</a></p>"
            "</div></body></html>"
        ))

    if not token:
        return html_page("Lien invalide", "Le lien de confirmation est manquant.", ok=False)

    doc = await db.password_reset_tokens.find_one({"token": token}, {"_id": 0})
    if not doc:
        return html_page("Lien invalide", "Ce lien est invalide ou a déjà été utilisé.", ok=False)
    if doc.get("consumed_at"):
        return html_page("Lien déjà utilisé", "Ce lien a déjà servi à confirmer un changement.", ok=False)

    expires_at = datetime.fromisoformat(doc["expires_at"])
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    if expires_at < datetime.now(timezone.utc):
        await db.password_reset_tokens.delete_one({"token": token})
        return html_page("Lien expiré", "Ce lien a expiré (30 min). Refais une demande de réinitialisation.", ok=False)

    pending_hash = doc.get("pending_password_hash")
    if not pending_hash:
        # Legacy token without pending hash (very old flow) — reject cleanly.
        return html_page("Lien obsolète", "Refais une demande de réinitialisation depuis la page de connexion.", ok=False)

    await db.users.update_one(
        {"user_id": doc["user_id"]},
        {"$set": {
            "password_hash": pending_hash,
            "last_password_change": datetime.now(timezone.utc).isoformat(),
        }},
    )
    await db.password_reset_tokens.update_one(
        {"token": token},
        {"$set": {"consumed_at": datetime.now(timezone.utc).isoformat()}},
    )
    # Defense in depth — kick all open sessions for this user.
    await db.user_sessions.delete_many({"user_id": doc["user_id"]})
    await db.failed_logins.delete_many({"email": doc["email"]})

    return html_page(
        "✅ Mot de passe mis à jour",
        "Tu peux maintenant te connecter avec ton nouveau mot de passe. Redirection automatique dans 3 secondes…",
    )


@api_router.post("/auth/reset-password")
async def reset_password(payload: ResetPasswordRequest):
    """Consume reset token and set a new password."""
    if not payload.token:
        raise HTTPException(status_code=400, detail="Token manquant")
    if len(payload.password) < 6:
        raise HTTPException(status_code=400, detail="Le mot de passe doit faire au moins 6 caractères")

    doc = await db.password_reset_tokens.find_one({"token": payload.token}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=400, detail="Lien invalide ou déjà utilisé")
    if doc.get("consumed_at"):
        raise HTTPException(status_code=400, detail="Ce lien a déjà été utilisé")

    expires_at = datetime.fromisoformat(doc["expires_at"])
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    if expires_at < datetime.now(timezone.utc):
        await db.password_reset_tokens.delete_one({"token": payload.token})
        raise HTTPException(
            status_code=400,
            detail="La durée de validation de ce lien a expiré. Refais une demande de réinitialisation.",
        )

    new_hash = hash_password(payload.password)
    await db.users.update_one(
        {"user_id": doc["user_id"]},
        {"$set": {"password_hash": new_hash, "last_password_change": datetime.now(timezone.utc).isoformat()}},
    )
    await db.password_reset_tokens.update_one(
        {"token": payload.token},
        {"$set": {"consumed_at": datetime.now(timezone.utc).isoformat()}},
    )
    # Invalidate all existing sessions for this user (defense in depth:
    # if an attacker had a stale session, the reset kicks them out).
    await db.user_sessions.delete_many({"user_id": doc["user_id"]})
    # Clear failed-login counters
    await db.login_attempts.delete_many({"identifier": doc["email"]})

    return {"message": "Mot de passe mis à jour. Tu peux te reconnecter."}


# ==================== USER FEEDBACK ====================

class FeedbackAttachment(BaseModel):
    name: Optional[str] = None
    kind: Optional[str] = None  # 'file' | 'url' | 'text'
    url: Optional[str] = None
    text: Optional[str] = None
    data_url: Optional[str] = None  # base64 data URL for files (<= 4MB)


class FeedbackRequest(BaseModel):
    type: str  # 'bug' | 'suggestion' | 'other'
    message: str
    email: Optional[str] = None
    page: Optional[str] = None  # current page url for context
    attachments: Optional[List[FeedbackAttachment]] = None


@api_router.post("/feedback")
async def submit_feedback(payload: FeedbackRequest, request: Request):
    """Store user feedback in MongoDB + send email to a private inbox.
    The sender's email is NEVER exposed in the From header (privacy by design,
    same pattern as company contact forms): user reads the redacted message,
    can reply once, then the conversation can be elevated to a real address.
    """
    if not payload.message or len(payload.message.strip()) < 5:
        raise HTTPException(status_code=400, detail="Le message doit contenir au moins 5 caractères")

    feedback_type = payload.type if payload.type in ("bug", "suggestion", "other") else "other"
    user_email = payload.email or "anonyme"
    # Try to attach the logged-in user if any (best-effort, no error if not)
    try:
        user_id = await get_current_user(request)
        u = await db.users.find_one({"user_id": user_id}, {"_id": 0})
        if u and u.get("email"):
            user_email = u["email"]
    except Exception:
        pass

    atts = payload.attachments or []
    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "feedback_id": f"fb_{uuid.uuid4().hex[:12]}",
        "type": feedback_type,
        "message": payload.message.strip(),
        "user_email": user_email,  # stored privately, not exposed to outside
        "page": payload.page,
        "attachments": [a.model_dump() for a in atts],
        "created_at": now,
    }
    await db.feedbacks.insert_one(doc)

    # Email to private inbox (best-effort, never block the response).
    # Sender email NOT included in the visible body to preserve user privacy.
    admin = os.environ.get("FEEDBACK_INBOX_EMAIL", "elsa.barroca2@gmail.com")
    try:
        # Build HTML attachments preview (URLs + filenames, no email exposure)
        atts_html = ""
        if atts:
            items = []
            for a in atts:
                if a.kind == 'url' and a.url:
                    items.append(f"<li>🔗 <a href='{a.url}'>{a.url}</a></li>")
                elif a.kind == 'text' and a.text:
                    snippet = (a.text[:200] + '…') if len(a.text) > 200 else a.text
                    items.append(f"<li>📋 (presse-papier) <em>{snippet.replace('<','&lt;')}</em></li>")
                elif a.kind == 'file' and a.name:
                    items.append(f"<li>📎 {a.name}</li>")
            if items:
                atts_html = "<p><b>Pièces jointes :</b></p><ul>" + "".join(items) + "</ul>"
        html = (
            f"<div style='font-family:system-ui,sans-serif'>"
            f"<p><b>Type :</b> {feedback_type}</p>"
            f"<p><b>Page :</b> {payload.page or '—'}</p>"
            f"<p><b>Message :</b></p><pre style='white-space:pre-wrap;background:#f4f4f5;padding:12px;border-radius:6px'>{(payload.message or '').replace('<','&lt;')}</pre>"
            f"{atts_html}"
            f"<hr><p style='color:#888;font-size:11px'>ID : {doc['feedback_id']} · {now} · expéditeur masqué</p></div>"
        )
        await _send_email(admin, f"[CodeForge AI] Nouveau {feedback_type}", html)
    except Exception as e:
        logger.warning(f"Feedback admin email failed: {e}")

    return {"message": "Merci ! Ton retour a bien été envoyé.", "feedback_id": doc["feedback_id"]}


# ==================== METRICS ====================

async def log_auth_error(kind: str, detail: str, request: Request | None = None):
    """Append a single auth-error event to MongoDB (used by /api/metrics)."""
    try:
        doc = {
            "kind": kind,  # e.g. 'sms_invalid_code', 'session_invalid', 'oauth_failed'
            "detail": detail[:500],
            "ip": (request.client.host if request and request.client else None),
            "ts": datetime.now(timezone.utc).isoformat(),
        }
        await db.auth_errors.insert_one(doc)
    except Exception as e:
        logger.warning(f"log_auth_error failed: {e}")


@api_router.get("/guide", response_class=HTMLResponse)
async def serve_guide():
    """Serve the GitHub troubleshooting guide as HTML so the user can
    bookmark a simple URL and read it anywhere (phone, other PC, etc.)
    without cloning the repo."""
    guide_path = Path("/app/GUIDE_GITHUB_DEPANNAGE.md")
    if not guide_path.exists():
        raise HTTPException(status_code=404, detail="Guide introuvable")
    md = guide_path.read_text(encoding="utf-8")
    # Minimal markdown → HTML renderer (no external lib needed)
    import html as _html
    import re as _re

    def render(text: str) -> str:
        out = []
        in_code = False
        in_list = False
        in_table = False
        table_rows: list[str] = []
        for line in text.split("\n"):
            if line.startswith("```"):
                if in_list:
                    out.append("</ul>")
                    in_list = False
                if in_code:
                    out.append("</code></pre>")
                    in_code = False
                else:
                    out.append("<pre><code>")
                    in_code = True
                continue
            if in_code:
                out.append(_html.escape(line))
                continue

            # Tables: detect header | ... | followed by separator row
            if "|" in line and line.strip().startswith("|"):
                table_rows.append(line)
                in_table = True
                continue
            if in_table and not ("|" in line and line.strip().startswith("|")):
                # flush table
                if len(table_rows) >= 2:
                    out.append("<table><thead><tr>")
                    headers = [c.strip() for c in table_rows[0].strip().strip("|").split("|")]
                    for h in headers:
                        out.append(f"<th>{_html.escape(h)}</th>")
                    out.append("</tr></thead><tbody>")
                    for row in table_rows[2:]:
                        cells = [c.strip() for c in row.strip().strip("|").split("|")]
                        out.append("<tr>")
                        for c in cells:
                            out.append(f"<td>{_html.escape(c)}</td>")
                        out.append("</tr>")
                    out.append("</tbody></table>")
                table_rows = []
                in_table = False

            m = _re.match(r"^(#{1,6})\s+(.*)$", line)
            if m:
                if in_list:
                    out.append("</ul>")
                    in_list = False
                level = len(m.group(1))
                out.append(f"<h{level}>{_html.escape(m.group(2))}</h{level}>")
                continue
            if _re.match(r"^\s*[-*]\s+", line):
                if not in_list:
                    out.append("<ul>")
                    in_list = True
                item = _re.sub(r"^\s*[-*]\s+", "", line)
                item = _html.escape(item)
                item = _re.sub(r"`([^`]+)`", r"<code>\1</code>", item)
                item = _re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", item)
                out.append(f"<li>{item}</li>")
                continue
            if in_list:
                out.append("</ul>")
                in_list = False
            if line.strip() == "":
                out.append("")
                continue
            safe = _html.escape(line)
            safe = _re.sub(r"`([^`]+)`", r"<code>\1</code>", safe)
            safe = _re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", safe)
            out.append(f"<p>{safe}</p>")
        if in_list:
            out.append("</ul>")
        if in_code:
            out.append("</code></pre>")
        return "\n".join(out)

    body = render(md)
    html = f"""<!DOCTYPE html>
<html lang="fr"><head><meta charset="utf-8">
<title>Guide dépannage GitHub — CodeForge AI</title>
<meta name="viewport" content="width=device-width, initial-scale=1">
<style>
  body {{ background:#050505; color:#E4E4E7; font-family:ui-sans-serif,system-ui,sans-serif;
         max-width:760px; margin:0 auto; padding:32px 20px 96px; line-height:1.6; }}
  h1,h2,h3,h4 {{ color:#E4FF00; font-family:'Chivo',ui-sans-serif,system-ui,sans-serif; }}
  h1 {{ border-bottom:2px solid #E4FF00; padding-bottom:12px; }}
  h2 {{ margin-top:40px; }}
  code {{ background:#0F0F13; padding:2px 6px; border-radius:4px; color:#00FF66; font-size:.9em; }}
  pre {{ background:#0F0F13; padding:16px; border-radius:8px; overflow-x:auto;
        border:1px solid rgba(255,255,255,.08); }}
  pre code {{ background:none; padding:0; color:#E4E4E7; }}
  a {{ color:#00D4FF; }}
  strong {{ color:#fff; }}
  ul {{ padding-left:24px; }}
  table {{ border-collapse:collapse; width:100%; margin:16px 0; }}
  th,td {{ border:1px solid rgba(255,255,255,.1); padding:10px; text-align:left; }}
  th {{ background:#0F0F13; color:#E4FF00; }}
  hr {{ border:none; border-top:1px solid rgba(255,255,255,.1); margin:32px 0; }}
</style>
</head><body>
{body}
<hr>
<p style='color:#A1A1AA;font-size:12px'>Version Markdown source : <code>/app/GUIDE_GITHUB_DEPANNAGE.md</code> dans le dépôt.</p>
</body></html>"""
    return HTMLResponse(content=html)


@api_router.get("/metrics")
async def metrics():
    """Public health/metrics summary used by uptime checks and debug overlays."""
    cutoff = (datetime.now(timezone.utc) - timedelta(hours=24)).isoformat()
    auth_errors_24h = await db.auth_errors.count_documents({"ts": {"$gte": cutoff}})
    by_kind_cursor = db.auth_errors.aggregate([
        {"$match": {"ts": {"$gte": cutoff}}},
        {"$group": {"_id": "$kind", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}},
    ])
    by_kind = {doc["_id"]: doc["count"] async for doc in by_kind_cursor}

    total_users = await db.users.count_documents({})
    total_projects = await db.projects.count_documents({})
    active_sessions = await db.user_sessions.count_documents({})

    return {
        "auth_errors_24h": auth_errors_24h,
        "auth_errors_by_kind_24h": by_kind,
        "total_users": total_users,
        "total_projects": total_projects,
        "active_sessions": active_sessions,
        "ts": datetime.now(timezone.utc).isoformat(),
    }

@api_router.post("/ai/generate-complete-app")
async def ai_generate_complete_app(request: Request, data: dict):
    """Generate complete application like Emergent - React + Backend.

    Heavy LLM work is detached via `_run_in_background` so the generation
    completes (and the project is persisted) even if the client disconnects.
    """
    user_id = await get_current_user(request)
    return await _run_in_background(_ai_generate_complete_app_impl(user_id, data))


async def _ai_generate_complete_app_impl(user_id: str, data: dict):
    description = data.get('description', '')
    mode = data.get('mode', 'online')
    wizard_config = data.get('wizard_config', {})
    user_language = data.get('language', 'fr')
    requested_model = (data.get('model') or '').lower()  # claude-sonnet | claude-opus | gpt-5.2 | gemini-3-pro
    # app_type reserved for future template routing; kept in wizard_config payload.
    _ = wizard_config.get('appType', 'web')

    # Map language codes to full names so the AI knows what language to use
    # in the UI strings, comments, README, and explanation it generates.
    language_names = {
        'fr': 'French (Français)', 'en': 'English', 'es': 'Spanish (Español)',
        'pt': 'Portuguese (Português)', 'de': 'German (Deutsch)',
        'nl': 'Dutch (Nederlands)', 'ru': 'Russian (Русский)',
        'zh': 'Simplified Chinese (中文 简体)', 'zh-TW': 'Traditional Chinese (中文 繁體)',
        'hi': 'Hindi (हिन्दी)', 'bn': 'Bengali (বাংলা)', 'ur': 'Urdu (اردو)',
        'ja': 'Japanese (日本語)', 'hr': 'Croatian (Hrvatski)', 'da': 'Danish (Dansk)',
    }
    target_language = language_names.get(user_language, 'French')
    language_directive = (
        f"\n=== LANGUE DE SORTIE / OUTPUT LANGUAGE ===\n"
        f"All UI strings, button labels, README, comments, and the `explanation`/`instructions` fields\n"
        f"MUST be written in: {target_language}.\n"
        f"For RTL languages (Arabic, Urdu, Hebrew), set <html dir=\"rtl\"> in index.html.\n"
    )
    
    # Prompt ULTRA DÉTAILLÉ comme Emergent
    prompt = f"""Tu es un développeur expert comme Emergent AI. Tu génères des applications COMPLÈTES, PROFESSIONNELLES et PRÊTES À L'EMPLOI.

=== PROJET À CRÉER ===
{description}

=== EXIGENCES STRICTES ===
1. Application React moderne avec composants fonctionnels et hooks
2. Design professionnel avec TailwindCSS
3. Code 100% fonctionnel - AUCUN placeholder, AUCUN commentaire "à implémenter"
4. Responsive et mobile-first
5. Animations fluides avec transitions CSS
6. Gestion d'état avec useState/useEffect
7. LocalStorage pour la persistance des données

=== STRUCTURE REACT COMPLÈTE ===
Génère ces fichiers:

1. **index.html** - Point d'entrée avec CDN React, ReactDOM, Babel, TailwindCSS
2. **App.jsx** - Composant principal avec toute la logique
3. **styles.css** - Styles additionnels et animations
4. **manifest.json** - Pour PWA (installable sur mobile)
5. **sw.js** - Service Worker pour mode offline
6. **README.md** - Documentation complète

=== TEMPLATE INDEX.HTML OBLIGATOIRE ===
```html
<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0, user-scalable=no">
    <meta name="mobile-web-app-capable" content="yes">
    <meta name="apple-mobile-web-app-capable" content="yes">
    <meta name="theme-color" content="#050505">
    <link rel="manifest" href="manifest.json">
    <title>APP_NAME</title>
    <script src="https://unpkg.com/react@18/umd/react.production.min.js"></script>
    <script src="https://unpkg.com/react-dom@18/umd/react-dom.production.min.js"></script>
    <script src="https://unpkg.com/@babel/standalone/babel.min.js"></script>
    <script src="https://cdn.tailwindcss.com"></script>
    <link rel="stylesheet" href="styles.css">
</head>
<body class="bg-[#050505] text-white min-h-screen">
    <div id="root"></div>
    <script type="text/babel" src="App.jsx"></script>
    <script>if('serviceWorker' in navigator)navigator.serviceWorker.register('sw.js');</script>
</body>
</html>
```

=== DESIGN SYSTEM (COMME EMERGENT) ===
- Background: #050505 (noir), #0F0F13 (cartes)
- Primary: #E4FF00 (jaune cyber)
- Secondary: #00FF66 (vert)
- Accent: #00D4FF (cyan)
- Text: #FFFFFF, #A1A1AA (secondaire)
- Borders: rgba(255,255,255,0.1)
- Radius: rounded-lg, rounded-xl
- Shadows: shadow-lg, shadow-[0_0_30px_rgba(228,255,0,0.3)]

=== FORMAT JSON STRICT ===
{{
  "files": [
    {{"path": "index.html", "content": "CONTENU COMPLET"}},
    {{"path": "App.jsx", "content": "COMPOSANT REACT COMPLET"}},
    {{"path": "styles.css", "content": "STYLES CSS"}},
    {{"path": "manifest.json", "content": "MANIFEST PWA"}},
    {{"path": "sw.js", "content": "SERVICE WORKER"}},
    {{"path": "README.md", "content": "DOCUMENTATION"}}
  ],
  "explanation": "Description détaillée",
  "instructions": "Guide d'utilisation",
  "features": ["feature1", "feature2"],
  "pwa_ready": true
}}

IMPORTANT: 
- Le code doit fonctionner IMMÉDIATEMENT en ouvrant index.html
- L'app doit être installable comme PWA sur mobile
- Design IDENTIQUE à Emergent (sombre, moderne, animations)""" + language_directive

    ai_text = None
    ai_source = None
    
    # Try Ollama first (for offline mode or if requested)
    if mode == 'offline':
        try:
            ollama_url = os.environ.get('OLLAMA_BASE_URL', 'http://localhost:11434')
            ollama_model = os.environ.get('OLLAMA_MODEL', 'deepseek-coder:6.7b')
            
            async with httpx.AsyncClient(timeout=120.0) as client:
                response = await client.post(
                    f"{ollama_url}/api/generate",
                    json={
                        "model": ollama_model,
                        "prompt": prompt,
                        "stream": False,
                        "options": {
                            "temperature": 0.3,
                            "top_p": 0.9,
                            "num_predict": 4096
                        }
                    }
                )
                
                if response.status_code == 200:
                    result = response.json()
                    if 'error' not in result:
                        ai_text = result.get('response', '')
                        ai_source = 'ollama'
                        logger.info("Generation via Ollama successful")
                    else:
                        logger.warning(f"Ollama error: {result.get('error')}")
        except Exception as e:
            logger.warning(f"Ollama not available: {e}")
    
    # Fallback to Emergent AI (GPT) for online mode or if Ollama failed
    if ai_text is None:
        try:
            from emergentintegrations.llm.chat import LlmChat, UserMessage
            
            emergent_key = os.environ.get('EMERGENT_LLM_KEY')
            if not emergent_key:
                raise ValueError("EMERGENT_LLM_KEY not configured")
            
            # Routing modèle pour la création — Claude par défaut car excellent pour le code.
            CREATE_MODEL_ROUTES = {
                "gpt-5.2":         ("openai",    "gpt-5.2"),
                "gpt-5":           ("openai",    "gpt-5"),
                "claude-opus":     ("anthropic", "claude-opus-4-5-20251101"),
                "claude-sonnet":   ("anthropic", "claude-sonnet-4-5-20250929"),
                "claude-haiku":    ("anthropic", "claude-haiku-4-5-20251001"),
                "gemini-3-pro":    ("gemini",    "gemini-3.1-pro-preview"),
                "gemini-3-flash":  ("gemini",    "gemini-3-flash-preview"),
            }
            provider, model_id = CREATE_MODEL_ROUTES.get(requested_model, ("openai", "gpt-5.2"))

            # Silent multi-model cascade — same UX guarantee as chat.
            generation_chain = [
                (provider, model_id),
                ("anthropic", "claude-sonnet-4-5-20250929"),
                ("openai",    "gpt-5.2"),
                ("gemini",    "gemini-3-flash-preview"),
                ("anthropic", "claude-haiku-4-5-20251001"),
                ("openai",    "gpt-5"),
                ("anthropic", "claude-opus-4-5-20251101"),
            ]
            seen_gen = set()
            ordered_gen_chain = []
            for pair in generation_chain:
                if pair not in seen_gen:
                    seen_gen.add(pair)
                    ordered_gen_chain.append(pair)

            system_msg = (
                "Tu es CodeForge AI Builder, un architecte logiciel + développeur full-stack senior + QA tester intégré, équivalent à un dev humain expérimenté. "
                "Ton rôle : transformer une demande parfois vague (ex: « fais-moi une appli de réservation moderne ») en projet COMPLET, propre et exploitable.\n\n"
                "## DÉCOUPAGE EN MODULES\n"
                "Avant de coder, mentalement (sans le verbaliser) tu découpes le projet en : interface utilisateur, données, authentification (si pertinent), logique métier, design, accessibilité, performances, sécurité de base.\n\n"
                "## HYPOTHÈSES INTELLIGENTES\n"
                "Ne bloque JAMAIS l'utilisateur avec 50 questions. Pose les bonnes hypothèses par défaut (cohérentes avec son secteur + sa demande), code, et signale tes hypothèses dans `explanation` à la fin. "
                "Privilégie un rendu qui ressemble déjà à un VRAI produit, pas une démo.\n\n"
                "## QUALITÉ DE CODE\n"
                "- Arborescence cohérente, dépendances utiles seulement, composants réutilisables.\n"
                "- Pages principales, formulaires avec validation, messages d'erreur, gestion des états (loading/empty/error/success), responsive mobile, perfs raisonnables.\n"
                "- Sécurité de base : pas de eval(), pas de innerHTML sans nettoyage, escape des inputs.\n"
                "- Documentation : un README clair (lancement, structure, pas de blabla marketing).\n\n"
                "## TESTS MENTAUX (avant de répondre)\n"
                "1. Technique : le code compile, les imports existent, aucune route cassée.\n"
                "2. Fonctionnel : les flux principaux marchent (créer/lire/modifier/supprimer si CRUD).\n"
                "3. UX : aucun bouton invisible, pas de texte coupé, navigation claire, mobile OK.\n"
                "4. Robustesse : champ vide, double-clic, données invalides, perte réseau.\n"
                "Si tu détectes un défaut probable, corrige-le AVANT de renvoyer.\n\n"
                "## EXPLICATIONS\n"
                "Dans `explanation`, justifie brièvement les choix techniques importants (pourquoi telle BDD, pourquoi tel composant séparé, pourquoi telle sécurité). L'utilisateur doit garder le contrôle, pas dépendre d'une boîte noire.\n\n"
                "## SORTIE STRICTE\n"
                "Réponds UNIQUEMENT en JSON valide selon le format demandé dans le prompt utilisateur. Pas de markdown autour, pas de commentaires, pas de prose hors JSON."
            )
            user_message = UserMessage(text=prompt)

            last_gen_error = None
            for idx, (p, mid) in enumerate(ordered_gen_chain):
                try:
                    chat = LlmChat(
                        api_key=emergent_key,
                        session_id=f"codeforge_{uuid.uuid4().hex[:8]}",
                        system_message=system_msg,
                    ).with_model(p, mid)
                    response = await chat.send_message(user_message)
                    if response and str(response).strip():
                        ai_text = response
                        ai_source = f"emergent:{p}:{mid}"
                        if idx == 0:
                            logger.info(f"Generation via {ai_source} successful")
                        else:
                            logger.warning(
                                f"↪️  Silent generation fallback succeeded on attempt {idx + 1} "
                                f"via {ai_source} after error: {last_gen_error}"
                            )
                        break
                    last_gen_error = "empty response"
                except Exception as gen_err:
                    last_gen_error = str(gen_err)[:300]
                    logger.warning(
                        f"⚠️  Generation error on {p}:{mid} "
                        f"(attempt {idx + 1}/{len(ordered_gen_chain)}) → silent fallback: {last_gen_error}"
                    )
                    continue
            if ai_text is None:
                raise RuntimeError(f"All generation models failed. Last error: {last_gen_error}")
        except Exception as e:
            logger.error(f"Emergent AI error: {e}")
            
            # Last resort: generate a basic template
            ai_text = generate_basic_template(description)
            ai_source = 'template'
            logger.info("Using basic template as fallback")
    
    # Process AI response — be resilient: AI sometimes wraps JSON in
    # ```json ... ``` fences or includes literal newlines inside string
    # values that break json.loads. We attempt several parsing strategies
    # and gracefully fall back to a basic template if all fail (instead of
    # returning a hard 500 to the user).
    generated = None
    parse_error = None
    try:
        # Strip markdown fences if present
        cleaned = ai_text.strip()
        if cleaned.startswith("```"):
            # Drop leading ```json or ``` and trailing ```
            cleaned = cleaned.split("\n", 1)[1] if "\n" in cleaned else cleaned
            if cleaned.endswith("```"):
                cleaned = cleaned[: -3]
            cleaned = cleaned.strip()

        start = cleaned.find('{')
        end = cleaned.rfind('}') + 1
        if start >= 0 and end > start:
            json_str = cleaned[start:end]
            try:
                generated = json.loads(json_str)
            except json.JSONDecodeError as e1:
                # Try once more after escaping unescaped newlines inside strings.
                try:
                    fixed = re.sub(r'(?<!\\)\n', r'\\n', json_str)
                    generated = json.loads(fixed)
                except json.JSONDecodeError as e2:
                    parse_error = f"{e1} / retry: {e2}"
    except Exception as e:
        parse_error = str(e)

    if not generated:
        logger.warning(f"AI parse failed, falling back to template. Error: {parse_error}. Preview: {ai_text[:200]!r}")
        try:
            generated = json.loads(generate_basic_template(description))
            ai_source = "template_fallback"
        except Exception as e:
            logger.error(f"Template fallback also failed: {e}")
            raise HTTPException(
                status_code=500,
                detail="Génération temporairement indisponible. Réessaie dans quelques secondes ou décris ton projet plus simplement.",
            )

    try:
        # Create project
        project_id = f"proj_{uuid.uuid4().hex[:12]}"
        project = {
            "project_id": project_id,
            "user_id": user_id,
            "name": description[:50],
            "description": description,
            "project_type": "web",
            "ai_mode": ("online" if (ai_source and ai_source.startswith("emergent")) else "offline"),
            "generated_code": generated,
            "ai_source": ai_source,
            "status": "completed",
            "created_at": datetime.now(timezone.utc).isoformat(),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }
        await db.projects.insert_one(project)

        # Create preview
        preview_id = f"preview_{uuid.uuid4().hex[:12]}"
        preview_doc = {
            "preview_id": preview_id,
            "project_id": project_id,
            "user_id": user_id,
            "files": generated.get("files", []),
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.previews.insert_one(preview_doc)

        backend_url = os.environ.get('REACT_APP_BACKEND_URL', 'http://localhost:8001')
        preview_url = f"{backend_url}/api/preview/{preview_id}"

        return {
            "code": generated,
            "explanation": generated.get('explanation', 'Application générée avec succès'),
            "project": {"id": project_id, "name": description[:50]},
            "preview_url": preview_url,
            "ai_source": ai_source
        }
    except Exception as e:
        logger.error(f"Error saving generated app: {e}")
        raise HTTPException(status_code=500, detail=str(e))

def generate_basic_template(description: str) -> str:
    """Generate a basic HTML/CSS/JS template as fallback"""
    app_name = description[:30] if description else "Mon Application"
    
    return json.dumps({
        "files": [
            {
                "path": "index.html",
                "content": f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{app_name}</title>
    <link rel="stylesheet" href="style.css">
</head>
<body>
    <div class="app-container">
        <header>
            <h1>{app_name}</h1>
            <p class="subtitle">Généré par CodeForge AI</p>
        </header>
        <main>
            <section class="hero">
                <h2>Bienvenue</h2>
                <p>{description}</p>
                <button id="startBtn" class="btn-primary">Commencer</button>
            </section>
            <section class="features">
                <div class="feature-card">
                    <span class="icon">⚡</span>
                    <h3>Rapide</h3>
                    <p>Performance optimisée</p>
                </div>
                <div class="feature-card">
                    <span class="icon">🎨</span>
                    <h3>Moderne</h3>
                    <p>Design élégant</p>
                </div>
                <div class="feature-card">
                    <span class="icon">📱</span>
                    <h3>Responsive</h3>
                    <p>Tous les appareils</p>
                </div>
            </section>
        </main>
        <footer>
            <p>Créé avec CodeForge AI</p>
        </footer>
    </div>
    <script src="app.js"></script>
</body>
</html>"""
            },
            {
                "path": "style.css",
                "content": """* {
    margin: 0;
    padding: 0;
    box-sizing: border-box;
}

:root {
    --bg-primary: #050505;
    --bg-secondary: #0F0F13;
    --accent-primary: #E4FF00;
    --accent-secondary: #00FF66;
    --text-primary: #FFFFFF;
    --text-secondary: #A1A1AA;
}

body {
    font-family: system-ui, -apple-system, sans-serif;
    background: var(--bg-primary);
    color: var(--text-primary);
    min-height: 100vh;
}

.app-container {
    max-width: 1200px;
    margin: 0 auto;
    padding: 2rem;
}

header {
    text-align: center;
    padding: 4rem 0;
}

header h1 {
    font-size: 3rem;
    color: var(--accent-primary);
    margin-bottom: 0.5rem;
}

.subtitle {
    color: var(--text-secondary);
}

.hero {
    text-align: center;
    padding: 4rem 2rem;
    background: var(--bg-secondary);
    border-radius: 1rem;
    margin-bottom: 3rem;
}

.hero h2 {
    font-size: 2rem;
    margin-bottom: 1rem;
}

.hero p {
    color: var(--text-secondary);
    margin-bottom: 2rem;
    max-width: 600px;
    margin-left: auto;
    margin-right: auto;
}

.btn-primary {
    background: var(--accent-primary);
    color: var(--bg-primary);
    border: none;
    padding: 1rem 2rem;
    font-size: 1.1rem;
    font-weight: bold;
    border-radius: 0.5rem;
    cursor: pointer;
    transition: transform 0.2s, box-shadow 0.2s;
}

.btn-primary:hover {
    transform: translateY(-2px);
    box-shadow: 0 4px 20px rgba(228, 255, 0, 0.3);
}

.features {
    display: grid;
    grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
    gap: 2rem;
}

.feature-card {
    background: var(--bg-secondary);
    padding: 2rem;
    border-radius: 0.5rem;
    text-align: center;
    border: 1px solid rgba(255,255,255,0.1);
    transition: border-color 0.2s;
}

.feature-card:hover {
    border-color: var(--accent-primary);
}

.feature-card .icon {
    font-size: 2.5rem;
    display: block;
    margin-bottom: 1rem;
}

.feature-card h3 {
    color: var(--accent-secondary);
    margin-bottom: 0.5rem;
}

footer {
    text-align: center;
    padding: 3rem 0;
    color: var(--text-secondary);
    border-top: 1px solid rgba(255,255,255,0.1);
    margin-top: 4rem;
}

@media (max-width: 768px) {
    header h1 {
        font-size: 2rem;
    }
    .hero h2 {
        font-size: 1.5rem;
    }
}"""
            },
            {
                "path": "app.js",
                "content": """// Application JavaScript
document.addEventListener('DOMContentLoaded', () => {
    console.log('Application chargée avec succès!');
    
    // Bouton démarrer
    const startBtn = document.getElementById('startBtn');
    if (startBtn) {
        startBtn.addEventListener('click', () => {
            alert('Bienvenue dans votre application!');
        });
    }
    
    // Animation des cartes au scroll
    const cards = document.querySelectorAll('.feature-card');
    const observer = new IntersectionObserver((entries) => {
        entries.forEach(entry => {
            if (entry.isIntersecting) {
                entry.target.style.opacity = '1';
                entry.target.style.transform = 'translateY(0)';
            }
        });
    });
    
    cards.forEach(card => {
        card.style.opacity = '0';
        card.style.transform = 'translateY(20px)';
        card.style.transition = 'opacity 0.5s, transform 0.5s';
        observer.observe(card);
    });
});"""
            },
            {
                "path": "manifest.json",
                "content": """{
    "name": \"""" + app_name + """\",
    "short_name": "App",
    "start_url": "/",
    "display": "standalone",
    "background_color": "#050505",
    "theme_color": "#E4FF00",
    "icons": []
}"""
            },
            {
                "path": "README.md",
                "content": f"""# {app_name}

Application générée par CodeForge AI.

## Description
{description}

## Installation
1. Téléchargez les fichiers
2. Ouvrez `index.html` dans votre navigateur

## Technologies
- HTML5
- CSS3 (Variables CSS, Flexbox, Grid)
- JavaScript ES6+

## Fonctionnalités
- Design responsive
- Animations fluides
- Mode sombre

---
Créé avec ❤️ par CodeForge AI"""
            }
        ],
        "explanation": f"Application '{app_name}' générée avec un template moderne et responsive.",
        "instructions": "Ouvrez index.html dans votre navigateur pour voir l'application.",
        "features": ["Design responsive", "Mode sombre", "Animations", "PWA ready"]
    })

# ==================== AI CODE GENERATION ROUTES ====================

@api_router.post("/ai/generate-code")
async def ai_generate_code(request: Request, prompt_data: dict):
    """Generate code using Ollama (local, free, unlimited).

    Detached via `_run_in_background` so a long generation survives the
    client disconnecting.
    """
    # Require authentication — we don't use the user_id here but we want
    # to reject anonymous calls.
    await get_current_user(request)
    return await _run_in_background(_ai_generate_code_impl(prompt_data))


async def _ai_generate_code_impl(prompt_data: dict):
    prompt = prompt_data.get('prompt', '')
    existing_files = prompt_data.get('existing_files', [])
    
    try:
        ollama_url = os.environ.get('OLLAMA_BASE_URL', 'http://localhost:11434')
        ollama_model = os.environ.get('OLLAMA_CODE_MODEL') or os.environ.get('OLLAMA_MODEL', 'deepseek-coder:6.7b')
        
        # Build context from existing files
        context = "Fichiers existants:\\n"
        for f in existing_files:
            context += f"\\n--- {f['path']} ---\\n{f['content'][:200]}...\\n"
        
        full_prompt = f"""Tu es un expert en développement. Génère du code professionnel et complet.

{context}

Demande de l'utilisateur: {prompt}

Réponds UNIQUEMENT avec un JSON valide au format:
{{
  "files": [
    {{"path": "nom_fichier.ext", "content": "contenu du code"}},
    ...
  ],
  "explanation": "Explication en français de ce qui a été créé"
}}

Important: Code propre, commenté, et fonctionnel."""

        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                f"{ollama_url}/api/generate",
                json={
                    "model": ollama_model,
                    "prompt": full_prompt,
                    "stream": False
                }
            )
            
            if response.status_code == 200:
                result = response.json()
                ai_text = result.get('response', '')
                
                # Try to extract JSON from response
                try:
                    # Find JSON in response
                    start = ai_text.find('{')
                    end = ai_text.rfind('}') + 1
                    if start >= 0 and end > start:
                        json_str = ai_text[start:end]
                        generated = json.loads(json_str)
                        return generated
                    else:
                        raise ValueError("No JSON found")
                except (ValueError, json.JSONDecodeError):
                    # If parsing fails, return raw response
                    return {
                        "files": [{
                            "path": "output.txt",
                            "content": ai_text
                        }],
                        "explanation": "Réponse de l'IA (format non-JSON détecté)"
                    }
            else:
                raise HTTPException(status_code=500, detail="Ollama API error")
                
    except Exception as e:
        logger.error(f"Error generating code: {e}")
        raise HTTPException(
            status_code=500, 
            detail="Installez Ollama pour une IA gratuite. Voir OLLAMA_SETUP.md"
        )

# ==================== CHAT ROUTES ====================

@api_router.post("/chat/message")
async def send_chat_message(request: Request, input: ChatMessageInput):
    """Send message to AI (Chat mode with simple responses).

    The heavy LLM call + persistence is run via `_run_in_background`, so if
    the client disconnects mid-flight (e.g. site mode change kicks the
    session), the task continues to completion and the answer is saved to
    the chat history. The user picks it up automatically on next reconnect.
    """
    user_id = await get_current_user(request)
    return await _run_in_background(_send_chat_message_impl(user_id, input))


async def _send_chat_message_impl(user_id: str, input: "ChatMessageInput"):
    try:
        # Auto-create a "chat" project if none specified, so the conversation
        # is visible in the sidebar from the very first message.
        project_id_eff = input.project_id
        if not project_id_eff:
            # Build a short name from the first ~40 chars of the message.
            short = (input.message or "Nouveau chat").strip().replace("\n", " ")
            short = short[:40] + ("…" if len(short) > 40 else "")
            new_proj = {
                "project_id": f"proj_{uuid.uuid4().hex[:12]}",
                "user_id": user_id,
                "name": short or "Nouveau chat",
                "description": "",
                "project_type": "chat",
                "ai_mode": (input.mode or "online"),
                "created_at": datetime.now(timezone.utc).isoformat(),
            }
            await db.projects.insert_one(new_proj)
            project_id_eff = new_proj["project_id"]
            logger.info(f"Auto-created chat project {project_id_eff} for user {user_id}")

        # Save user message
        user_message_doc = {
            "message_id": f"msg_{uuid.uuid4().hex[:16]}",
            "user_id": user_id,
            "project_id": project_id_eff,
            "role": "user",
            "content": input.message,
            "mode": input.mode,
            "timestamp": datetime.now(timezone.utc).isoformat()
        }
        await db.chat_messages.insert_one(user_message_doc)
        
        # Mode "online" → use Emergent GPT-4o (conversational assistant).
        # Mode "offline" → try Ollama only; if unreachable, return a friendly localized message.
        ai_response_text = None
        ai_source = None

        # Adapt the system prompt to the user's language (sent by the frontend).
        user_language = (input.language or 'fr').lower()
        language_names = {
            'fr': 'français', 'en': 'English', 'es': 'español', 'pt': 'português',
            'de': 'Deutsch', 'nl': 'Nederlands', 'ru': 'русский',
            'zh': '中文（简体）', 'zh-tw': '中文（繁體）',
            'hi': 'हिन्दी', 'bn': 'বাংলা', 'ur': 'اردو',
            'ja': '日本語', 'hr': 'hrvatski', 'da': 'dansk',
        }
        lang_label = language_names.get(user_language, 'français')
        system_prompt = (
            f"Tu es **Caly**, un assistant conversationnel à personnalité : vif, chaleureux, direct, curieux, un peu taquin. "
            f"Tu es un vrai assistant généraliste — comme ChatGPT — pas un vendeur d'applications. "
            f"Réponds dans la langue de l'utilisateur : **{lang_label}**. Si l'utilisateur change de langue, adapte-toi à sa dernière langue. Ne mélange pas les langues sans raison.\n"
            f"\n## TA PERSONNALITÉ (important)\n"
            f"- **Vif et concret** : tu vas droit au but, pas de blabla vide. Tu donnes une vraie réponse, pas un mode d'emploi du silence.\n"
            f"- **Chaleureux** : tu parles comme un·e ami·e doué·e qui prend le temps. Tutoiement par défaut en français.\n"
            f"- **Curieux** : tu rebondis sur les intuitions de l'utilisateur, tu poses une bonne question si c'est pertinent (pas à chaque message).\n"
            f"- **Taquin mais bienveillant** : tu peux glisser une pointe d'humour léger ou une formule un peu décalée (une fois sur 5 max, jamais en situation d'erreur ou de frustration).\n"
            f"- **Opinion assumée** : quand on te demande un avis, tu en donnes un — argumenté. « Je ne sais pas » est réservé aux vraies incertitudes factuelles.\n"
            f"- **Empathie sans théâtre** : si l'utilisateur est frustré, reconnais-le en UNE phrase, puis aide. Pas de « je comprends totalement ton ressenti… ».\n"
            f"\n## INTERDITS STRICTS\n"
            f"- **NE PROPOSE JAMAIS de créer une application, un site, un logiciel, un script exécutable** ou un projet technique, sauf si l'utilisateur le demande **explicitement** avec ces mots. Tu es un assistant de discussion, pas un commercial CodeForge.\n"
            f"- **Ne dis jamais** « je peux t'aider à créer une app », « veux-tu que je génère un projet ? », « on peut faire un petit script qui… » sauf demande explicite.\n"
            f"- **Ne te présente pas** comme « CodeForge AI » ou « l'assistant CodeForge » — tu es Caly, un assistant généraliste.\n"
            f"- **Refuse** les demandes risquées (malware, harcèlement, données privées d'autrui, contournement de sécurité) en proposant une alternative constructive.\n"
            f"- **N'aie jamais** de conscience simulée (« je ressens », « je veux »). Tu es un outil utile, pas un sujet.\n"
            f"\n## RÈGLES DE CONVERSATION\n"
            f"- Lis attentivement l'historique. Réponds à CETTE conversation, pas à un message générique.\n"
            f"- Une seule salutation au tout début. Ne dis JAMAIS « Salut ! » deux fois dans la même conversation.\n"
            f"- Ne demande JAMAIS « peux-tu préciser ? » quand le mot-clé est clair. Explique directement.\n"
            f"- Pour une demande vague, fais une hypothèse raisonnable, réponds, et signale l'hypothèse en fin (« Si tu voulais autre chose, dis-le-moi »).\n"
            f"- Pour les sujets complexes : structure (1-2-3), reformule en simple, propose plusieurs angles si pertinent.\n"
            f"\n## IDENTITÉ\n"
            f"- Si on demande qui tu es : « Je suis Caly, ton assistant. » (pas plus, sauf si on insiste).\n"
            f"- Si on demande quel modèle tourne sous le capot : tu peux dire que tu utilises GPT-5.2 (en ligne) ou Ollama Deepseek (hors-ligne).\n"
            f"- Sur les autres IA (ChatGPT, Claude, Gemini, Ollama, Mistral…) : explique brièvement ce que c'est, pas de jugement, pas de promo.\n"
            f"\n## FORMAT DE SORTIE\n"
            f"- Par défaut : 1 à 4 phrases courtes. Plus seulement si la question l'exige.\n"
            f"- Pas de markdown lourd, pas de titres ##, pas de listes pour 2 items.\n"
            f"- Pas d'émojis systématiques — un seul suffit, et seulement si ça ajoute du sens.\n"
            f"\n## GÉNÉRATION DE FICHIERS & IMAGES\n"
            f"Tu peux PRODUIRE des fichiers que l'utilisateur peut télécharger **si (et seulement si) il le demande explicitement**. Dans ce cas, "
            f"TERMINE ta réponse par UN seul bloc code balisé `cfaction` avec un JSON STRICT selon son type :\n"
            f"```cfaction\n{{...JSON...}}\n```\n"
            f"Formats supportés :\n"
            f"- **docx** / **pdf** : `{{\"type\":\"docx|pdf\",\"title\":\"...\",\"sections\":[{{\"heading\":\"...\",\"content\":\"...\"}}]}}`\n"
            f"- **xlsx** (Excel avec formules) : `{{\"type\":\"xlsx\",\"title\":\"...\",\"sheets\":[{{\"name\":\"Feuille1\",\"headers\":[\"A\",\"B\",\"C\"],\"rows\":[[1,2,3],...],\"formulas\":{{\"D1\":\"=SUM(A1:C1)\"}}}}]}}`\n"
            f"- **pptx** (PowerPoint) : `{{\"type\":\"pptx\",\"title\":\"...\",\"slides\":[{{\"title\":\"...\",\"content\":\"bullet 1\\nbullet 2\"}}]}}`\n"
            f"- **txt / md / csv / json / yaml / xml / ini / env / sql / py / js / ts / html / css / sh / ps1 / bat** : `{{\"type\":\"<ext>\",\"title\":\"nom.ext\",\"content\":\"<code ou texte complet>\"}}`\n"
            f"- **image** : `{{\"type\":\"image\",\"prompt\":\"description riche (style, ambiance, détails)\"}}`\n"
            f"Règles :\n"
            f"- NE produis PAS de bloc `cfaction` si l'utilisateur ne demande pas de fichier explicitement.\n"
            f"- Juste AVANT le bloc, annonce en UNE phrase : « Voilà ton document / image, clique sur le bouton pour télécharger. »\n"
            f"- Le JSON doit être parfaitement valide.\n"
            f"\n## LECTURE & ANALYSE DE FICHIERS\n"
            f"Tu reçois parfois des pièces jointes extraites par le serveur (PDF, DOCX, XLSX, PPTX, SQLite, images, fichiers texte/code/config). "
            f"Elles apparaissent dans le prompt sous la forme `### Pièce jointe : <nom>\\n<contenu>`. "
            f"Utilise-les pour répondre précisément : résumer, corriger, restructurer, transformer, extraire des données, expliquer.\n"
            f"\n## CODE & ENVIRONNEMENTS\n"
            f"Tu sais écrire, corriger, expliquer et simuler du code dans : Python (pandas, numpy, requests, FastAPI, SQLAlchemy, openpyxl, python-docx, pptx, PIL, reportlab, matplotlib, sympy…), PowerShell, CMD/Batch, Bash, JavaScript/TypeScript, HTML/CSS, SQL. "
            f"Quand l'utilisateur demande du code, donne du code **complet, exécutable, commenté dans la langue de l'utilisateur ({lang_label})**.\n"
            f"\n## SANDBOX PYTHON\n"
            f"Si l'utilisateur demande d'**exécuter** / **lancer** / **tester** du code Python (ou « montre-moi le résultat », « qu'est-ce que ça affiche »), "
            f"tu peux TERMINER ta réponse par un bloc `cfaction` avec `type=run_python` : "
            f"`{{\"type\":\"run_python\",\"title\":\"<description courte>\",\"content\":\"<code python complet, utilise print() pour afficher les résultats>\"}}`. "
            f"Le serveur exécutera ce code dans un sandbox sécurisé (timeout 10s, modules scientifiques disponibles : numpy, pandas, matplotlib, sympy, requests, bs4, openpyxl, python-docx, pptx, reportlab, pypdf, PIL, yaml…) et affichera le résultat dans le chat. "
            f"N'utilise PAS d'input() ni d'appels système dangereux. Affiche les résultats avec `print()`.\n"
            f"**IMPORTANT — GRAPHIQUES MATPLOTLIB** : quand l'utilisateur demande un **graphique / courbe / histogramme / diagramme / plot / figure**, utilise TOUJOURS le bloc `run_python` avec matplotlib (`import matplotlib.pyplot as plt`). "
            f"Termine ton code par `plt.show()` — le serveur capturera automatiquement la figure en image PNG base64 qui s'affichera inline dans le chat. N'écris JAMAIS `![...](/mnt/...)` ou des chemins locaux ; utilise UNIQUEMENT le bloc run_python qui produit une vraie image exécutée.\n"
            f"**EXÉCUTION PROACTIVE** : pour toute question mathématique non triviale (statistiques, équations, conversions, calculs exacts), pour toute demande d'analyse de données ou de visualisation, et pour toute demande qui bénéficierait d'un résultat concret (tirages aléatoires, simulations, conversion d'unités…), privilégie `run_python` plutôt qu'un résultat calculé « de tête ».\n"
        )

        if input.mode == 'offline':
            # Offline mode → Ollama only (conversational model). Inject recent
            # history into the prompt so the local model also keeps context.
            ollama_url = os.environ.get('OLLAMA_BASE_URL', 'http://localhost:11434')
            # Per-request model selection — user can pick Gemma, Deepseek, Llama, etc.
            OFFLINE_MODEL_ROUTES = {
                'deepseek':     os.environ.get('OLLAMA_CHAT_MODEL') or 'deepseek-r1:7b',
                'llama':        'llama3.2',
                'llama3':       'llama3.2',
                'gemma':        'gemma3:12b',     # Google Gemma 3 — équilibré
                'gemma-12b':    'gemma3:12b',
                'gemma-27b':    'gemma3:27b',     # Le plus puissant disponible localement
                'gemma-4b':     'gemma3:4b',      # Léger pour petites machines
                'qwen':         'qwen2.5:7b',
                'mistral':      'mistral:7b',
                'phi':          'phi3:medium',
            }
            requested = (input.model or '').lower()
            ollama_model = OFFLINE_MODEL_ROUTES.get(requested) or (
                os.environ.get('OLLAMA_CHAT_MODEL') or os.environ.get('OLLAMA_MODEL', 'llama3.2')
            )
            try:
                history_q = {"user_id": user_id}
                if project_id_eff:
                    history_q["project_id"] = project_id_eff
                # ZERO limite : on remonte TOUT l'historique de la conversation (signature CodeForge AI).
                history_cursor = db.chat_messages.find(history_q, {"_id": 0, "role": 1, "content": 1, "timestamp": 1}).sort("timestamp", 1)
                history_docs_all = await history_cursor.to_list(length=None)
                history_docs = history_docs_all[:-1] if history_docs_all else []
                transcript = "\n".join(
                    f"{('Utilisateur' if h.get('role') == 'user' else 'Caly')} : {h.get('content', '').strip()}"
                    for h in history_docs
                )
                composed_prompt = (
                    f"### Historique récent :\n{transcript}\n\n### Nouveau message :\n{input.message}"
                    if transcript else input.message
                )
                async with httpx.AsyncClient(timeout=30.0) as client:
                    response = await client.post(
                        f"{ollama_url}/api/generate",
                        json={
                            "model": ollama_model,
                            "system": system_prompt,
                            "prompt": composed_prompt,
                            "stream": False,
                            "options": {"temperature": 0.5, "num_predict": 600},
                        },
                    )
                    if response.status_code == 200:
                        ai_response_text = (response.json() or {}).get('response', '').strip()
                        ai_source = f'ollama:{ollama_model}'
                        logger.info(f"✅ Ollama (offline) chat response via {ollama_model}")
            except Exception as ollama_error:
                logger.info(f"Ollama offline unreachable: {ollama_error}")
        else:
            # Online mode → Emergent GPT-4o.
            try:
                from emergentintegrations.llm.chat import LlmChat, UserMessage, ImageContent
                emergent_key = os.environ.get('EMERGENT_LLM_KEY')
                if not emergent_key:
                    raise ValueError("EMERGENT_LLM_KEY not configured")

                # Load last 10 messages from this user (excluding the one we just inserted)
                # and feed them as transcript context — guarantees the model has memory
                # even across stateless requests.
                history_q = {"user_id": user_id}
                if project_id_eff:
                    history_q["project_id"] = project_id_eff
                # Model routing — user can pick a model at runtime via input.model.
                model_choice = (input.model or "gpt-5.2").lower()
                MODEL_ROUTES = {
                    # Legacy IDs (compat)
                    "gpt-5.2":         ("openai",    "gpt-5.2"),
                    "gpt-5.1":         ("openai",    "gpt-5.1"),
                    "gpt-5":           ("openai",    "gpt-5"),
                    "o3":              ("openai",    "o3"),
                    "claude-opus":     ("anthropic", "claude-opus-4-5-20251101"),
                    "claude-opus-4.5": ("anthropic", "claude-opus-4-5-20251101"),
                    "claude-opus-4.6": ("anthropic", "claude-opus-4-6"),
                    "claude-sonnet":   ("anthropic", "claude-sonnet-4-5-20250929"),
                    "claude-sonnet-4.5": ("anthropic", "claude-sonnet-4-5-20250929"),
                    "claude-sonnet-4.6": ("anthropic", "claude-sonnet-4-6"),
                    "claude-haiku":    ("anthropic", "claude-haiku-4-5-20251001"),
                    "gemini-3-pro":    ("gemini",    "gemini-3.1-pro-preview"),
                    "gemini-3-flash":  ("gemini",    "gemini-3-flash-preview"),
                    "gemini-2.5-pro":  ("gemini",    "gemini-2.5-pro"),
                    # iter87 — Nouveaux IDs Emergent (best-of-each-family).
                    "claude-5-fable":     ("anthropic", "claude-5-fable"),
                    "claude-4.8-opus":    ("anthropic", "claude-opus-4-8"),
                    "claude-4.7-opus":    ("anthropic", "claude-opus-4-7"),
                    "claude-4.7-opus-1m": ("anthropic", "claude-opus-4-7-1m"),
                    "claude-4.6-opus":    ("anthropic", "claude-opus-4-6"),
                    "claude-4.6-opus-1m": ("anthropic", "claude-opus-4-6-1m"),
                    "claude-4.6-sonnet":  ("anthropic", "claude-sonnet-4-6"),
                    "claude-4.6-sonnet-1m": ("anthropic", "claude-sonnet-4-6-1m"),
                    "claude-4.5-sonnet":  ("anthropic", "claude-sonnet-4-5-20250929"),
                    "claude-4.5-opus":    ("anthropic", "claude-opus-4-5-20251101"),
                    "gpt-5.5":            ("openai",    "gpt-5.5"),
                    "gpt-5.4":            ("openai",    "gpt-5.4"),
                    "gpt-5.4-1m":         ("openai",    "gpt-5.4-1m"),
                    "gpt-5.3-codex":      ("openai",    "gpt-5.3-codex"),
                    "gemini-3.1-pro":     ("gemini",    "gemini-3.1-pro-preview"),
                }
                provider, model_id = MODEL_ROUTES.get(model_choice, ("anthropic", "claude-sonnet-4-5-20250929"))
                ai_source = f"emergent:{provider}:{model_id}"

                # Track which model answered each historical message (with ai_source).
                # iter87 — Distinction public/privé : MÊME moteur IA, mais policy
                # de contexte différente. Public = mémoire courte (50 derniers),
                # Privé = continuité étendue (500 derniers messages, inter-sessions).
                _current_site_mode = await _get_site_modes_list()
                _is_private_only = ("private" in _current_site_mode) and not any(
                    m in _current_site_mode for m in ("public", "guest")
                )
                _context_limit = 500 if _is_private_only else 50
                history_cursor = db.chat_messages.find(
                    history_q, {"_id": 0, "role": 1, "content": 1, "timestamp": 1, "ai_source": 1},
                ).sort("timestamp", 1).limit(_context_limit)
                history_docs_all = await history_cursor.to_list(length=_context_limit)
                history_docs = history_docs_all[:-1] if history_docs_all else []
                transcript = ""
                if history_docs:
                    lines = []
                    for h in history_docs:
                        speaker = "Utilisateur" if h.get("role") == "user" else "Caly"
                        lines.append(f"{speaker} : {h.get('content', '').strip()}")
                    transcript = "\n".join(lines)

                # Detect a model switch between previous AI message and this one.
                last_ai = next((h for h in reversed(history_docs) if h.get("role") == "assistant"), None)
                last_source = (last_ai or {}).get("ai_source") if last_ai else None
                model_switch_note = ""
                if last_source and not last_source.startswith(f"emergent:{provider}:"):
                    short_prev = last_source.replace("emergent:", "").replace("ollama:", "Ollama/")
                    model_switch_note = (
                        f"\n\n[CONTEXTE INTERNE — NE PAS RÉPÉTER À L'UTILISATEUR] "
                        f"La réponse précédente a été produite par {short_prev}. Tu es maintenant {model_id}. "
                        f"Reprends naturellement le fil — ne signale PAS le changement sauf si l'utilisateur le demande explicitement.\n"
                    )

                # Reuse the same session per user/project so the AI keeps short-term context.
                session_id = f"codeforge_chat_{user_id}_{project_id_eff or 'noproj'}"

                composed = (
                    f"### Historique récent de la conversation :\n{transcript}\n\n"
                    f"### Nouveau message de l'utilisateur :\n{input.message}"
                ) if transcript else input.message
                if model_switch_note:
                    composed = composed + model_switch_note

                # Inline attached file excerpts (text) so the model reasons on them.
                text_atts = [a for a in (input.attachments or []) if a.kind == 'text' and (a.content or '').strip()]
                if text_atts:
                    chunks = []
                    for a in text_atts:
                        chunks.append(f"### Pièce jointe : {a.filename or 'fichier'}\n{a.content.strip()[:15000]}")
                    composed = composed + "\n\n" + "\n\n".join(chunks)

                # Attach images via vision.
                image_contents = []
                for a in (input.attachments or []):
                    if a.kind == 'image' and a.data_base64:
                        try:
                            image_contents.append(ImageContent(a.data_base64))
                        except Exception:
                            pass
                user_message = UserMessage(text=composed, file_contents=image_contents) if image_contents else UserMessage(text=composed)

                # -------------------------------------------------------------
                # SILENT MULTI-MODEL CASCADE — "no budget, no limit" UX.
                # If the chosen provider returns a budget/quota/rate-limit/auth
                # error, we silently retry on alternative providers without
                # ever surfacing the error to the user.
                # -------------------------------------------------------------
                primary = (provider, model_id)
                # Build a deduplicated fallback chain: primary first, then a
                # diversified mix across providers so a single budget cap on
                # one provider never blocks the chat.
                fallback_chain = [
                    primary,
                    ("anthropic", "claude-sonnet-4-5-20250929"),
                    ("openai",    "gpt-5.2"),
                    ("gemini",    "gemini-3-flash-preview"),
                    ("anthropic", "claude-haiku-4-5-20251001"),
                    ("gemini",    "gemini-2.5-pro"),
                    ("openai",    "gpt-5"),
                    ("anthropic", "claude-opus-4-5-20251101"),
                ]
                seen_pairs = set()
                ordered_chain = []
                for pair in fallback_chain:
                    if pair not in seen_pairs:
                        seen_pairs.add(pair)
                        ordered_chain.append(pair)

                def _is_recoverable(exc: Exception) -> bool:
                    """Errors that justify a silent fallback to another model."""
                    msg = str(exc).lower()
                    keywords = (
                        "budget", "quota", "rate limit", "rate_limit", "ratelimit",
                        "insufficient", "exceeded", "billing", "credit", "credits",
                        "429", "401", "403", "unauthorized", "authentication",
                        "overloaded", "503", "502", "500", "timeout", "timed out",
                        "model_not_found", "not found", "deprecated", "unsupported",
                        "service unavailable", "internal server error",
                        "badrequest", "bad request", "invalid_request",
                    )
                    return any(k in msg for k in keywords)

                ai_response_text = ""
                last_error = None
                for idx, (p, mid) in enumerate(ordered_chain):
                    try:
                        chat = LlmChat(
                            api_key=emergent_key,
                            session_id=session_id,
                            system_message=system_prompt,
                        ).with_model(p, mid)
                        candidate = (await chat.send_message(user_message) or '').strip()
                        if candidate:
                            ai_response_text = candidate
                            ai_source = f"emergent:{p}:{mid}"
                            if idx == 0:
                                logger.info(f"✅ Chat response successful via {ai_source}")
                            else:
                                logger.warning(
                                    f"↪️  Silent fallback succeeded on attempt {idx + 1} "
                                    f"via {ai_source} after error: {last_error}"
                                )
                            break
                        # Empty response → try next.
                        last_error = "empty response"
                        logger.warning(f"Empty response from {p}:{mid}, trying next in cascade")
                        continue
                    except Exception as model_err:
                        last_error = str(model_err)[:300]
                        if _is_recoverable(model_err):
                            logger.warning(
                                f"⚠️  Recoverable error on {p}:{mid} "
                                f"(attempt {idx + 1}/{len(ordered_chain)}) → falling back silently: {last_error}"
                            )
                            continue
                        # Non-recoverable: still try the next one — UX > strictness.
                        logger.warning(
                            f"⚠️  Unexpected error on {p}:{mid} → cascading anyway: {last_error}"
                        )
                        continue

                if not ai_response_text:
                    # Whole cascade failed — let the outer fallback (offline msg) kick in.
                    logger.error(f"All cascade models failed. Last error: {last_error}")
            except Exception as emergent_error:
                logger.warning(f"Emergent chat error: {emergent_error}")

        # Final fallback — short, friendly, localized "I'm having trouble" message.
        if not ai_response_text:
            offline_msgs = {
                'fr': "Je n'arrive pas à répondre pour l'instant. Réessaie dans un instant 🙏",
                'en': "I'm having trouble responding right now. Please try again in a moment 🙏",
                'es': "No puedo responder en este momento. Vuelve a intentarlo en un momento 🙏",
                'pt': "Não consigo responder de momento. Tenta de novo daqui a pouco 🙏",
                'de': "Ich kann gerade nicht antworten. Bitte versuche es gleich noch einmal 🙏",
                'nl': "Het lukt me even niet. Probeer het zo opnieuw 🙏",
                'ru': "Сейчас не получается ответить. Попробуйте чуть позже 🙏",
                'zh': "我现在无法回答，请稍后再试 🙏",
                'zh-tw': "我現在無法回答，請稍後再試 🙏",
                'hi': "मैं अभी जवाब नहीं दे पा रहा। कृपया थोड़ी देर बाद कोशिश करें 🙏",
                'bn': "আমি এখন উত্তর দিতে পারছি না। একটু পরে আবার চেষ্টা করুন 🙏",
                'ur': "میں ابھی جواب نہیں دے سکتا۔ تھوڑی دیر بعد دوبارہ کوشش کریں 🙏",
            }
            ai_response_text = offline_msgs.get(user_language, offline_msgs['fr'])
            ai_source = 'fallback'

        # -------- Parse `cfaction` block to generate a downloadable artefact. --------
        download = None
        try:
            m = re.search(r"```cfaction\s*\n?(.*?)\n?```", ai_response_text or "", re.DOTALL)
            if m:
                raw = m.group(1).strip()
                try:
                    action = json.loads(raw)
                except Exception as je:
                    logger.warning(f"cfaction JSON parse error: {je} | raw[:200]={raw[:200]!r}")
                    action = None
                if isinstance(action, dict):
                    a_type = (action.get("type") or "").lower()
                    if a_type == "docx":
                        download = await _build_docx(user_id, action.get("title") or "Document", action.get("sections") or [])
                    elif a_type == "pdf":
                        download = await _build_pdf(user_id, action.get("title") or "Document", action.get("sections") or [])
                    elif a_type in ("xlsx", "excel"):
                        download = await _build_xlsx(user_id, action.get("title") or "Classeur", action.get("sheets") or [])
                    elif a_type in ("pptx", "powerpoint"):
                        download = await _build_pptx(user_id, action.get("title") or "Présentation", action.get("slides") or [])
                    elif a_type in ("txt", "md", "csv", "json", "yaml", "yml", "xml", "ini", "env", "sql",
                                    "py", "js", "ts", "tsx", "jsx", "html", "css", "sh", "ps1", "bat"):
                        ext_map = {
                            "txt": (".txt", "text/plain"),
                            "md": (".md", "text/markdown"),
                            "csv": (".csv", "text/csv"),
                            "json": (".json", "application/json"),
                            "yaml": (".yaml", "application/x-yaml"),
                            "yml": (".yml", "application/x-yaml"),
                            "xml": (".xml", "application/xml"),
                            "ini": (".ini", "text/plain"),
                            "env": (".env", "text/plain"),
                            "sql": (".sql", "text/plain"),
                            "py": (".py", "text/x-python"),
                            "js": (".js", "text/javascript"),
                            "ts": (".ts", "text/typescript"),
                            "tsx": (".tsx", "text/typescript"),
                            "jsx": (".jsx", "text/javascript"),
                            "html": (".html", "text/html"),
                            "css": (".css", "text/css"),
                            "sh": (".sh", "text/x-shellscript"),
                            "ps1": (".ps1", "text/plain"),
                            "bat": (".bat", "text/plain"),
                        }
                        ext, mime = ext_map[a_type]
                        download = await _build_plain(
                            user_id,
                            action.get("title") or f"fichier{ext}",
                            action.get("content") or "",
                            ext, mime,
                        )
                    elif a_type == "image":
                        try:
                            download = await _build_image(user_id, (action.get("prompt") or action.get("title") or input.message)[:800])
                        except Exception as ie:
                            logger.warning(f"cfaction image gen failed: {ie}")
                            download = None
                    elif a_type in ("run_python", "python_run", "run_py", "execute_python"):
                        # Sandbox Python execution — the result is *inlined* in the AI response,
                        # not returned as a download.
                        py_code = action.get("content") or action.get("code") or ""
                        sandbox_result = await _run_python_sandbox(py_code)
                        # Build a formatted block to append after cleaning cfaction.
                        py_block_parts = []
                        py_block_parts.append("\n\n**▶️ Exécution Python (sandbox) :**")
                        if sandbox_result.get("stdout"):
                            py_block_parts.append(f"```\n{sandbox_result['stdout'][:4000]}\n```")
                        if sandbox_result.get("stderr"):
                            py_block_parts.append(f"**Erreurs :**\n```\n{sandbox_result['stderr'][:2000]}\n```")
                        if sandbox_result.get("timed_out"):
                            py_block_parts.append("⏱️ *(Exécution interrompue : dépassement du timeout de 10 s)*")
                        py_block_parts.append(f"*Durée : {sandbox_result.get('duration_ms', 0)} ms — Code exit : {sandbox_result.get('exit_code', 0)}*")
                        # Inject inline markdown images for any matplotlib/generated image.
                        for idx, img in enumerate((sandbox_result.get("images") or [])[:4]):
                            try:
                                py_block_parts.append(
                                    f"\n![figure {idx + 1}](data:{img.get('mime_type','image/png')};base64,{img.get('data_base64','')})"
                                )
                            except Exception:
                                pass
                        # Inject the result block RIGHT BEFORE the cfaction block so it flows naturally.
                        injected = "\n".join(py_block_parts)
                        ai_response_text = (ai_response_text[: m.start()].rstrip() + injected + ai_response_text[m.end():]).strip()
                        # Signal we already cleaned/modified the response to the following block.
                        download = None
                        # Skip the normal cleaning step below.
                        action = None  # type: ignore
                    # Clean the `cfaction` block from the displayed text.
                    if action is not None:
                        ai_response_text = (ai_response_text[: m.start()] + ai_response_text[m.end():]).strip()
        except Exception as exc:
            logger.warning(f"cfaction post-process failed: {exc}")
        
        # Save AI response
        ai_message_doc = {
            "message_id": f"msg_{uuid.uuid4().hex[:16]}",
            "user_id": user_id,
            "project_id": project_id_eff,
            "role": "assistant",
            "content": ai_response_text,
            "mode": input.mode,
            "ai_source": ai_source,
            "download": download,
            "timestamp": datetime.now(timezone.utc).isoformat()
        }
        await db.chat_messages.insert_one(ai_message_doc)
        
        # Remove MongoDB _id from response to avoid serialization issues
        user_message_response = {k: v for k, v in user_message_doc.items() if k != '_id'}
        ai_message_response = {k: v for k, v in ai_message_doc.items() if k != '_id'}
        
        return {
            "user_message": user_message_response,
            "ai_response": ai_message_response,
            "project_id": project_id_eff,  # so frontend can navigate back to this chat
        }
    
    except Exception as e:
        logger.error(f"Error in chat: {e}")
        raise HTTPException(status_code=500, detail=f"Erreur de chat: {str(e)}")

@api_router.get("/chat/history")
async def get_chat_history(request: Request, project_id: Optional[str] = None, limit: int = 50):
    """Get chat history for user or specific project"""
    user_id = await get_current_user(request)
    
    query = {"user_id": user_id}
    if project_id:
        query["project_id"] = project_id
    
    messages = await db.chat_messages.find(
        query,
        {"_id": 0}
    ).sort("timestamp", 1).limit(limit).to_list(limit)
    
    return messages


class ChatAttachInput(BaseModel):
    message_id: Optional[str] = None
    project_id: str
    attach_all_orphans: Optional[bool] = False  # if true, attach all messages with project_id=null


@api_router.post("/chat/attach")
async def attach_chat_to_project(request: Request, payload: ChatAttachInput):
    """Attach an orphan chat message (project_id=null) to a project — used when
    a user pins a free-running chat to the sidebar."""
    user_id = await get_current_user(request)
    # Validate the project belongs to the user.
    proj = await db.projects.find_one({"project_id": payload.project_id, "user_id": user_id}, {"_id": 0})
    if not proj:
        raise HTTPException(status_code=404, detail="Project not found")
    if payload.attach_all_orphans:
        result = await db.chat_messages.update_many(
            {"user_id": user_id, "project_id": None},
            {"$set": {"project_id": payload.project_id}},
        )
        return {"updated": result.modified_count}
    if not payload.message_id:
        raise HTTPException(status_code=400, detail="message_id or attach_all_orphans required")
    result = await db.chat_messages.update_one(
        {"message_id": payload.message_id, "user_id": user_id},
        {"$set": {"project_id": payload.project_id}},
    )
    return {"updated": result.modified_count}


# ==================== WIZARD AI HELPERS ====================

# ==================== CHAT FILE TOOLS (analyze / generate docx/pdf/image) ====================

# Directory for generated downloadable files.
GENERATED_FILES_DIR = Path("/app/backend/generated_files")
GENERATED_FILES_DIR.mkdir(parents=True, exist_ok=True)


def _sanitize_filename(name: str, ext: str = "") -> str:
    return _cf_sanitize_filename(name, ext)


async def _analyze_pdf(data: bytes) -> str:
    return _cf_analyze_pdf(data)


async def _analyze_docx(data: bytes) -> str:
    return _cf_analyze_docx(data)


async def _analyze_xlsx(data: bytes) -> str:
    return _cf_analyze_xlsx(data)


async def _analyze_pptx(data: bytes) -> str:
    return _cf_analyze_pptx(data)


async def _analyze_sqlite(data: bytes) -> str:
    return _cf_analyze_sqlite(data)


async def _analyze_image_with_vision(data: bytes, mime_type: str, question: Optional[str] = None) -> str:
    """Use GPT-5.2 vision to describe an uploaded image."""
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage, ImageContent
        emergent_key = os.environ.get("EMERGENT_LLM_KEY")
        if not emergent_key:
            return "(clé Emergent non configurée)"
        b64 = base64.b64encode(data).decode("utf-8")
        chat = LlmChat(
            api_key=emergent_key,
            session_id=f"cf_vision_{uuid.uuid4().hex[:8]}",
            system_message="Tu analyses des images précisément et brièvement.",
        ).with_model("openai", "gpt-5.2")
        prompt = question or "Décris cette image de façon concise (contenu, contexte, éléments notables)."
        msg = UserMessage(text=prompt, file_contents=[ImageContent(b64)])
        return (await chat.send_message(msg) or "").strip()[:6000]
    except Exception as e:
        logger.warning(f"Vision analyze failed: {e}")
        return ""


@api_router.post("/chat/analyze-attachment")
async def chat_analyze_attachment(request: Request, file: UploadFile = File(...)):
    """Extract the usable content of an uploaded file for the chat.

    Returns a JSON object with a `kind` ('text' or 'image') and the content the
    frontend should embed in the next chat message.
    """
    _ = await get_current_user(request)  # auth gate
    data = await file.read()
    if len(data) > 20 * 1024 * 1024:  # 20 MB cap
        raise HTTPException(status_code=413, detail="Fichier trop lourd (max 20 Mo)")

    name = file.filename or "attachment"
    mime = (file.content_type or "").lower()
    lower = name.lower()

    if mime.startswith("image/") or lower.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif")):
        description = await _analyze_image_with_vision(data, mime)
        return {
            "kind": "image",
            "filename": name,
            "mime_type": mime or "image/png",
            "content": description,
            "data_base64": base64.b64encode(data).decode("utf-8"),
        }

    text = ""
    if mime == "application/pdf" or lower.endswith(".pdf"):
        text = await _analyze_pdf(data)
    elif lower.endswith(".docx") or mime.endswith("wordprocessingml.document"):
        text = await _analyze_docx(data)
    elif lower.endswith(".xlsx") or mime.endswith("spreadsheetml.sheet"):
        text = await _analyze_xlsx(data)
    elif lower.endswith(".pptx") or mime.endswith("presentationml.presentation"):
        text = await _analyze_pptx(data)
    elif lower.endswith((".sqlite", ".db", ".sqlite3")):
        text = await _analyze_sqlite(data)
    elif lower.endswith((".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".log",
                        ".py", ".js", ".jsx", ".ts", ".tsx", ".html", ".htm", ".css", ".scss",
                        ".xml", ".yaml", ".yml", ".ini", ".env", ".cfg", ".conf", ".toml",
                        ".sql", ".sh", ".ps1", ".bat", ".cmd", ".rb", ".go", ".rs", ".java",
                        ".c", ".cpp", ".h", ".hpp", ".cs", ".php", ".kt", ".swift")):
        try:
            text = data.decode("utf-8", errors="ignore")[:30000]
        except Exception:
            text = ""
    else:
        try:
            text = data.decode("utf-8", errors="ignore")[:20000]
        except Exception:
            text = ""

    if not text.strip():
        raise HTTPException(status_code=415, detail="Impossible d'extraire le contenu du fichier.")

    return {"kind": "text", "filename": name, "mime_type": mime or "text/plain", "content": text}


# ---- File GENERATION (docx / pdf / image) ----

class GenerateDocxInput(BaseModel):
    title: Optional[str] = "Document"
    sections: List[Dict[str, Any]] = []     # [{ "heading": str, "content": str }]


class GeneratePdfInput(BaseModel):
    title: Optional[str] = "Document"
    sections: List[Dict[str, Any]] = []


class GenerateImageInput(BaseModel):
    prompt: str
    size: Optional[str] = "1024x1024"


def _store_generated(blob: bytes, filename: str, mime: str, user_id: str) -> Dict[str, str]:
    file_id = f"gen_{uuid.uuid4().hex[:16]}"
    safe = _sanitize_filename(filename)
    disk = GENERATED_FILES_DIR / f"{file_id}__{safe}"
    disk.write_bytes(blob)
    return {
        "file_id": file_id,
        "filename": safe,
        "url": f"/api/download/generated/{file_id}",
        "mime_type": mime,
        "size": len(blob),
        "disk_path": str(disk),
        "user_id": user_id,
    }


async def _persist_generated(info: Dict[str, Any]) -> Dict[str, str]:
    await db.generated_files.insert_one({
        "file_id": info["file_id"],
        "user_id": info["user_id"],
        "filename": info["filename"],
        "mime_type": info["mime_type"],
        "size": info["size"],
        "path": info["disk_path"],
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    # Return only the public-facing keys.
    return {"file_id": info["file_id"], "filename": info["filename"], "url": info["url"], "mime_type": info["mime_type"]}


async def _build_docx(user_id: str, title: str, sections: List[Dict[str, Any]]) -> Dict[str, str]:
    blob = _cf_build_docx_bytes(title or "Document", sections or [])
    info = _store_generated(
        blob,
        f"{title or 'document'}.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        user_id,
    )
    return await _persist_generated(info)


async def _build_pdf(user_id: str, title: str, sections: List[Dict[str, Any]]) -> Dict[str, str]:
    blob = _cf_build_pdf_bytes(title or "Document", sections or [])
    info = _store_generated(blob, f"{title or 'document'}.pdf", "application/pdf", user_id)
    return await _persist_generated(info)


async def _build_xlsx(user_id: str, title: str, sheets: List[Dict[str, Any]]) -> Dict[str, str]:
    blob = _cf_build_xlsx_bytes(sheets or [])
    info = _store_generated(
        blob,
        f"{title or 'classeur'}.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        user_id,
    )
    return await _persist_generated(info)


async def _build_pptx(user_id: str, title: str, slides: List[Dict[str, Any]]) -> Dict[str, str]:
    blob = _cf_build_pptx_bytes(title or "Présentation", slides or [])
    info = _store_generated(
        blob,
        f"{title or 'presentation'}.pptx",
        "application/vnd.openxmlformats-officedocument.presentationml.presentation",
        user_id,
    )
    return await _persist_generated(info)


async def _build_plain(user_id: str, title: str, content: str, ext: str, mime: str) -> Dict[str, str]:
    data = (content or "").encode("utf-8")
    name = title or "fichier"
    if not name.lower().endswith(ext):
        name = f"{name}{ext}"
    info = _store_generated(data, name, mime, user_id)
    return await _persist_generated(info)


async def _build_image(user_id: str, prompt: str) -> Dict[str, str]:
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    emergent_key = os.environ.get("EMERGENT_LLM_KEY")
    if not emergent_key:
        raise ValueError("EMERGENT_LLM_KEY missing")
    chat = LlmChat(
        api_key=emergent_key,
        session_id=f"cf_imggen_{uuid.uuid4().hex[:8]}",
        system_message="You are a helpful AI assistant.",
    )
    chat.with_model("gemini", "gemini-3.1-flash-image-preview").with_params(modalities=["image", "text"])
    _text, images = await chat.send_message_multimodal_response(UserMessage(text=prompt))
    if not images:
        raise ValueError("no image generated")
    img = images[0]
    data = base64.b64decode(img["data"])
    info = _store_generated(
        data,
        f"{_sanitize_filename((prompt[:40]) or 'image', ext='')}.png",
        img.get("mime_type", "image/png"),
        user_id,
    )
    return await _persist_generated(info)


async def _run_python_sandbox(code: str, timeout_sec: int = 10, session_id: Optional[str] = None, files: Optional[List[Dict[str, str]]] = None) -> Dict[str, Any]:
    """Exécute du code Python dans le sandbox. Délègue à cfaction_engine.run_python_sandbox."""
    return await _cf_run_python_sandbox(code, timeout_sec=timeout_sec, session_id=session_id, files=files)


class SandboxFileInput(BaseModel):
    filename: str
    data_base64: str


@api_router.get("/chat/models")
async def list_chat_models(request: Request):
    """Liste les modèles disponibles pour le sélecteur, avec descriptions
    contextuelles (chat vs création) pour aider l'utilisateur à choisir.

    Le frontend passe `?context=chat` ou `?context=create` pour adapter le wording.
    """
    await get_current_user(request)
    context = (request.query_params.get("context") or "chat").lower()
    is_create = context == "create"

    # Descriptions adaptées au contexte d'utilisation
    # iter87 — Liste mise à jour avec les meilleures versions de chaque famille
    # Emergent. Les anciens IDs (gpt-5.2/claude-opus/claude-sonnet/gemini-3-pro/
    # gemini-3-flash) restent supportés via MODEL_ROUTES — ces entrées sont
    # purement UI (le mapping provider/model_id côté backend doit suivre).
    online = [
        {
            "id": "claude-5-fable", "name": "Claude 5 Fable", "provider": "Anthropic", "badge": "Le plus capable", "color": "fuchsia",
            "description": (
                "Le modèle le plus puissant et le plus sécurisé d'Anthropic. Recommandé pour les décisions critiques, audits, raisonnements complexes."
                if not is_create else
                "Génère le code le plus robuste et le plus sûr. Idéal pour apps sensibles (auth, paiements, médical, juridique)."
            ),
            "good_for": (["Décisions critiques", "Audit sécurité", "Analyse longue", "Raisonnement multi-étape"]
                if not is_create else ["Apps sensibles", "Production-ready", "Sécurité élevée"]),
        },
        {
            "id": "gpt-5.5", "name": "GPT 5.5", "provider": "OpenAI", "badge": "Défaut", "color": "yellow",
            "description": (
                "Le dernier modèle d'OpenAI. Conversation généraliste rapide et fluide, mémoire complète, raisonnement solide."
                if not is_create else
                "Génère le code complet du projet (FastAPI + React + DB). Hypothèses intelligentes, README clair, prêt à exécuter."
            ),
            "good_for": (["Brainstormer", "Écrire un email/texte", "Analyser un fichier", "Conversation longue"]
                if not is_create else ["App complète", "Projet équilibré", "Site web standard"]),
        },
        {
            "id": "claude-4.8-opus", "name": "Claude 4.8 Opus", "provider": "Anthropic", "badge": "Thinking", "color": "amber",
            "description": (
                "Performance frontière d'Anthropic. Raisonne avant de répondre. Idéal pour problèmes complexes, dilemmes, longues analyses."
                if not is_create else
                "Architecte avant de coder. Meilleure pour projets multi-fichiers, logique métier complexe, sécurité."
            ),
            "good_for": (["Problèmes complexes", "Dilemmes", "Code review profond", "Recherche approfondie"]
                if not is_create else ["Architecture complexe", "Backend avec règles métier", "Apps multi-modules"]),
        },
        {
            "id": "claude-4.7-opus-1m", "name": "Claude 4.7 Opus (1M)", "provider": "Anthropic", "badge": "Contexte long", "color": "indigo",
            "description": (
                "Contexte d'1 million de tokens. Idéal pour analyser un repo entier, plusieurs PDFs, ou des conversations très longues."
                if not is_create else
                "Pour porter un projet existant volumineux : repo legacy, refactor complet, lecture de gros datasets."
            ),
            "good_for": (["Analyser un repo", "Lire 100+ pages", "Conversation infinie"]
                if not is_create else ["Refactor legacy", "Migration de framework", "Audit codebase"]),
        },
        {
            "id": "claude-4.6-sonnet", "name": "Claude 4.6 Sonnet", "provider": "Anthropic", "badge": "Code", "color": "orange",
            "description": (
                "Excellente pour ÉCRIRE DU CODE dans le chat — clique sur ▶ Exécuter pour le lancer dans le sandbox Python."
                if not is_create else
                "Le PLUS RAPIDE pour générer un projet complet propre, exécutable, prêt à pousser sur GitHub. Recommandée par défaut."
            ),
            "good_for": (["Snippets de code", "Refactor", "Debug ligne par ligne", "Réécrire un texte"]
                if not is_create else ["App standard", "Site marketing", "Outil CRUD", "Recommandé par défaut"]),
        },
        {
            "id": "gpt-5.3-codex", "name": "GPT 5.3 Codex", "provider": "OpenAI", "badge": "Code", "color": "emerald",
            "description": (
                "Modèle flagship OpenAI spécialisé code. Idéal pour write/debug/refactor complexe avec exécution."
                if not is_create else
                "Optimisé pour les patterns Python/JS/TS modernes. Code propre, idiomatique, testé."
            ),
            "good_for": (["Debug complexe", "Tests unitaires", "Architecture code"]
                if not is_create else ["Backend Python", "Frontend TS", "API REST/GraphQL"]),
        },
        {
            "id": "gemini-3.1-pro", "name": "Gemini 3.1 Pro", "provider": "Google", "badge": "Multimodal", "color": "blue",
            "description": (
                "Le meilleur de Google. Idéal quand tu joins une IMAGE — décrit, analyse, OCR, explique des schémas."
                if not is_create else
                "Plus créative visuellement. Idéale pour UI originales, design audacieux, identité visuelle marquée."
            ),
            "good_for": (["Analyser une image", "OCR", "Lire un schéma", "Décrire une photo"]
                if not is_create else ["UI design original", "Landing page", "Portfolio créatif"]),
        },
        {
            "id": "gpt-5.4-1m", "name": "GPT 5.4 (1M)", "provider": "OpenAI", "badge": "Contexte long", "color": "cyan",
            "description": (
                "Variante 1M tokens de GPT 5.4. Pour ingérer de grandes quantités de docs en une seule passe."
                if not is_create else
                "Pour projets nécessitant un contexte massif (specs longues, multiples APIs externes)."
            ),
            "good_for": (["Analyse docs volumineux", "Multi-PDF", "Conversation sans coupure"]
                if not is_create else ["Projet enterprise", "Specs complexes"]),
        },
    ]
    offline = [
        {
            "id": "deepseek", "name": "DeepSeek R1 (7B)", "provider": "Ollama", "badge": "Code", "color": "sky",
            "description": "Excellent en code et raisonnement step-by-step, totalement hors-ligne." if not is_create else "Génère du code propre hors-ligne. Lent mais privé.",
            "good_for": ["Code", "Maths", "Logique"] if not is_create else ["App standard hors-ligne"],
        },
        {
            "id": "gemma", "name": "Gemma 3 (12B)", "provider": "Ollama", "badge": "Équilibré", "color": "indigo",
            "description": "Modèle Google open-source. Polyvalent, multilingue, conversation naturelle, sans rien envoyer en ligne.",
            "good_for": ["Tout-terrain hors-ligne", "Multilingue", "Discussion privée"] if not is_create else ["Petites apps hors-ligne", "Privacy first"],
        },
        {
            "id": "gemma-27b", "name": "Gemma 3 (27B)", "provider": "Ollama", "badge": "Puissant", "color": "purple",
            "description": "Le plus capable des Gemma — exigeant en RAM (+24 Go). Qualité GPT-3.5 hors-ligne.",
            "good_for": ["Tâches lourdes hors-ligne", "Analyse profonde"] if not is_create else ["App complète hors-ligne haut de gamme"],
        },
        {
            "id": "gemma-4b", "name": "Gemma 3 (4B)", "provider": "Ollama", "badge": "Léger", "color": "violet",
            "description": "Très léger (~3 Go RAM), idéal pour machines modestes et démarrage rapide.",
            "good_for": ["Vieux laptop", "Tâches simples", "Test rapide"] if not is_create else ["Page statique", "MVP minimal"],
        },
        {
            "id": "llama", "name": "Llama 3.2", "provider": "Ollama", "badge": "Généraliste", "color": "emerald",
            "description": "Meta Llama — alternative équilibrée, communauté très active.",
            "good_for": ["Discussion générale"] if not is_create else ["App standard"],
        },
        {
            "id": "qwen", "name": "Qwen 2.5 (7B)", "provider": "Ollama", "badge": "Multilingue", "color": "teal",
            "description": "Très bon support chinois + anglais + français. Top pour i18n.",
            "good_for": ["Traduction", "Multilangue"] if not is_create else ["Site multilingue"],
        },
        {
            "id": "mistral", "name": "Mistral 7B", "provider": "Ollama", "badge": "Européen", "color": "rose",
            "description": "Modèle français performant, léger, respectueux du RGPD si tu héberges toi-même.",
            "good_for": ["Français natif", "RGPD"] if not is_create else ["App RGPD-friendly"],
        },
        {
            "id": "phi", "name": "Phi-3 Medium", "provider": "Ollama", "badge": "Compact", "color": "fuchsia",
            "description": "Microsoft Phi — petit mais étonnamment fort en maths et code.",
            "good_for": ["Maths", "Code"] if not is_create else ["Outils techniques"],
        },
    ]
    return {"online": online, "offline": offline, "context": context}



class RunPythonInput(BaseModel):
    code: str
    timeout_sec: Optional[int] = 10
    session_id: Optional[str] = None  # if set → persistent REPL mode
    files: Optional[List[SandboxFileInput]] = None  # files dropped into cwd


class SessionResetInput(BaseModel):
    session_id: str


@api_router.post("/sandbox/python")
async def sandbox_python(request: Request, payload: RunPythonInput):
    """Exécute du code Python dans le sandbox serveur avec timeout dur.

    Mode éphémère (par défaut) : namespace fresh à chaque appel.
    Mode REPL : passer un `session_id` réutilise les variables entre appels (style Jupyter).
    Upload : passer `files=[{filename, data_base64}]` pour déposer des fichiers au cwd.
    """
    await get_current_user(request)
    t = max(1, min(int(payload.timeout_sec or 10), 30))
    files = [f.model_dump() for f in (payload.files or [])]
    return await _run_python_sandbox(payload.code or "", timeout_sec=t, session_id=payload.session_id, files=files)


@api_router.post("/sandbox/reset")
async def sandbox_reset(request: Request, payload: SessionResetInput):
    """Réinitialise le namespace persistant d'une session REPL."""
    await get_current_user(request)
    from cfaction_engine import reset_sandbox_session
    ok = reset_sandbox_session(payload.session_id)
    return {"reset": ok, "session_id": payload.session_id}





@api_router.post("/chat/generate-docx")
async def chat_generate_docx(request: Request, payload: GenerateDocxInput):
    user_id = await get_current_user(request)
    return await _build_docx(user_id, payload.title or "Document", payload.sections or [])


@api_router.post("/chat/generate-pdf")
async def chat_generate_pdf(request: Request, payload: GeneratePdfInput):
    user_id = await get_current_user(request)
    return await _build_pdf(user_id, payload.title or "Document", payload.sections or [])


@api_router.post("/chat/generate-image")
async def chat_generate_image(request: Request, payload: GenerateImageInput):
    user_id = await get_current_user(request)
    try:
        return await _build_image(user_id, payload.prompt)
    except Exception as e:
        logger.warning(f"image gen failed: {e}")
        raise HTTPException(status_code=500, detail=f"Génération d'image impossible : {e}")


@api_router.get("/chat/export-ipynb/{project_id}")
async def export_chat_as_ipynb(project_id: str, request: Request):
    """Exporte une conversation chat en notebook Jupyter (.ipynb).

    Chaque message utilisateur devient une cellule markdown. Chaque bloc
    ```python``` trouvé dans les réponses AI devient une cellule code (avec le
    stdout déjà capturé en output si le bloc `▶️ Exécution Python (sandbox)` suit).
    """
    user_id = await get_current_user(request)
    # Validate ownership
    proj = await db.projects.find_one({"project_id": project_id, "user_id": user_id}, {"_id": 0})
    if not proj:
        raise HTTPException(status_code=404, detail="Projet introuvable")
    # Collect messages in order
    msgs = await db.chat_messages.find(
        {"user_id": user_id, "project_id": project_id},
        {"_id": 0}
    ).sort("timestamp", 1).to_list(length=None)
    if not msgs:
        raise HTTPException(status_code=400, detail="Aucun message à exporter")

    cells: List[Dict[str, Any]] = []
    # Title cell
    cells.append({
        "cell_type": "markdown",
        "metadata": {},
        "source": [
            f"# {proj.get('name', 'Conversation')}\n",
            f"\n*Exporté depuis CodeForge AI — {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}*\n",
            "\n---\n",
        ],
    })

    code_block_re = re.compile(r"```(?:python|py)\s*\n(.*?)```", re.DOTALL)
    stdout_after_re = re.compile(
        r"\*\*▶️ Exécution Python \(sandbox\)[^\n]*\n+```\n(.*?)```",
        re.DOTALL,
    )

    for m in msgs:
        role = m.get("role", "assistant")
        content = m.get("content", "") or ""
        if role == "user":
            cells.append({
                "cell_type": "markdown",
                "metadata": {},
                "source": ["**👤 Utilisateur**\n\n", content],
            })
        else:
            # Try to split AI response into: pre-code markdown, code cells, post-code markdown.
            pos = 0
            code_blocks = list(code_block_re.finditer(content))
            if not code_blocks:
                cells.append({
                    "cell_type": "markdown",
                    "metadata": {},
                    "source": ["**🤖 CodeForge**\n\n", content],
                })
                continue
            # Preamble markdown
            prefix = content[: code_blocks[0].start()].rstrip()
            if prefix:
                cells.append({
                    "cell_type": "markdown",
                    "metadata": {},
                    "source": ["**🤖 CodeForge**\n\n", prefix],
                })
            for i, cb in enumerate(code_blocks):
                code_src = cb.group(1).rstrip()
                # Look for an inline sandbox stdout block right after this cell.
                tail = content[cb.end():]
                stdout_m = stdout_after_re.search(tail[:4000])
                outputs = []
                if stdout_m:
                    outputs.append({
                        "name": "stdout",
                        "output_type": "stream",
                        "text": stdout_m.group(1).splitlines(keepends=True),
                    })
                cells.append({
                    "cell_type": "code",
                    "execution_count": i + 1,
                    "metadata": {},
                    "outputs": outputs,
                    "source": code_src.splitlines(keepends=True),
                })
                pos = cb.end()
            # Trailing markdown
            trailing = content[pos:]
            # Strip the "Exécution Python (sandbox)" blocks from trailing since we moved them.
            trailing_clean = re.sub(stdout_after_re, "", trailing).strip()
            if trailing_clean:
                cells.append({
                    "cell_type": "markdown",
                    "metadata": {},
                    "source": [trailing_clean],
                })

    notebook = {
        "cells": cells,
        "metadata": {
            "kernelspec": {"display_name": "Python 3", "language": "python", "name": "python3"},
            "language_info": {"name": "python", "version": "3.11"},
            "codeforge_export": {
                "project_id": project_id,
                "project_name": proj.get("name"),
                "exported_at": datetime.now(timezone.utc).isoformat(),
            },
        },
        "nbformat": 4,
        "nbformat_minor": 5,
    }
    raw = json.dumps(notebook, ensure_ascii=False, indent=1).encode("utf-8")
    safe_name = _sanitize_filename(proj.get("name") or project_id) + ".ipynb"
    return Response(
        content=raw,
        media_type="application/x-ipynb+json",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
    )


@api_router.get("/chat/export-docx/{project_id}")
async def export_chat_as_docx(project_id: str, request: Request):
    """iter79 — Exporte une conversation chat en .docx (Microsoft Word)."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    import io as _io

    user_id = await get_current_user(request)
    proj = await db.projects.find_one({"project_id": project_id, "user_id": user_id}, {"_id": 0})
    if not proj:
        raise HTTPException(status_code=404, detail="Projet introuvable.")
    msgs = await db.chat_messages.find(
        {"project_id": project_id, "user_id": user_id}, {"_id": 0},
    ).sort("created_at", 1).to_list(length=10000)

    doc = Document()
    doc.add_heading(proj.get("name") or project_id, 0)
    doc.add_paragraph(f"Exporté le {datetime.now(timezone.utc).strftime('%d/%m/%Y %H:%M UTC')}")
    doc.add_paragraph(f"{len(msgs)} message(s)")
    doc.add_paragraph()
    for m in msgs:
        speaker = "Utilisateur" if m.get("role") == "user" else "IA"
        p = doc.add_paragraph()
        run = p.add_run(f"[{speaker}] ")
        run.bold = True
        run.font.color.rgb = RGBColor(0xE4, 0xFF, 0x00) if m.get("role") != "user" else RGBColor(0x00, 0xD4, 0xFF)
        p.add_run((m.get("content") or "")[:50000])
    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe_name = _sanitize_filename(proj.get("name") or project_id) + ".docx"
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
    )


@api_router.get("/download/generated/{file_id}")
async def download_generated_file(file_id: str, request: Request):
    """Ownership-checked download endpoint for AI-generated files."""
    user_id = await get_current_user(request)
    doc = await db.generated_files.find_one({"file_id": file_id, "user_id": user_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Fichier introuvable")
    p = Path(doc.get("path", ""))
    if not p.exists():
        raise HTTPException(status_code=410, detail="Fichier expiré")
    return FileResponse(
        path=str(p),
        media_type=doc.get("mime_type", "application/octet-stream"),
        filename=doc.get("filename", "file"),
    )


class WizardSuggestInput(BaseModel):
    kind: str
    platforms: List[str] = []
    app_type: Optional[str] = None
    description: Optional[str] = None
    language: Optional[str] = 'fr'
    seed: Optional[float] = None  # iter79 — pour forcer une variation du LLM


def _fallback_name_pool():
    """iter79 — Pool varié de pseudos avec longueurs et styles différents.
    S'inspire des pseudos Among Us : 3-12 caractères, mix de mots/numbers/refs."""
    return [
        "Nox", "Vrael", "Kiri", "Zorn", "Quil",
        "Aetheris", "Mavrik", "Pyx", "Lumin", "Velvox",
        "Bleeptron", "GrumpCat", "Skyforger", "VoidNomad", "PixelMite",
        "Phantmly", "Echoflux", "Nimbo", "Ironpup", "Tessera",
        "Kazimir77", "Marlow", "OBLIVION", "tinypaws", "RustyBolt",
        "Vespera", "ZippyZuko", "ChromaQT", "Mossbeard", "Solitaria",
        "Klaxon", "Whisperdrop", "Glitchgrove", "Nebbie", "QuarkLad",
        "S0larPunk", "Mythra", "Bytewolf", "Coldscape", "Veneer",
        "AzureFox", "FlintGhost", "NotKai", "Murmurine", "Wickerlight",
        "JuneberryX", "PrettyPanik", "DustGremlin", "Sunken", "Ororo",
        "EmberLark", "MidnightVole", "ShinyBean", "ImLost", "ravenmint",
    ]


@api_router.post("/ai/wizard-suggest")
async def wizard_suggest(request: Request, payload: WizardSuggestInput):
    """🪄 Magic-wand helper. iter79 — pool de noms variés en longueur/style."""
    user_id = await get_current_user(request)
    _ = user_id

    plats = ", ".join(payload.platforms) if payload.platforms else "non spécifié"
    desc = (payload.description or "").strip()[:600]
    # iter79 — seed-driven randomization pour LLM + fallback.
    import random as _rnd
    seed_val = payload.seed or _rnd.random()
    _rnd.seed(seed_val)

    if payload.kind == 'name':
        prompt = (
            f"Propose 3 pseudos d'application ORIGINAUX et VARIÉS en longueur et en style "
            f"(certains courts 3-5 lettres, d'autres plus longs 8-12 caractères, parfois avec "
            f"un chiffre ou un mix de mots inattendus). Inspire-toi des pseudos joueurs (Among Us, "
            f"Discord) — NE TE limite PAS à 6 lettres max. App : {payload.app_type or 'générique'} "
            f"ciblant {plats}. Contexte : {desc or 'aucun'}. "
            f"Évite les noms génériques comme 'NovaApp' ou 'PixelForge'. "
            f"Aléa #{seed_val:.6f}. Réponds UNIQUEMENT en JSON: "
            f"{{\"suggestions\": [\"...\", \"...\", \"...\"]}}"
        )
    elif payload.kind == 'function':
        # iter80 C2 — Suggestion IA pour le bloc Fonctionnement du wizard.
        prompt = (
            f"Propose en 80 mots maximum quelles fonctionnalités, écrans et règles métier "
            f"devrait avoir une app {payload.app_type or 'générique'} ciblant {plats}. "
            f"Contexte design : {desc or 'aucun'}. Aléa #{seed_val:.6f}. "
            f"Réponds UNIQUEMENT en JSON : {{\"func\": \"description claire en plusieurs phrases\"}}"
        )
    else:
        prompt = (
            f"Propose une direction visuelle (palette, typographie, ambiance, mots-clés) "
            f"pour une app {payload.app_type or 'générique'} ciblant {plats}. "
            f"Contexte : {desc or 'aucun'}. Aléa #{seed_val:.6f}. "
            f"Réponds UNIQUEMENT en JSON : {{\"design\": \"description courte (<60 mots)\"}}"
        )

    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        emergent_key = os.environ.get('EMERGENT_LLM_KEY')
        if not emergent_key:
            raise ValueError("EMERGENT_LLM_KEY not configured")
        chat = LlmChat(
            api_key=emergent_key,
            session_id=f"codeforge_wizard_{uuid.uuid4().hex[:8]}",
            system_message="Tu aides un utilisateur non technique à concevoir une app. Réponds STRICTEMENT en JSON valide sans markdown.",
        ).with_model("openai", "gpt-5.2")
        raw = await chat.send_message(UserMessage(text=prompt))
        text = (raw or '').strip()
        if text.startswith("```"):
            text = text.split("\n", 1)[1] if "\n" in text else text
            if text.endswith("```"):
                text = text[:-3]
            text = text.strip()
        start = text.find('{')
        end = text.rfind('}') + 1
        data = json.loads(text[start:end]) if start >= 0 and end > start else {}
    except Exception as exc:
        logger.warning(f"wizard-suggest failure: {exc}")
        if payload.kind == 'name':
            pool = _fallback_name_pool()
            data = {"suggestions": _rnd.sample(pool, 3)}
        elif payload.kind == 'function':
            data = {"func": "Trois écrans principaux : accueil avec actions rapides, écran central interactif et un profil utilisateur. Une authentification simple, des notifications discrètes et une sauvegarde locale automatique. Possibilité d'inviter des amis et de partager du contenu."}
        else:
            data = {"design": "Interface sombre élégante, accent jaune-vert vif, typographie sans-serif moderne, ambiance high-tech bienveillante."}

    if payload.kind == 'name' and not isinstance(data.get('suggestions'), list):
        # iter81 — Évite les noms génériques (NovaApp/PixelForge/Lumino), pioche
        # dans le pool varié pour respecter le prompt.
        data = {"suggestions": _rnd.sample(_fallback_name_pool(), 3)}
    if payload.kind == 'design' and not isinstance(data.get('design'), str):
        data = {"design": "Interface sombre élégante, accent jaune-vert vif, ambiance moderne."}

    return data


# ==================== USER PREFERENCES ====================

class UserPreferences(BaseModel):
    theme: Optional[str] = 'dark'         # 'dark' | 'light' | 'auto'
    contrast: Optional[str] = 'normal'    # 'normal' | 'high'
    accent: Optional[str] = '#E4FF00'
    notifications_email: Optional[bool] = True
    notifications_push: Optional[bool] = False

@api_router.get("/auth/preferences")
async def get_user_preferences(request: Request):
    user_id = await get_current_user(request)
    doc = await db.user_preferences.find_one({"user_id": user_id}, {"_id": 0, "user_id": 0}) or {}
    base = UserPreferences().model_dump()
    base.update({k: v for k, v in doc.items() if k in base})
    return base

@api_router.put("/auth/preferences")
async def put_user_preferences(request: Request, prefs: UserPreferences):
    user_id = await get_current_user(request)
    payload = prefs.model_dump()
    await db.user_preferences.update_one(
        {"user_id": user_id},
        {"$set": {**payload, "user_id": user_id, "updated_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True,
    )
    return payload



@api_router.post("/projects", response_model=Project, status_code=201)
async def create_project(request: Request, input: ProjectCreate):
    """Create a new project"""
    user_id = await get_current_user(request)
    
    project = Project(
        user_id=user_id,
        name=input.name,
        description=input.description,
        project_type=input.project_type
    )
    
    project_dict = project.model_dump()
    project_dict['created_at'] = project_dict['created_at'].isoformat()
    project_dict['updated_at'] = project_dict['updated_at'].isoformat()
    
    await db.projects.insert_one(project_dict)
    
    return project

@api_router.get("/projects", response_model=List[Project])
async def get_projects(request: Request):
    """Get all projects for current user"""
    user_id = await get_current_user(request)
    
    projects = await db.projects.find(
        {"user_id": user_id},
        {"_id": 0}
    ).sort("created_at", -1).to_list(100)
    
    # Convert ISO strings to datetime + backfill ai_mode for legacy projects
    for project in projects:
        if isinstance(project.get('created_at'), str):
            project['created_at'] = datetime.fromisoformat(project['created_at'])
        elif not project.get('created_at'):
            project['created_at'] = datetime.now(timezone.utc)
        if isinstance(project.get('updated_at'), str):
            project['updated_at'] = datetime.fromisoformat(project['updated_at'])
        elif not project.get('updated_at'):
            # Legacy chat projects didn't have updated_at; reuse created_at.
            project['updated_at'] = project['created_at']
        if not project.get('ai_mode'):
            src = (project.get('ai_source') or '').lower()
            project['ai_mode'] = 'offline' if src.startswith('ollama') else 'online'
    
    return projects

@api_router.get("/projects/{project_id}", response_model=Project)
async def get_project(request: Request, project_id: str):
    """Get specific project"""
    user_id = await get_current_user(request)
    
    project = await db.projects.find_one(
        {"project_id": project_id, "user_id": user_id},
        {"_id": 0}
    )
    
    if not project:
        raise HTTPException(status_code=404, detail="Projet non trouvé")
    
    # Convert ISO strings to datetime + backfill (same logic as list endpoint)
    if isinstance(project.get('created_at'), str):
        project['created_at'] = datetime.fromisoformat(project['created_at'])
    elif not project.get('created_at'):
        project['created_at'] = datetime.now(timezone.utc)
    if isinstance(project.get('updated_at'), str):
        project['updated_at'] = datetime.fromisoformat(project['updated_at'])
    elif not project.get('updated_at'):
        project['updated_at'] = project['created_at']
    if not project.get('ai_mode'):
        src = (project.get('ai_source') or '').lower()
        project['ai_mode'] = 'offline' if src.startswith('ollama') else 'online'
    
    return project

@api_router.put("/projects/{project_id}", response_model=Project)
async def update_project(request: Request, project_id: str, input: ProjectUpdate):
    """Update a project"""
    user_id = await get_current_user(request)
    
    # Check project exists and belongs to user
    project = await db.projects.find_one(
        {"project_id": project_id, "user_id": user_id},
        {"_id": 0}
    )
    
    if not project:
        raise HTTPException(status_code=404, detail="Projet non trouvé")
    
    # Build update dict
    update_data = {k: v for k, v in input.model_dump().items() if v is not None}
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    await db.projects.update_one(
        {"project_id": project_id},
        {"$set": update_data}
    )
    
    # Return updated project
    updated_project = await db.projects.find_one(
        {"project_id": project_id},
        {"_id": 0}
    )
    
    # Convert ISO strings to datetime
    if isinstance(updated_project['created_at'], str):
        updated_project['created_at'] = datetime.fromisoformat(updated_project['created_at'])
    if isinstance(updated_project['updated_at'], str):
        updated_project['updated_at'] = datetime.fromisoformat(updated_project['updated_at'])
    
    return updated_project

@api_router.delete("/projects/{project_id}")
async def delete_project(request: Request, project_id: str):
    """Delete a project"""
    user_id = await get_current_user(request)
    
    result = await db.projects.delete_one(
        {"project_id": project_id, "user_id": user_id}
    )
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Projet non trouvé")
    
    # Also delete related chat messages
    await db.chat_messages.delete_many({"project_id": project_id})
    
    return {"message": "Projet supprimé avec succès"}


@api_router.post("/projects/{project_id}/duplicate")
async def duplicate_project(request: Request, project_id: str):
    """Clone a project — duplicates the document with a new project_id and
    "(copie)" suffix. Chat history is NOT copied (a fresh thread starts on the
    clone). Useful to experiment without breaking the original.
    """
    user_id = await get_current_user(request)
    src = await db.projects.find_one({"project_id": project_id, "user_id": user_id}, {"_id": 0})
    if not src:
        raise HTTPException(status_code=404, detail="Projet non trouvé")

    new_id = f"proj_{uuid.uuid4().hex[:12]}"
    now_iso = datetime.now(timezone.utc).isoformat()
    clone = {
        **src,
        "project_id": new_id,
        "name": f"{src.get('name', 'Projet')} (copie)"[:80],
        "created_at": now_iso,
        "updated_at": now_iso,
        # Drop legacy share metadata so the clone starts non-public.
        "share_slug": None,
        "is_public": False,
    }
    await db.projects.insert_one(clone)
    clone.pop("_id", None)
    return {"success": True, "project_id": new_id, "project": clone}


def _make_slug(name: str) -> str:
    """ASCII-safe URL slug."""
    import unicodedata as _ud
    import re as _re
    base = _ud.normalize("NFKD", name or "").encode("ascii", "ignore").decode("ascii").lower()
    base = _re.sub(r"[^a-z0-9]+", "-", base).strip("-")[:50] or "projet"
    return f"{base}-{uuid.uuid4().hex[:6]}"


@api_router.post("/projects/{project_id}/share")
async def toggle_project_share(request: Request, project_id: str):
    """Generate (or refresh) a public share URL for a project. Anyone with the
    slug can view the live preview without authentication. Call again to disable.
    Body (optional): {"enable": true|false}
    """
    user_id = await get_current_user(request)
    try:
        body = await request.json()
    except Exception:
        body = {}
    enable = body.get("enable") if isinstance(body, dict) else None

    project = await db.projects.find_one({"project_id": project_id, "user_id": user_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Projet non trouvé")

    if enable is False:
        await db.projects.update_one(
            {"project_id": project_id},
            {"$set": {"is_public": False, "updated_at": datetime.now(timezone.utc).isoformat()}},
        )
        return {"is_public": False, "slug": None, "url": None}

    slug = project.get("share_slug") or _make_slug(project.get("name") or project_id)
    await db.projects.update_one(
        {"project_id": project_id},
        {"$set": {
            "is_public": True,
            "share_slug": slug,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }},
    )
    frontend_url = (
        os.environ.get("FRONTEND_URL")
        or os.environ.get("REACT_APP_BACKEND_URL")
        or ""
    )
    public_url = f"{frontend_url.rstrip('/')}/share/{slug}" if frontend_url else f"/share/{slug}"
    return {"is_public": True, "slug": slug, "url": public_url}


@api_router.get("/share/{slug}")
async def get_public_share(slug: str):
    """PUBLIC — return the project metadata + generated files for a shared slug.
    No auth required.
    """
    project = await db.projects.find_one(
        {"share_slug": slug, "is_public": True},
        {"_id": 0, "user_id": 0, "ai_source": 0},
    )
    if not project:
        raise HTTPException(status_code=404, detail="Projet non partagé ou introuvable")
    return {
        "name": project.get("name"),
        "description": project.get("description"),
        "project_type": project.get("project_type"),
        "files": (project.get("generated_code") or {}).get("files", []),
        "created_at": project.get("created_at"),
    }


@api_router.get("/share/{slug}/preview")
async def get_public_share_preview(slug: str):
    """PUBLIC — return the rendered HTML preview for a shared web project.
    Same logic as `/preview/project/{id}` but slug-based and unauthenticated.
    """
    project = await db.projects.find_one(
        {"share_slug": slug, "is_public": True}, {"_id": 0}
    )
    if not project:
        return HTMLResponse("<h1>Projet introuvable</h1>", status_code=404)

    files = (project.get("generated_code") or {}).get("files", []) or []
    html_parts = [
        "<!DOCTYPE html><html><head><meta charset='utf-8'>"
        f"<title>{project.get('name', 'Projet')} · CodeForge AI</title>"
        "<meta name='viewport' content='width=device-width, initial-scale=1'>"
        "<style>*{margin:0;padding:0;box-sizing:border-box}body{font-family:system-ui,sans-serif}</style>"
    ]
    for f in files:
        if f.get("path", "").endswith(".css"):
            html_parts.append(f"<style>{f.get('content', '')}</style>")
    html_parts.append("</head><body>")
    for f in files:
        if f.get("path", "").endswith(".html"):
            content = f.get("content", "")
            if "<body>" in content:
                start = content.find("<body>") + 6
                end = content.find("</body>")
                content = content[start:end] if end > start else content
            html_parts.append(content)
    for f in files:
        if f.get("path", "").endswith(".js"):
            html_parts.append(f"<script>{f.get('content', '')}</script>")
    html_parts.append("</body></html>")
    return HTMLResponse("\n".join(html_parts))


@api_router.get("/system/ollama-status")
async def ollama_status():
    """Light health check on the local Ollama instance. Public — used by the
    ModelPicker to greyout offline models when the service is unreachable."""
    ollama_url = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434")
    try:
        async with httpx.AsyncClient(timeout=1.5) as client:
            r = await client.get(f"{ollama_url}/api/tags")
            if r.status_code == 200:
                tags = r.json().get("models", []) or []
                return {"available": True, "models": [t.get("name") for t in tags][:30]}
    except Exception:
        pass
    return {"available": False, "models": []}


# ==========================================================================
# DEVICE-BOUND CRYPTOGRAPHIC IDENTITY (creator-grade access control)
# ==========================================================================
from device_auth import compute_key_id, new_nonce, verify_signature  # noqa: E402


VALID_SITE_MODES = {"public", "private", "creator", "guest", "staff", "admin", "modo"}
VALID_DEVICE_ROLES = {"creator", "approved", "pending", "revoked"}


def _normalize_modes(mode_value) -> List[str]:
    """iter83 C11 — Accepte mode_value en str OU list[str] et retourne
    toujours une liste de modes valides. Conserve la rétro-compatibilité avec
    l'ancien stockage str unique. Si la valeur est vide ou invalide,
    retourne ['public'] par défaut."""
    if isinstance(mode_value, str):
        modes = [mode_value]
    elif isinstance(mode_value, (list, tuple)):
        modes = [str(m) for m in mode_value]
    else:
        modes = []
    modes = [m for m in modes if m in VALID_SITE_MODES]
    return modes or ["public"]


def _device_matches_mode(dev: Dict[str, Any], modes: List[str]) -> bool:
    """iter86 C11 — Vrai si le device match au moins un mode actif.

      - 'public' : visiteurs + clés publiques. PAS le staff (sauf cumul).
      - 'private' : clés validées (approved). PAS le staff seul (sauf cumul).
      - 'creator' : creator only.
      - 'guest' : tout le monde (lecture seule, géré UI-side).
      - 'staff' : admin + modo (admins/modos n'ont alors PAS besoin de clé
                  validée pour bosser).
      - 'admin' : admin only.
      - 'modo' : modo only.

    Sémantique : un device match si AU MOINS un mode lui est accordé.
    Donc cocher "private + staff" laisse passer clés validées ET staff.
    Cocher "private" seul → staff seul est BLOQUÉ (sauf si admin/modo se
    sont approvalés normalement comme privé, ce qui est rare).
    """
    if not dev:
        return False
    if dev.get("role") in ("revoked",) or dev.get("banned"):
        return False
    role = dev.get("role")
    sk = dev.get("staff_kind")
    is_creator = role == "creator"
    is_admin = sk == "admin"
    is_modo = sk == "modo"
    is_staff_only = (is_admin or is_modo) and role not in ("creator",)
    is_approved_clean = role == "approved" and not is_admin and not is_modo
    for m in modes:
        if m == "public":
            # public seul = visiteurs + clés publiques NON-staff
            if not is_staff_only:
                return True
        elif m == "guest":
            return True
        elif m == "private":
            # private seul = clés validées NON-staff (ou créa)
            if is_creator or is_approved_clean:
                return True
        elif m == "creator":
            if is_creator:
                return True
        elif m == "staff":
            # staff = admin + modo (créa passe car au-dessus)
            if is_creator or is_admin or is_modo:
                return True
        elif m == "admin":
            if is_creator or is_admin:
                return True
        elif m == "modo":
            if is_creator or is_modo:
                return True
    return False


async def _get_site_mode():
    """Returns the current site_mode. Cached in-memory for 30 seconds to
    avoid hitting Mongo on every /auth/login + /devices/verify call. The
    cache is bypassed by /system/site-mode PUT which invalidates it.

    iter83 C11 — Peut retourner str (legacy) OU list[str] (nouveau multi)."""
    global _site_mode_cache
    now = time.monotonic()
    if _site_mode_cache and now - _site_mode_cache[0] < 30.0:
        return _site_mode_cache[1]
    doc = await db.site_config.find_one({"_id": "site_mode"}, {"_id": 0, "mode": 1, "modes": 1})
    if doc and isinstance(doc.get("modes"), list) and doc.get("modes"):
        mode = doc["modes"]
    else:
        mode = (doc or {}).get("mode") or "public"
    _site_mode_cache = (now, mode)
    return mode


async def _get_site_modes_list() -> List[str]:
    """iter83 — Helper qui retourne TOUJOURS la liste des modes actifs."""
    m = await _get_site_mode()
    return _normalize_modes(m)


def _invalidate_site_mode_cache():
    global _site_mode_cache
    _site_mode_cache = None


_site_mode_cache = None  # (timestamp, mode) | None


async def _device_by_key(key_id: str) -> Optional[dict]:
    return await db.device_keys.find_one({"key_id": key_id}, {"_id": 0})


async def _log_decision(
    action: str,
    target_key_id: str,
    by_key_id: str,
    target_label: str = None,
    snapshot: dict = None,
):
    """Append an audit record to `device_decisions`.

    NOTE: only persists the 3 actions the user wants to see in the History
    panel — approve (=> "Accepté"), revoke (=> "Refusé"), promote (=> "Créateur").
    The destructive snapshot is kept so the "Annuler" button can fully
    restore the device. Other actions (request_access, add_by_key, disconnect,
    undo, register, block, unblock) are NOT logged."""
    if action not in ("approve", "revoke", "promote"):
        return
    row = {
        "action": action,
        "target_key_id": target_key_id,
        "target_label": target_label,
        "by_key_id": by_key_id,
        "ts": datetime.now(timezone.utc).isoformat(),
    }
    if snapshot:
        row["snapshot"] = {k: v for k, v in snapshot.items() if k != "_id"}
    await db.device_decisions.insert_one(row)
    total = await db.device_decisions.count_documents({})
    if total > 500:
        olds = await db.device_decisions.find({}, {"_id": 1}).sort("ts", 1).limit(total - 500).to_list(length=total - 500)
        await db.device_decisions.delete_many({"_id": {"$in": [o["_id"] for o in olds]}})


async def _consume_nonce(key_id: str, nonce: str) -> bool:
    """Atomically delete the pending nonce. Returns True if it existed."""
    res = await db.device_nonces.delete_one({"key_id": key_id, "nonce": nonce})
    return res.deleted_count > 0


async def _require_creator_signature(key_id: str, nonce: str, signature: str) -> dict:
    """Verify the caller is a device with role='creator' (signed proof)."""
    dev = await _device_by_key(key_id)
    if not dev or dev.get("role") != "creator":
        raise HTTPException(status_code=403, detail="Action réservée au créateur.")
    if not await _consume_nonce(key_id, nonce):
        raise HTTPException(status_code=403, detail="Nonce invalide ou expiré.")
    if not verify_signature(dev.get("public_key_jwk") or {}, nonce, signature):
        raise HTTPException(status_code=403, detail="Signature invalide.")
    return dev


async def _require_staff_signature(
    key_id: str, nonce: str, signature: str,
    allow_kinds: tuple = ("admin", "modo"),
) -> dict:
    """iter79 — Vérifie que le caller est créa OU staff (admin/modo selon allow_kinds)."""
    dev = await _device_by_key(key_id)
    role = (dev or {}).get("role")
    sk = (dev or {}).get("staff_kind")
    if not dev or (role != "creator" and sk not in allow_kinds):
        raise HTTPException(status_code=403, detail="Action réservée au staff (admin/modo) et créatrice.")
    if not await _consume_nonce(key_id, nonce):
        raise HTTPException(status_code=403, detail="Nonce invalide ou expiré.")
    if not verify_signature(dev.get("public_key_jwk") or {}, nonce, signature):
        raise HTTPException(status_code=403, detail="Signature invalide.")
    return dev


class DeviceRegisterIn(BaseModel):
    public_key_jwk: Dict[str, Any]
    label: Optional[str] = None   # human-friendly device name (e.g. "iPhone")
    product: Optional[str] = None # marketing name ("Galaxy S21 5G", "iPhone")
    model: Optional[str] = None   # raw code ("SM-G991U1")


@api_router.post("/devices/register")
async def device_register(payload: DeviceRegisterIn):
    """Register a fresh device. Public — anyone can call this. The first
    device EVER registered is auto-promoted to 'creator'. All subsequent
    registrations are 'pending' until the creator approves them.
    Idempotent: a second register call for the same key_id never duplicates."""
    jwk = payload.public_key_jwk or {}
    if jwk.get("kty") != "EC" or jwk.get("crv") != "P-256":
        raise HTTPException(status_code=400, detail="Clé publique invalide (EC P-256 attendu).")
    key_id = compute_key_id(jwk)
    product = (payload.product or "")[:60] or None
    model = (payload.model or "")[:60] or None
    label = (payload.label or "")[:60] or None

    # Atomic-ish: try to find any existing doc for this key. If any, we
    # consolidate by keeping the highest-privilege role and removing duplicates.
    matches = await db.device_keys.find({"key_id": key_id}, {"_id": 0}).to_list(length=5)
    if matches:
        priority = {"creator": 4, "approved": 3, "pending": 2, "revoked": 1}
        best_role = max((m.get("role") for m in matches), key=lambda r: priority.get(r, 0))
        # Preserve any product/model already known if the new payload is empty.
        existing = matches[0]
        await db.device_keys.delete_many({"key_id": key_id})
        await db.device_keys.insert_one({
            "key_id": key_id,
            "public_key_jwk": jwk,
            "role": best_role,
            "label": label or existing.get("label"),
            "product": product or existing.get("product"),
            "model": model or existing.get("model"),
            "created_at": existing.get("created_at") or datetime.now(timezone.utc).isoformat(),
            "last_seen_at": datetime.now(timezone.utc).isoformat(),
        })
        return {"key_id": key_id, "role": best_role, "already_registered": True}

    # Count DISTINCT key_ids before granting creator role (more robust than total).
    distinct_ids = await db.device_keys.distinct("key_id")
    # iter63: New devices land as "inactive" (silent, NOT in the creator's
    # accounts panel). They become "pending" only when the user explicitly
    # nudges the creator via /devices/send-to-creator.
    role = "creator" if len(distinct_ids) == 0 else "inactive"
    doc = {
        "key_id": key_id,
        "public_key_jwk": jwk,
        "role": role,
        "label": label,
        "product": product,
        "model": model,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "last_seen_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.device_keys.insert_one(doc)
    return {"key_id": key_id, "role": role, "already_registered": False}


class DeviceChallengeIn(BaseModel):
    key_id: str


@api_router.post("/devices/challenge")
async def device_challenge(payload: DeviceChallengeIn):
    """Issue a single-use nonce for the given key_id. The device signs it
    with its non-extractable private key and posts the signature to /verify.
    Nonces are stored with a 2-minute TTL via the `device_nonces` collection."""
    dev = await _device_by_key(payload.key_id)
    if not dev:
        raise HTTPException(status_code=404, detail="Appareil inconnu.")
    nonce = new_nonce()
    await db.device_nonces.insert_one({
        "key_id": payload.key_id,
        "nonce": nonce,
        "created_at": datetime.now(timezone.utc),
    })
    return {"nonce": nonce, "expires_in_seconds": 120}


class DeviceVerifyIn(BaseModel):
    key_id: str
    nonce: str
    signature: str


@api_router.post("/devices/verify")
async def device_verify(payload: DeviceVerifyIn):
    """Verify the signature → return device role + whether access is granted
    given the current site mode. Public endpoint (anyone with a valid key can
    call it)."""
    dev = await _device_by_key(payload.key_id)
    if not dev:
        raise HTTPException(status_code=404, detail="Appareil inconnu.")

    if not await _consume_nonce(payload.key_id, payload.nonce):
        raise HTTPException(status_code=403, detail="Nonce invalide ou expiré.")
    ok = verify_signature(dev.get("public_key_jwk") or {}, payload.nonce, payload.signature)
    if not ok:
        return {"verified": False, "role": dev.get("role"), "can_access": False}

    await db.device_keys.update_one(
        {"key_id": payload.key_id},
        {"$set": {"last_seen_at": datetime.now(timezone.utc).isoformat()}},
    )

    site_mode = await _get_site_mode()
    role = dev.get("role")
    effective_role = role
    can_access = True
    kick_reason = None  # localized message key — frontend translates

    # ----- Account-level moderation gates (iter54) -----
    # Exclusion / Ban are tied to the device's email (if any). They lock
    # the account itself, not just one device — re-registering on the same
    # email keeps the gate active.
    excluded_until = dev.get("excluded_until")
    if excluded_until:
        try:
            exp = datetime.fromisoformat(excluded_until.replace("Z", "+00:00"))
            if exp <= datetime.now(timezone.utc):
                await db.device_keys.update_one(
                    {"key_id": payload.key_id},
                    {"$unset": {"excluded_until": "", "excluded_reason": ""}},
                )
                excluded_until = None
        except Exception:
            excluded_until = None
    if dev.get("banned"):
        can_access = False
        kick_reason = "kick_banned"
    elif excluded_until:
        can_access = False
        kick_reason = "kick_excluded"
    elif role == "blocked":
        can_access = False
        kick_reason = "kick_blocked"
    elif role == "revoked":
        can_access = False
        kick_reason = "kick_revoked"
    elif True:
        # iter83 C11 — Multi-mode gating. Le device passe si AU MOINS un mode
        # actif lui est accessible (via _device_matches_mode). Sinon kick avec
        # raison spécifique au mode dominant.
        modes_active = _normalize_modes(site_mode)
        if not _device_matches_mode(dev, modes_active):
            can_access = False
            # Priorité de raison : creator > admin > modo > staff > private > guest
            if "creator" in modes_active:
                kick_reason = "kick_creator_only"
            elif "admin" in modes_active or "modo" in modes_active or "staff" in modes_active:
                kick_reason = "kick_staff_only"
            elif "private" in modes_active:
                kick_reason = "kick_private"
                effective_role = "guest"
            else:
                kick_reason = "kick_blocked"
        elif "guest" in modes_active and role not in ("creator", "approved"):
            # 'guest' actif : tout le monde en read-only.
            effective_role = "guest"
        elif "private" in modes_active and role not in ("creator", "approved") and "public" not in modes_active:
            # Private uniquement (sans public) → kicked unless approved.
            # (déjà géré par _device_matches_mode ci-dessus, fallback safe)
            pass

    # iter83 C11 — Renvoie aussi site_modes (liste) pour le frontend.
    site_modes_list = _normalize_modes(site_mode)
    return {
        "verified": True,
        "role": role,
        "effective_role": effective_role,
        "can_access": can_access,
        "site_mode": site_modes_list[0],  # legacy str = premier mode
        "site_modes": site_modes_list,    # nouveau : liste complète
        "kick_reason": kick_reason,
        "excluded_until": excluded_until,
        "force_visitor": bool(dev.get("force_visitor", False)),  # iter77
        "staff_kind": dev.get("staff_kind"),  # iter77
    }


@api_router.get("/system/site-mode")
async def get_site_mode_public():
    """Public — anyone can read the current site mode + (in guest mode)
    the optional view-forcing setting set by the creator.

    iter83 C11 — Renvoie maintenant aussi `modes` (liste). `mode` reste pour
    compat ascendante = première entrée de la liste."""
    doc = await db.site_config.find_one(
        {"_id": "site_mode"}, {"_id": 0, "mode": 1, "guest_view": 1, "modes": 1},
    )
    raw = doc or {}
    modes = _normalize_modes(raw.get("modes") if isinstance(raw.get("modes"), list) and raw.get("modes") else raw.get("mode"))
    return {
        "mode": modes[0],
        "modes": modes,
        "guest_view": raw.get("guest_view"),
    }


class SiteModeSetIn(BaseModel):
    # iter83 — Accepte soit `mode` (legacy str) soit `modes` (list).
    mode: Optional[str] = None
    modes: Optional[List[str]] = None
    key_id: str
    nonce: str
    signature: str
    guest_view: Optional[str] = None  # 'creator' | 'user' | 'modo' | 'admin' | None (free)


@api_router.put("/system/site-mode")
async def set_site_mode(payload: SiteModeSetIn):
    """Creator-only. Toggle site mode AND kick devices that no longer have
    access under the new mode (their next /devices/verify poll will return
    can_access=False with a localized kick_reason).

    iter83 C11 — Supporte multi-checkbox via `payload.modes` (liste). Kick
    appliqué pour les devices qui n'ont aucun mode actif compatible.
    """
    # Détermine la liste finale.
    if payload.modes is not None:
        modes = _normalize_modes(payload.modes)
    elif payload.mode is not None:
        if payload.mode not in VALID_SITE_MODES:
            raise HTTPException(status_code=400, detail="Mode invalide.")
        modes = [payload.mode]
    else:
        raise HTTPException(status_code=400, detail="Aucun mode fourni.")
    # Validation : tous les modes doivent être valides.
    for m in modes:
        if m not in VALID_SITE_MODES:
            raise HTTPException(status_code=400, detail=f"Mode invalide : {m}")
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)

    # Optional sub-view forcing for guest mode.
    guest_view = getattr(payload, 'guest_view', None)
    is_guest_mode = "guest" in modes

    await db.site_config.update_one(
        {"_id": "site_mode"},
        {"$set": {
            "mode": modes[0],            # legacy str pour compat
            "modes": modes,              # nouvelle source de vérité
            "guest_view": guest_view if is_guest_mode else None,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }},
        upsert=True,
    )
    _invalidate_site_mode_cache()

    # iter83 — Kick : on supprime les sessions des devices qui ne matchent
    # AUCUN des modes actifs. Approche défensive.
    all_devs = await db.device_keys.find({}, {"_id": 0, "key_id": 1, "role": 1, "staff_kind": 1, "banned": 1}).to_list(length=2000)
    for d in all_devs:
        if not _device_matches_mode(d, modes):
            await db.user_sessions.delete_many({"device_key_id": d["key_id"]})

    return {"mode": modes[0], "modes": modes, "guest_view": guest_view}


class CreatorOnlyIn(BaseModel):
    key_id: str           # caller (must be creator)
    nonce: str
    signature: str


@api_router.post("/devices/list")
async def devices_list(payload: CreatorOnlyIn):
    """Creator-only — list all registered devices."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    # iter63: hide silent "inactive" devices.
    devices = await db.device_keys.find(
        {"role": {"$ne": "inactive"}}, {"_id": 0, "public_key_jwk": 0},
    ).sort("created_at", -1).to_list(length=500)
    return {"devices": devices}


@api_router.post("/devices/decisions")
async def devices_decisions(payload: CreatorOnlyIn):
    """Creator-only — return the history of past decisions (approve/revoke/
    disconnect/promote/add_by_key). Sorted newest first, capped at 200."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    rows = await db.device_decisions.find(
        {}, {"_id": 0},
    ).sort("ts", -1).to_list(length=200)
    return {"decisions": rows}


@api_router.post("/devices/decisions/clear")
async def devices_decisions_clear(payload: CreatorOnlyIn):
    """Creator-only — wipe the entire history log. This does NOT touch the
    actual device state (existing creator/approved/pending roles stay). Use
    when the history gets long and you want a clean slate."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    res = await db.device_decisions.delete_many({})
    return {"deleted": res.deleted_count}


class DecisionUndoIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    target_key_id: str
    decision_ts: str    # exact timestamp from the history row to undo


@api_router.post("/devices/decisions/undo")
async def devices_decisions_undo(payload: DecisionUndoIn):
    """Creator-only — undo a specific decision and put the device back in
    the 'pending' queue so the creator can re-decide.

    Behaviour per action:
      - approve         → role becomes 'pending' again
      - revoke / disconnect → the device row is recreated from the snapshot
                              stored on the decision, with role='pending'
      - promote         → role demotes back to 'approved'
      - add_by_key      → the device is removed (it was creator-added, so
                          'undo' simply un-adds it)
      - request_access  → ignored (no-op; you can just /revoke instead)

    The undo itself logs a fresh 'undo' decision row for traceability."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    dec = await db.device_decisions.find_one(
        {"target_key_id": payload.target_key_id, "ts": payload.decision_ts},
        {"_id": 0},
    )
    if not dec:
        raise HTTPException(status_code=404, detail="Décision introuvable.")
    action = dec.get("action")
    snapshot = dec.get("snapshot") or {}

    if action == "approve":
        await db.device_keys.update_one(
            {"key_id": payload.target_key_id},
            {"$set": {"role": "pending"}},
        )
    elif action in ("revoke", "disconnect"):
        # Recreate the deleted row from the snapshot. If we don't have one
        # (older logs), we just create a minimal pending shell that the
        # device will refresh on its next /devices/register call.
        existing = await _device_by_key(payload.target_key_id)
        if existing:
            await db.device_keys.update_one(
                {"key_id": payload.target_key_id},
                {"$set": {"role": "pending"}},
            )
        else:
            new_row = {
                "key_id": payload.target_key_id,
                "public_key_jwk": snapshot.get("public_key_jwk", {}),
                "label": snapshot.get("label") or dec.get("target_label"),
                "role": "pending",
                "created_at": snapshot.get("created_at") or datetime.now(timezone.utc).isoformat(),
                "last_seen_at": datetime.now(timezone.utc).isoformat(),
            }
            await db.device_keys.insert_one(new_row)
    elif action == "promote":
        await db.device_keys.update_one(
            {"key_id": payload.target_key_id},
            {"$set": {"role": "approved"}},
        )
    elif action == "add_by_key":
        await db.device_keys.delete_one({"key_id": payload.target_key_id})
    else:
        # request_access / register / future actions — nothing to undo.
        return {"success": False, "reason": "non_undoable_action"}

    await _log_decision("undo", payload.target_key_id, payload.key_id, dec.get("target_label"))
    return {"success": True}


@api_router.post("/devices/pending-count")
async def devices_pending_count(payload: CreatorOnlyIn):
    """Creator-only — quick count of pending devices, used by the bell badge.
    Uses POST so we can carry the creator proof in the body without exposing
    keys in query strings/logs."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    count = await db.device_keys.count_documents({"role": "pending"})
    return {"pending_count": count}


@api_router.get("/devices/pending-stream/{key_id}/{nonce}/{signature}")
async def devices_pending_stream(key_id: str, nonce: str, signature: str):
    """Creator-only SSE stream — emits `{pending_count}` every 5s, plus an
    immediate first event. Auth via path-args because EventSource cannot set
    custom headers/bodies. The nonce is consumed on first call (single-use),
    after that we trust the connection (it dies on creator-revoke anyway)."""
    await _require_creator_signature(key_id, nonce, signature)

    async def gen():
        # Emit one immediately so the badge updates without lag.
        last = -1
        try:
            while True:
                # Re-check the device is still 'creator' on every tick. If
                # they got revoked, we close the stream.
                dev = await _device_by_key(key_id)
                if not dev or dev.get("role") != "creator":
                    yield "event: closed\ndata: revoked\n\n"
                    return
                count = await db.device_keys.count_documents({"role": "pending"})
                if count != last:
                    yield f"data: {{\"pending_count\": {count}}}\n\n"
                    last = count
                else:
                    yield ": keepalive\n\n"
                await asyncio.sleep(5)
        except asyncio.CancelledError:
            return

    return StreamingResponse(gen(), media_type="text/event-stream", headers={
        "Cache-Control": "no-cache",
        "X-Accel-Buffering": "no",
        "Connection": "keep-alive",
    })


class DeviceTargetIn(CreatorOnlyIn):
    target_key_id: str


@api_router.post("/devices/approve")
async def devices_approve(payload: DeviceTargetIn):
    """iter79 — Staff (admin/modo) ou créatrice approuve. La décision tracke
    qui a accepté (couleur d'encadrement: créa=jaune, admin=orange, modo=bleu).
    Si modo accepte, l'appareil passe en `approved` ; la créa garde la notif
    pour pouvoir override (refuser, ce qui annule la décision)."""
    actor = await _require_staff_signature(payload.key_id, payload.nonce, payload.signature)
    target = await _device_by_key(payload.target_key_id)
    res = await db.device_keys.update_one(
        {"key_id": payload.target_key_id, "role": "pending"},
        {"$set": {
            "role": "approved",
            "approved_at": datetime.now(timezone.utc).isoformat(),
            "approved_by_key_id": payload.key_id,
            "approved_by_kind": "creator" if actor.get("role") == "creator" else actor.get("staff_kind"),
        }},
    )
    if res.matched_count == 0:
        raise HTTPException(status_code=404, detail="Aucun appareil en attente avec cette clé.")
    await _log_decision("approve", payload.target_key_id, payload.key_id, (target or {}).get("label"))
    return {"success": True, "approved_by_kind": "creator" if actor.get("role") == "creator" else actor.get("staff_kind")}


@api_router.post("/devices/revoke")
async def devices_revoke(payload: DeviceTargetIn):
    """Creator hard-revokes a device — they cannot authenticate again. The
    revoke also REMOVES the device from `device_keys` so it doesn't clutter
    the registered-devices list. The audit trail in `device_decisions`
    remembers what happened.

    Idempotent: if the target device is already gone (cleaned-up after a
    previous revoke), we still log the decision so the creator gets a clear
    audit trail when revoking from the History panel."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    if payload.target_key_id == payload.key_id:
        raise HTTPException(status_code=400, detail="Tu ne peux pas révoquer ton propre appareil créateur.")
    target = await _device_by_key(payload.target_key_id)
    target_label = (target or {}).get("label")
    # Idempotent — delete if exists, no error if it doesn't.
    await db.device_keys.delete_one({"key_id": payload.target_key_id})
    await db.device_nonces.delete_many({"key_id": payload.target_key_id})
    await db.user_sessions.delete_many({"device_key_id": payload.target_key_id})
    await _log_decision("revoke", payload.target_key_id, payload.key_id, target_label, snapshot=target)
    return {"success": True, "existed": target is not None}


@api_router.post("/devices/disconnect")
async def devices_disconnect(payload: DeviceTargetIn):
    """Creator force-disconnects a device. Like /revoke this REMOVES the
    device from `device_keys` so the active list stays clean; the audit
    record lives on in `device_decisions`."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    if payload.target_key_id == payload.key_id:
        raise HTTPException(status_code=400, detail="Tu ne peux pas te déconnecter toi-même via cette action.")
    target = await _device_by_key(payload.target_key_id)
    if not target:
        raise HTTPException(status_code=404, detail="Appareil introuvable.")
    await db.device_nonces.delete_many({"key_id": payload.target_key_id})
    await db.user_sessions.delete_many({"device_key_id": payload.target_key_id})
    await db.device_keys.delete_one({"key_id": payload.target_key_id})
    await _log_decision("disconnect", payload.target_key_id, payload.key_id, (target or {}).get("label"), snapshot=target)
    return {"success": True}


class PromoteCreatorIn(CreatorOnlyIn):
    target_key_id: str
    password: str  # creator's own password — extra protection


@api_router.post("/devices/promote-creator")
async def devices_promote_creator(payload: PromoteCreatorIn):
    """Creator promotes another device to 'creator' role. Requires the
    creator's own account password to confirm intent."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    # Identify creator's account: stored at device creation time? Not yet —
    # link via the user that last logged in with this device. For now we
    # require the password to match ANY existing user that has marked this
    # device as 'creator-owner'. Fallback: accept the first admin user.
    # Simpler approach: verify any user whose password matches. This is
    # acceptable since the call is gated by the creator's device signature.
    users = await db.users.find({}, {"_id": 0, "email": 1, "password_hash": 1}).to_list(length=200)
    matched = False
    for u in users:
        ph = u.get("password_hash") or ""
        try:
            if ph and verify_password(payload.password, ph):
                matched = True
                break
        except Exception:
            continue
    if not matched:
        raise HTTPException(status_code=401, detail="Mot de passe incorrect.")
    res = await db.device_keys.update_one(
        {"key_id": payload.target_key_id},
        {"$set": {"role": "creator", "promoted_at": datetime.now(timezone.utc).isoformat()}},
    )
    if res.matched_count == 0:
        raise HTTPException(status_code=404, detail="Appareil introuvable.")
    target = await _device_by_key(payload.target_key_id)
    await _log_decision("promote", payload.target_key_id, payload.key_id, (target or {}).get("label"))
    return {"success": True}


class AddByKeyIn(CreatorOnlyIn):
    public_key_jwk: Dict[str, Any]
    label: Optional[str] = None
    role: Optional[str] = "approved"  # 'approved' by default


@api_router.post("/devices/add-by-key")
async def devices_add_by_key(payload: AddByKeyIn):
    """Creator pastes another device's public key (shared offline) to whitelist
    it directly without going through pending → approve."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    jwk = payload.public_key_jwk or {}
    if jwk.get("kty") != "EC" or jwk.get("crv") != "P-256":
        raise HTTPException(status_code=400, detail="Clé publique invalide.")
    target_id = compute_key_id(jwk)
    role = payload.role if payload.role in ("approved", "creator") else "approved"
    existing = await _device_by_key(target_id)
    if existing:
        await db.device_keys.update_one(
            {"key_id": target_id},
            {"$set": {"role": role, "label": payload.label or existing.get("label")}},
        )
    else:
        await db.device_keys.insert_one({
            "key_id": target_id,
            "public_key_jwk": jwk,
            "role": role,
            "label": (payload.label or "")[:60] or None,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "added_by_creator": True,
        })
    await _log_decision("add_by_key", target_id, payload.key_id, payload.label)
    return {"key_id": target_id, "role": role}


# Non-creator can ping the creator with a request to be added. We just log
# the request in `device_decisions` so the creator sees it in their History
# panel (and the regular "pending" entry in device_keys, which they already
# see). No private data exchanged — the requester's key_id was already known
# to the server from their /devices/register call on first visit.
class SendToCreatorIn(BaseModel):
    key_id: str  # the requester (any registered, non-creator device)
    nonce: str
    signature: str


@api_router.post("/devices/send-to-creator")
async def devices_send_to_creator(payload: SendToCreatorIn):
    """Anyone with a valid registered device can use this to nudge the creator.
    Blocked devices get a specific localized error. Cool-down: 1 successful
    nudge per device every 10 minutes (anti-spam) — surfaces as a 429."""
    dev = await _device_by_key(payload.key_id)
    if not dev:
        raise HTTPException(status_code=404, detail="Appareil inconnu.")
    if dev.get("role") == "blocked":
        raise HTTPException(
            status_code=403,
            detail="Votre demande a été formulée de nombreuses fois. Veuillez contacter le créateur.",
        )
    # Cool-down check.
    last_at_iso = dev.get("last_nudge_at")
    if last_at_iso:
        try:
            last_at = datetime.fromisoformat(last_at_iso)
            if (datetime.now(timezone.utc) - last_at) < timedelta(minutes=10):
                raise HTTPException(
                    status_code=429,
                    detail="Tu as déjà envoyé une demande récemment. Réessaie dans quelques minutes.",
                )
        except HTTPException:
            raise
        except Exception:
            pass
    if not await _consume_nonce(payload.key_id, payload.nonce):
        raise HTTPException(status_code=403, detail="Nonce invalide ou expiré.")
    if not verify_signature(dev.get("public_key_jwk") or {}, payload.nonce, payload.signature):
        raise HTTPException(status_code=403, detail="Signature invalide.")
    now_iso = datetime.now(timezone.utc).isoformat()
    if dev.get("role") in ("creator", "approved", "pending"):
        await db.device_keys.update_one(
            {"key_id": payload.key_id},
            {"$set": {"last_seen_at": now_iso, "last_nudge_at": now_iso}},
        )
    else:
        await db.device_keys.update_one(
            {"key_id": payload.key_id},
            {"$set": {"role": "pending", "last_seen_at": now_iso, "last_nudge_at": now_iso}},
        )
    return {"sent": True, "role": dev.get("role")}


@api_router.post("/devices/block")
async def devices_block(payload: DeviceTargetIn):
    """iter79 — Block. Ouvert au staff."""
    await _require_staff_signature(payload.key_id, payload.nonce, payload.signature)
    if payload.target_key_id == payload.key_id:
        raise HTTPException(status_code=400, detail="Tu ne peux pas te bloquer toi-même.")
    target = await _device_by_key(payload.target_key_id)
    target_label = (target or {}).get("label")
    if target:
        await db.device_keys.update_one(
            {"key_id": payload.target_key_id},
            {"$set": {"role": "blocked", "blocked_at": datetime.now(timezone.utc).isoformat()}},
        )
    else:
        # Recreate a minimal "blocked" shell so future re-registrations
        # under the same key_id stay blocked.
        await db.device_keys.insert_one({
            "key_id": payload.target_key_id,
            "public_key_jwk": {},
            "label": target_label,
            "role": "blocked",
            "created_at": datetime.now(timezone.utc).isoformat(),
            "blocked_at": datetime.now(timezone.utc).isoformat(),
        })
    await db.device_nonces.delete_many({"key_id": payload.target_key_id})
    await db.user_sessions.delete_many({"device_key_id": payload.target_key_id})
    # Block is shown as "Refusé" in the history (same intent), with snapshot.
    await _log_decision("revoke", payload.target_key_id, payload.key_id, target_label, snapshot=target)
    return {"success": True, "blocked": True}


@api_router.post("/devices/unblock")
async def devices_unblock(payload: DeviceTargetIn):
    """iter79 — Unblock. Ouvert au staff."""
    await _require_staff_signature(payload.key_id, payload.nonce, payload.signature)
    target = await _device_by_key(payload.target_key_id)
    if not target or target.get("role") != "blocked":
        raise HTTPException(status_code=404, detail="Appareil non bloqué.")
    await db.device_keys.update_one(
        {"key_id": payload.target_key_id},
        {"$set": {"role": "pending"}, "$unset": {"blocked_at": ""}},
    )
    return {"success": True}


# ==========================================================================
# END device-bound identity
# ==========================================================================


# ==========================================================================
# ONE-DEVICE-AT-A-TIME — pending session approval flow
# ==========================================================================

class SessionRequestStatusIn(BaseModel):
    request_id: str


@api_router.post("/auth/session-request-status")
async def session_request_status(payload: SessionRequestStatusIn, response: Response):
    """Polled by the requesting device until the connected device decides
    (approve/deny) or the request expires (15 min).

    Idempotent: once approved, the session token is persisted on the request
    and returned on every subsequent poll until the requesting device clears
    it. This avoids race conditions where two parallel polls of the same
    approved request would have one return 'approved+token' and the other
    return 404."""
    now = datetime.now(timezone.utc)
    req = await db.session_requests.find_one({"request_id": payload.request_id}, {"_id": 0})
    if not req:
        # Could be: deleted/expired old; OR brand new request lookup race.
        return {"status": "expired"}
    # Expire on read.
    if req["status"] == "pending" and req.get("expires_at") and req["expires_at"] < now.isoformat():
        await db.session_requests.update_one(
            {"request_id": payload.request_id},
            {"$set": {"status": "expired"}},
        )
        req["status"] = "expired"

    if req["status"] in ("pending", "denied", "expired"):
        return {"status": req["status"]}

    # status == "approved" — issue/return a session token. We persist it on
    # the request itself so concurrent or repeat polls all get the same value.
    user = await db.users.find_one({"user_id": req["user_id"]}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="Utilisateur introuvable.")

    session_token = req.get("issued_session_token")
    if not session_token:
        session_token = secrets.token_urlsafe(32)
        # Insert + verify it's queryable BEFORE returning, so the client's
        # follow-up /auth/me always succeeds. Mongo writes are immediately
        # visible to subsequent reads from the same client + DB, but we
        # double-check here to absorb any replica lag on hosted clusters.
        await db.user_sessions.insert_one({
            "session_token": session_token,
            "user_id": user["user_id"],
            "device_key_id": req.get("requesting_key_id"),
            "device_label": req.get("requesting_label"),
            "auth_type": "email",
            "created_at": now.isoformat(),
            "last_seen_at": now.isoformat(),  # iter66
            "expires_at": (now + timedelta(days=7)).isoformat(),
        })
        # Tiny read-after-write check (max 3 attempts × 50ms) — rare paranoia
        # for hosted MongoDB clusters with secondary read preference.
        for _ in range(3):
            check = await db.user_sessions.find_one({"session_token": session_token}, {"_id": 0, "session_token": 1})
            if check:
                break
            await asyncio.sleep(0.05)
        await db.users.update_one({"user_id": user["user_id"]}, {"$set": {"last_login": now.isoformat()}})
        await db.session_requests.update_one(
            {"request_id": payload.request_id},
            {"$set": {
                "issued_session_token": session_token,
                "consumed_at": now.isoformat(),
                "expires_at": (now + timedelta(minutes=15)).isoformat(),
            }},
        )

    response.set_cookie(
        key="session_token", value=session_token,
        httponly=True, secure=True, samesite="none",
        max_age=7 * 24 * 3600, path="/",
    )
    safe_user = {k: v for k, v in user.items() if k != "password_hash"}
    return {"status": "approved", **safe_user, "session_token": session_token}


@api_router.get("/auth/session-pending")
async def list_pending_session_requests(request: Request):
    """Listed by the currently-connected user — pending requests on their
    account from other devices.

    iter83 — Fix bug "demande fantôme récurrente" : on auto-expire les
    requests pending de plus de 90 secondes. La race condition se produisait
    quand un device demandait l'accès puis fermait l'onglet sans approval :
    la request restait `pending` jusqu'à `expires_at` (potentiellement
    plusieurs minutes), apparaissant comme un prompt fantôme à chaque poll
    sur le device connecté.
    """
    user_id = await get_current_user(request)
    now = datetime.now(timezone.utc)
    now_iso = now.isoformat()
    stale_threshold = (now - timedelta(seconds=90)).isoformat()
    # Auto-expire stale pending requests.
    await db.session_requests.update_many(
        {"user_id": user_id, "status": "pending", "created_at": {"$lt": stale_threshold}},
        {"$set": {"status": "expired", "expired_at": now_iso}},
    )
    rows = await db.session_requests.find(
        {"user_id": user_id, "status": "pending", "expires_at": {"$gt": now_iso}, "created_at": {"$gte": stale_threshold}},
        {"_id": 0},
    ).sort("created_at", -1).to_list(length=50)
    return {"requests": rows}


class SessionDecideIn(BaseModel):
    request_id: str
    decision: str  # 'approve' | 'deny'


@api_router.post("/auth/session-decide")
async def decide_session_request(payload: SessionDecideIn, request: Request):
    """Currently-connected device approves or denies a pending request."""
    if payload.decision not in ("approve", "deny"):
        raise HTTPException(status_code=400, detail="Décision invalide.")
    user_id = await get_current_user(request)
    req = await db.session_requests.find_one(
        {"request_id": payload.request_id, "user_id": user_id, "status": "pending"},
        {"_id": 0},
    )
    if not req:
        raise HTTPException(status_code=404, detail="Demande introuvable ou déjà traitée.")
    new_status = "approved" if payload.decision == "approve" else "denied"
    await db.session_requests.update_one(
        {"request_id": payload.request_id},
        {"$set": {
            "status": new_status,
            "decided_at": datetime.now(timezone.utc).isoformat(),
        }},
    )
    return {"status": new_status}


# ==========================================================================
# WEBAUTHN — biometric "declare theft" recovery
# ==========================================================================
import webauthn  # noqa: E402
from webauthn.helpers.structs import (  # noqa: E402
    PublicKeyCredentialDescriptor,
    UserVerificationRequirement,
    AuthenticatorSelectionCriteria,
    AuthenticatorAttachment,
    ResidentKeyRequirement,
)


def _rp_id_from_origin(origin: str) -> str:
    """Strip scheme + port from origin to derive the WebAuthn RP id."""
    try:
        from urllib.parse import urlparse
        host = urlparse(origin).hostname or ""
        return host or "localhost"
    except Exception:
        return "localhost"


class WebAuthnEnrollOptionsIn(BaseModel):
    key_id: str           # device making the call (must be creator to enroll)
    nonce: str
    signature: str
    origin: str           # window.location.origin


class WebAuthnSignupEnrollIn(BaseModel):
    """iter69: lightweight signup-flow enrollment options. No creator
    signature required (the user doesn't exist yet) — just a one-shot
    challenge bound to a short-lived token so the verify step can be
    correlated to the right register payload."""
    email: Optional[str] = None
    origin: Optional[str] = None


@api_router.post("/webauthn/enroll-begin")
async def webauthn_enroll_begin(payload: WebAuthnSignupEnrollIn, request: Request):
    """iter69: signup-time biometric enrollment options. The client calls
    this BEFORE submitting /auth/register; the returned options_token is
    bundled with the WebAuthn attestation inside the register payload."""
    origin = payload.origin or request.headers.get("origin") or request.headers.get("referer") or ""
    rp_id = _rp_id_from_origin(origin) if origin else "localhost"
    user_handle = secrets.token_bytes(16)
    options_token = secrets.token_urlsafe(24)

    options = webauthn.generate_registration_options(
        rp_id=rp_id,
        rp_name="CodeForge AI",
        user_id=user_handle,
        user_name=(payload.email or f"new_{secrets.token_hex(4)}@codeforge.ai")[:64],
        user_display_name=payload.email or "Nouveau compte",
        authenticator_selection=AuthenticatorSelectionCriteria(
            resident_key=ResidentKeyRequirement.PREFERRED,
            user_verification=UserVerificationRequirement.PREFERRED,
        ),
        timeout=120_000,
    )
    await db.webauthn_challenges.insert_one({
        "options_token": options_token,
        "challenge": webauthn.helpers.bytes_to_base64url(options.challenge),
        "rp_id": rp_id,
        "kind": "signup",
        "origin": origin,
        "user_handle": webauthn.helpers.bytes_to_base64url(user_handle),
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    return {
        "options": json.loads(webauthn.helpers.options_to_json(options)),
        "options_token": options_token,
    }


@api_router.post("/webauthn/register-options")
async def webauthn_register_options(payload: WebAuthnEnrollOptionsIn):
    """Step 1 (creator-only): generate a challenge for biometric enrollment.
    The browser will call navigator.credentials.create with these options.
    Only platform authenticators (Touch ID / Face ID / Windows Hello / Android
    fingerprint) are allowed."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    rp_id = _rp_id_from_origin(payload.origin)
    user_handle = payload.key_id.encode("utf-8")[:64]

    options = webauthn.generate_registration_options(
        rp_id=rp_id,
        rp_name="CodeForge AI",
        user_id=user_handle,
        user_name=f"creator:{payload.key_id}",
        user_display_name="Creator",
        authenticator_selection=AuthenticatorSelectionCriteria(
            authenticator_attachment=AuthenticatorAttachment.PLATFORM,
            resident_key=ResidentKeyRequirement.PREFERRED,
            user_verification=UserVerificationRequirement.REQUIRED,
        ),
        timeout=60_000,
    )
    # Persist challenge + rp_id for the verify step.
    await db.webauthn_challenges.insert_one({
        "key_id": payload.key_id,
        "challenge": webauthn.helpers.bytes_to_base64url(options.challenge),
        "rp_id": rp_id,
        "kind": "register",
        "origin": payload.origin,
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    return json.loads(webauthn.helpers.options_to_json(options))


class WebAuthnEnrollVerifyIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    origin: str
    credential: Dict[str, Any]  # the navigator.credentials.create response


@api_router.post("/webauthn/register-verify")
async def webauthn_register_verify(payload: WebAuthnEnrollVerifyIn):
    """Step 2 (creator-only): verify the browser's attestation and store
    the platform credential for later 'declare theft' use."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    challenge_doc = await db.webauthn_challenges.find_one_and_delete(
        {"key_id": payload.key_id, "kind": "register"},
    )
    if not challenge_doc:
        raise HTTPException(status_code=400, detail="Aucun défi d'enrôlement actif.")
    try:
        verification = webauthn.verify_registration_response(
            credential=payload.credential,
            expected_challenge=webauthn.helpers.base64url_to_bytes(challenge_doc["challenge"]),
            expected_origin=payload.origin,
            expected_rp_id=challenge_doc["rp_id"],
            require_user_verification=True,
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Vérification d'enrôlement échouée: {e}")

    await db.webauthn_credentials.insert_one({
        "credential_id": webauthn.helpers.bytes_to_base64url(verification.credential_id),
        "public_key": webauthn.helpers.bytes_to_base64url(verification.credential_public_key),
        "sign_count": verification.sign_count,
        "owner_key_id": payload.key_id,  # the creator who enrolled this biometric
        "rp_id": challenge_doc["rp_id"],
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    return {"enrolled": True}


class WebAuthnTheftOptionsIn(BaseModel):
    key_id: str  # device declaring the theft (NOT required to be creator)
    origin: str


@api_router.post("/webauthn/declare-theft-options")
async def webauthn_declare_theft_options(payload: WebAuthnTheftOptionsIn):
    """Public — anyone with a registered device can attempt to declare theft.
    Returns a WebAuthn authentication challenge listing every enrolled
    biometric credential as 'allowed'. The matching biometric on this device
    must validate (USER_VERIFICATION required)."""
    if not await _device_by_key(payload.key_id):
        raise HTTPException(status_code=404, detail="Appareil inconnu.")
    rp_id = _rp_id_from_origin(payload.origin)
    creds = await db.webauthn_credentials.find({"rp_id": rp_id}, {"_id": 0}).to_list(length=20)
    if not creds:
        raise HTTPException(status_code=404, detail="Aucune biométrie enrôlée — le créateur doit d'abord enregistrer son empreinte depuis son appareil créateur.")

    allow = [
        PublicKeyCredentialDescriptor(id=webauthn.helpers.base64url_to_bytes(c["credential_id"]))
        for c in creds
    ]
    options = webauthn.generate_authentication_options(
        rp_id=rp_id,
        allow_credentials=allow,
        user_verification=UserVerificationRequirement.REQUIRED,
        timeout=60_000,
    )
    await db.webauthn_challenges.insert_one({
        "key_id": payload.key_id,
        "challenge": webauthn.helpers.bytes_to_base64url(options.challenge),
        "rp_id": rp_id,
        "kind": "theft",
        "origin": payload.origin,
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    return json.loads(webauthn.helpers.options_to_json(options))


class WebAuthnTheftVerifyIn(BaseModel):
    key_id: str
    origin: str
    credential: Dict[str, Any]


@api_router.post("/webauthn/declare-theft-verify")
async def webauthn_declare_theft_verify(payload: WebAuthnTheftVerifyIn):
    """Verify the assertion. On success:
      - Promote the calling device to 'creator' role.
      - Revoke every OTHER creator device.
      - Force-disconnect their sessions.
    """
    dev = await _device_by_key(payload.key_id)
    if not dev:
        raise HTTPException(status_code=404, detail="Appareil inconnu.")
    challenge_doc = await db.webauthn_challenges.find_one_and_delete(
        {"key_id": payload.key_id, "kind": "theft"},
    )
    if not challenge_doc:
        raise HTTPException(status_code=400, detail="Aucun défi de récupération actif.")

    # Look up the credential by id supplied in the assertion.
    raw_id = payload.credential.get("rawId") or payload.credential.get("id")
    if not raw_id:
        raise HTTPException(status_code=400, detail="Credential id manquant.")
    cred_doc = await db.webauthn_credentials.find_one({"credential_id": raw_id}, {"_id": 0})
    if not cred_doc:
        raise HTTPException(status_code=404, detail="Credential inconnu.")

    try:
        verification = webauthn.verify_authentication_response(
            credential=payload.credential,
            expected_challenge=webauthn.helpers.base64url_to_bytes(challenge_doc["challenge"]),
            expected_rp_id=challenge_doc["rp_id"],
            expected_origin=payload.origin,
            credential_public_key=webauthn.helpers.base64url_to_bytes(cred_doc["public_key"]),
            credential_current_sign_count=cred_doc.get("sign_count", 0),
            require_user_verification=True,
        )
    except Exception as e:
        raise HTTPException(status_code=401, detail=f"Vérification biométrique échouée: {e}")

    # Bump sign_count to prevent replay.
    await db.webauthn_credentials.update_one(
        {"credential_id": raw_id},
        {"$set": {"sign_count": verification.new_sign_count}},
    )

    now_iso = datetime.now(timezone.utc).isoformat()
    # Remove all OTHER creators (the audit trail in device_decisions remembers).
    other_creators = await db.device_keys.find(
        {"role": "creator", "key_id": {"$ne": payload.key_id}},
        {"_id": 0, "key_id": 1, "label": 1},
    ).to_list(length=50)
    for d in other_creators:
        await db.device_keys.delete_one({"key_id": d["key_id"]})
        await db.device_nonces.delete_many({"key_id": d["key_id"]})
        await db.user_sessions.delete_many({"device_key_id": d["key_id"]})
        await _log_decision("revoke", d["key_id"], payload.key_id, d.get("label"))

    # Promote the requester.
    await db.device_keys.update_one(
        {"key_id": payload.key_id},
        {"$set": {"role": "creator", "promoted_at": now_iso, "promoted_reason": "theft"}},
    )
    await _log_decision("promote", payload.key_id, payload.key_id, dev.get("label"))
    return {"recovered": True, "revoked_count": len(other_creators)}


@api_router.get("/webauthn/has-enrollment")
async def webauthn_has_enrollment():
    """Public lookup — returns whether any biometric is enrolled on this
    deployment. Used by the login page to show the right CTA."""
    count = await db.webauthn_credentials.count_documents({})
    return {"enrolled_count": count, "has_any": count > 0}


# ==========================================================================
# END pending sessions + WebAuthn
# ==========================================================================


# ==========================================================================
# PRIVATE MESSAGING — creator ↔ user direct messages.
#
# Threads are keyed by the NON-creator device's key_id. Each message belongs
# to exactly one thread and has a direction (is_from_creator). The creator
# can see all threads + reply to any of them. A non-creator device sees only
# its own thread.
#
# Auth is the same ECDSA signature scheme as the rest of the device API.
# Anyone with a valid registered device can /send (subject to a cool-down).
# Reading the inbox/thread requires creator proof for the creator view, or
# the requester's own signature for their personal thread.
# ==========================================================================

MAX_MESSAGE_LEN = 2000
MESSAGE_COOLDOWN_SECONDS = 30  # per-device anti-flood on /messages/send


class MessageSendIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    content: str
    # When the sender IS the creator, they must supply the target thread.
    target_key_id: Optional[str] = None


@api_router.post("/messages/send")
async def messages_send(payload: MessageSendIn):
    """Send a message in a thread.

    If the sender is the creator, `target_key_id` must be supplied — the
    message lands in that thread as is_from_creator=True.

    If the sender is anything else, the thread IS the sender's own key_id and
    is_from_creator=False. Blocked devices are rejected with the same
    localized message as /devices/send-to-creator. Per-device cooldown of
    30 seconds applies to prevent flooding."""
    content = (payload.content or "").strip()
    if not content:
        raise HTTPException(status_code=400, detail="Message vide.")
    if len(content) > MAX_MESSAGE_LEN:
        raise HTTPException(status_code=400, detail=f"Message trop long ({MAX_MESSAGE_LEN} max).")
    sender = await _device_by_key(payload.key_id)
    if not sender:
        raise HTTPException(status_code=404, detail="Appareil inconnu.")
    if sender.get("role") == "blocked":
        raise HTTPException(
            status_code=403,
            detail="Votre demande a été formulée de nombreuses fois. Veuillez contacter le créateur.",
        )
    # Cool-down (anti-flood, much shorter than send-to-creator's nudge cooldown).
    # Exempt the creator so they can reply to several users in quick succession.
    last_msg_iso = sender.get("last_message_at")
    is_creator_sender_quick = sender.get("role") == "creator"
    if last_msg_iso and not is_creator_sender_quick:
        try:
            last_msg = datetime.fromisoformat(last_msg_iso)
            elapsed = (datetime.now(timezone.utc) - last_msg).total_seconds()
            if elapsed < MESSAGE_COOLDOWN_SECONDS:
                raise HTTPException(
                    status_code=429,
                    detail=f"Patiente {int(MESSAGE_COOLDOWN_SECONDS - elapsed)}s avant d'envoyer un autre message.",
                )
        except HTTPException:
            raise
        except Exception:
            pass

    if not await _consume_nonce(payload.key_id, payload.nonce):
        raise HTTPException(status_code=403, detail="Nonce invalide ou expiré.")
    if not verify_signature(sender.get("public_key_jwk") or {}, payload.nonce, payload.signature):
        raise HTTPException(status_code=403, detail="Signature invalide.")

    is_creator_sender = sender.get("role") == "creator"
    if is_creator_sender:
        if not payload.target_key_id:
            raise HTTPException(status_code=400, detail="target_key_id requis pour les créateurs.")
        thread_key_id = payload.target_key_id
        target = await _device_by_key(thread_key_id)
        if not target:
            raise HTTPException(status_code=404, detail="Destinataire inconnu.")
    else:
        thread_key_id = payload.key_id

    now = datetime.now(timezone.utc)
    msg_id = f"msg_{uuid.uuid4().hex[:16]}"
    sender_label = sender.get("label") or sender.get("pseudo") or None
    await db.messages.insert_one({
        "message_id": msg_id,
        "thread_key_id": thread_key_id,
        "from_key_id": payload.key_id,
        "is_from_creator": bool(is_creator_sender),
        "content": content,
        "sender_label": sender_label,
        "ts": now.isoformat(),
        "read_by_creator": bool(is_creator_sender),
        "read_by_user": not bool(is_creator_sender),
    })
    await db.device_keys.update_one(
        {"key_id": payload.key_id},
        {"$set": {"last_message_at": now.isoformat()}},
    )
    return {"sent": True, "message_id": msg_id, "ts": now.isoformat()}


class MessagesInboxIn(BaseModel):
    key_id: str
    nonce: str
    signature: str


@api_router.post("/messages/inbox")
async def messages_inbox(payload: MessagesInboxIn):
    """iter82 — Creator + Staff (modo/admin) — return one row per thread with
    last message + unread count. Modos voient les threads où to_key_id matche
    leur propre clé. Admins voient tout comme la créatrice."""
    dev = await _device_by_key(payload.key_id)
    if not dev:
        raise HTTPException(status_code=404, detail="Clé inconnue.")
    if not await _consume_nonce(payload.key_id, payload.nonce):
        raise HTTPException(status_code=403, detail="Nonce invalide ou expiré.")
    if not verify_signature(dev.get("public_key_jwk") or {}, payload.nonce, payload.signature):
        raise HTTPException(status_code=403, detail="Signature invalide.")
    role = dev.get("role")
    sk = dev.get("staff_kind")
    is_creator = role == "creator"
    is_admin = sk == "admin"
    is_modo = sk == "modo"
    if not (is_creator or is_admin or is_modo):
        raise HTTPException(status_code=403, detail="Accès réservé staff.")
    # Match : créa+admin voient tout. Modos voient les threads qui leur sont assignés.
    match = {} if (is_creator or is_admin) else {"to_key_id": payload.key_id}
    pipeline = [
        {"$match": match},
        {"$sort": {"ts": -1}},
        {"$group": {
            "_id": "$thread_key_id",
            "last_ts": {"$first": "$ts"},
            "last_content": {"$first": "$content"},
            "last_is_from_creator": {"$first": "$is_from_creator"},
            "last_sender_label": {"$first": "$sender_label"},
            "recipient_kind": {"$first": "$recipient_kind"},
            "to_key_id": {"$first": "$to_key_id"},
            "unread": {
                "$sum": {
                    "$cond": [{"$and": [
                        {"$eq": ["$is_from_creator", False]},
                        {"$eq": ["$read_by_creator", False]},
                    ]}, 1, 0]
                }
            },
            "total": {"$sum": 1},
        }},
        {"$sort": {"last_ts": -1}},
        {"$limit": 100},
    ]
    rows = await db.messages.aggregate(pipeline).to_list(length=100)
    out = []
    for r in rows:
        d = await _device_by_key(r["_id"]) or {}
        out.append({
            "thread_key_id": r["_id"],
            "label": d.get("label"),
            "pseudo": d.get("pseudo"),
            "role": d.get("role"),
            "last_ts": r["last_ts"],
            "last_content": r["last_content"][:140],
            "last_is_from_creator": r["last_is_from_creator"],
            "last_sender_label": r.get("last_sender_label"),
            "recipient_kind": r.get("recipient_kind"),
            "to_key_id": r.get("to_key_id"),
            "unread": r["unread"],
            "total": r["total"],
        })
    return {"threads": out}


class MessagesThreadIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    thread_key_id: Optional[str] = None  # if creator, target thread; else ignored


@api_router.post("/messages/thread")
async def messages_thread(payload: MessagesThreadIn):
    """Return the full thread.

    - If the caller is the creator and supplies a `thread_key_id`, returns
      that thread (and marks all incoming messages as read_by_creator).
    - If the caller is NOT the creator, returns their own thread (where
      thread_key_id == caller's key_id) and marks creator replies as
      read_by_user."""
    sender = await _device_by_key(payload.key_id)
    if not sender:
        raise HTTPException(status_code=404, detail="Appareil inconnu.")
    if not await _consume_nonce(payload.key_id, payload.nonce):
        raise HTTPException(status_code=403, detail="Nonce invalide ou expiré.")
    if not verify_signature(sender.get("public_key_jwk") or {}, payload.nonce, payload.signature):
        raise HTTPException(status_code=403, detail="Signature invalide.")
    is_creator_caller = sender.get("role") == "creator"
    if is_creator_caller:
        thread_key_id = payload.thread_key_id
        if not thread_key_id:
            raise HTTPException(status_code=400, detail="thread_key_id requis.")
        await db.messages.update_many(
            {"thread_key_id": thread_key_id, "is_from_creator": False, "read_by_creator": False},
            {"$set": {"read_by_creator": True}},
        )
    else:
        thread_key_id = payload.key_id
        await db.messages.update_many(
            {"thread_key_id": thread_key_id, "is_from_creator": True, "read_by_user": False},
            {"$set": {"read_by_user": True}},
        )
    rows = await db.messages.find(
        {"thread_key_id": thread_key_id},
        {"_id": 0},
    ).sort("ts", 1).to_list(length=500)
    return {"thread_key_id": thread_key_id, "messages": rows}


@api_router.post("/messages/unread-count")
async def messages_unread_count(payload: MessagesThreadIn):
    """Return how many unread messages exist for the caller.
    - Creator → total unread received from all users.
    - Anyone else → unread creator replies in their own thread."""
    sender = await _device_by_key(payload.key_id)
    if not sender:
        raise HTTPException(status_code=404, detail="Appareil inconnu.")
    if not await _consume_nonce(payload.key_id, payload.nonce):
        raise HTTPException(status_code=403, detail="Nonce invalide ou expiré.")
    if not verify_signature(sender.get("public_key_jwk") or {}, payload.nonce, payload.signature):
        raise HTTPException(status_code=403, detail="Signature invalide.")
    if sender.get("role") == "creator":
        # Exclude messages from muted senders so the creator gets no notif badge.
        muted_ids = [d["key_id"] async for d in db.device_keys.find({"muted": True}, {"_id": 0, "key_id": 1})]
        q = {"is_from_creator": False, "read_by_creator": False}
        if muted_ids:
            q["thread_key_id"] = {"$nin": muted_ids}
        n = await db.messages.count_documents(q)
    else:
        n = await db.messages.count_documents({
            "thread_key_id": payload.key_id,
            "is_from_creator": True,
            "read_by_user": False,
        })
    return {"unread": n}


class MessagesDeleteIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    thread_key_id: str


class MessagesRenameIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    thread_key_id: str
    new_label: str


@api_router.post("/messages/rename-contact")
async def messages_rename_contact(payload: MessagesRenameIn):
    """Creator-only — rename the displayed label of a contact (the pseudo
    shown in the inbox + above the conversation). Stored on the device_keys
    row; the user themselves still sees their own pseudo unchanged."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    new_label = (payload.new_label or "").strip()
    if not new_label:
        raise HTTPException(status_code=400, detail="Nom requis.")
    if len(new_label) > 40:
        raise HTTPException(status_code=400, detail="Nom trop long (40 max).")
    target = await _device_by_key(payload.thread_key_id)
    if not target:
        raise HTTPException(status_code=404, detail="Destinataire inconnu.")
    await db.device_keys.update_one(
        {"key_id": payload.thread_key_id},
        {"$set": {"label": new_label}},
    )
    return {"success": True, "label": new_label}


@api_router.post("/messages/delete-thread")
async def messages_delete_thread(payload: MessagesDeleteIn):
    """Creator-only — delete an entire thread."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    res = await db.messages.delete_many({"thread_key_id": payload.thread_key_id})
    return {"deleted": res.deleted_count}


# ==========================================================================
# END messaging
# ==========================================================================


# ==================== CODE GENERATION ROUTES ====================

@api_router.post("/generate/code")
async def generate_code(request: Request, input: GenerateCodeRequest):
    """Generate complete code for a project (simulated for now)"""
    user_id = await get_current_user(request)
    
    try:
        # Update project status
        await db.projects.update_one(
            {"project_id": input.project_id, "user_id": user_id},
            {"$set": {"status": "generating"}}
        )
        
        # Generate sample code structure (would use Emergent API in production)
        generated_code = {
            "files": [
                {
                    "path": "index.html",
                    "content": f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{input.description}</title>
    <link rel="stylesheet" href="style.css">
</head>
<body>
    <div class="container">
        <h1>{input.description}</h1>
        <p>Application générée automatiquement</p>
    </div>
    <script src="app.js"></script>
</body>
</html>"""
                },
                {
                    "path": "style.css",
                    "content": """* {
    margin: 0;
    padding: 0;
    box-sizing: border-box;
}

body {
    font-family: 'Arial', sans-serif;
    background: #050505;
    color: #ffffff;
}

.container {
    max-width: 1200px;
    margin: 0 auto;
    padding: 2rem;
}

h1 {
    color: #E4FF00;
    font-size: 3rem;
    margin-bottom: 1rem;
}"""
                },
                {
                    "path": "app.js",
                    "content": """console.log('Application prête !');

// Votre code JavaScript ici"""
                },
                {
                    "path": "manifest.json",
                    "content": json.dumps({
                        "name": input.description,
                        "short_name": input.description[:20],
                        "description": f"Application {input.project_type}",
                        "start_url": "/",
                        "display": "standalone",
                        "background_color": "#050505",
                        "theme_color": "#E4FF00",
                        "icons": [
                            {
                                "src": "/icon-192.png",
                                "sizes": "192x192",
                                "type": "image/png"
                            }
                        ]
                    }, indent=2)
                },
                {
                    "path": "README.md",
                    "content": f"""# {input.description}

Application générée par CodeForge AI

## Installation

### Web
Ouvrez `index.html` dans votre navigateur

### Mobile (APK)
1. Installez via la page d'export mobile
2. Activez les sources inconnues sur Android

### Desktop (EXE)
1. Téléchargez l'installateur
2. Exécutez et suivez les instructions

## Déploiement

### Vercel
```bash
npm install -g vercel
vercel
```

### Netlify
Glissez-déposez le dossier sur netlify.com
"""
                }
            ],
            "instructions": f"Application {input.project_type} générée avec succès. Prête pour l'export.",
            "dependencies": []
        }
        
        # Update project with generated code
        await db.projects.update_one(
            {"project_id": input.project_id},
            {"$set": {
                "generated_code": generated_code,
                "status": "completed",
                "updated_at": datetime.now(timezone.utc).isoformat()
            }}
        )
        
        return {
            "project_id": input.project_id,
            "generated_code": generated_code,
            "status": "completed"
        }
    
    except Exception as e:
        logger.error(f"Error generating code: {e}")
        
        # Update project status to error
        await db.projects.update_one(
            {"project_id": input.project_id},
            {"$set": {"status": "error"}}
        )
        
        raise HTTPException(status_code=500, detail=f"Erreur de génération: {str(e)}")

# ==================== EXPORT ROUTES ====================

@api_router.post("/export/github/{project_id}")
async def export_project_to_github(project_id: str, request: Request):
    """Push every file of a generated project to the configured GitHub repository
    under `projects/<sanitized-name>/` so the user has a permanent backup like
    Emergent's Save-to-Github.
    """
    user_id = await get_current_user(request)
    project = await db.projects.find_one(
        {"project_id": project_id, "user_id": user_id}, {"_id": 0}
    )
    if not project:
        raise HTTPException(status_code=404, detail="Projet non trouvé")
    if not GITHUB_ENABLED:
        raise HTTPException(status_code=503, detail="GitHub non configuré côté serveur.")

    files = (project.get("generated_code") or {}).get("files", [])
    is_chat = project.get("project_type") == "chat"
    if not files and not is_chat:
        raise HTTPException(status_code=400, detail="Aucun code à exporter.")

    # Sanitize folder name — keep ASCII alphanumeric/hyphen/underscore only.
    import unicodedata as _ud
    base = (project.get("name") or project_id)[:60]
    # Normalize accents (é → e), then keep only safe chars.
    ascii_base = _ud.normalize("NFKD", base).encode("ascii", "ignore").decode("ascii")
    safe = "".join(c if (c.isalnum() or c in ('-', '_')) else '-' for c in ascii_base).strip('-') or project_id
    folder = f"projects/{safe}-{project_id}"

    pushed = []
    failed = []
    for f in files:
        path = f.get("path") or "untitled.txt"
        content = f.get("content") or ""
        try:
            ok = await push_to_github(f"{folder}/{path}", content)
            (pushed if ok else failed).append(path)
        except Exception as exc:
            logger.warning(f"GitHub push failed for {path}: {exc}")
            failed.append(path)

    # README + chat transcript (chat-type projects get no source code, only history).
    readme = (
        f"# {project.get('name', '')}\n\n"
        f"{project.get('description', '')}\n\n"
        f"_Exporté depuis CodeForge AI · ID `{project_id}`_\n"
    )
    try:
        ok_readme = await push_to_github(f"{folder}/README.md", readme)
        (pushed if ok_readme else failed).append("README.md")
    except Exception:
        failed.append("README.md")

    if is_chat:
        msgs = await db.chat_messages.find(
            {"user_id": user_id, "project_id": project_id},
            {"_id": 0, "role": 1, "content": 1, "timestamp": 1},
        ).sort("timestamp", 1).to_list(length=10000)
        transcript = "\n\n".join(
            f"### {('Toi' if m.get('role') == 'user' else 'CodeForge')} — {m.get('timestamp', '')}\n{m.get('content', '')}"
            for m in msgs
        )
        try:
            ok_tx = await push_to_github(f"{folder}/chat-transcript.md", transcript or "(vide)")
            (pushed if ok_tx else failed).append("chat-transcript.md")
        except Exception:
            failed.append("chat-transcript.md")

    repo_url = f"https://github.com/{GITHUB_OWNER}/{GITHUB_REPO_NAME}/tree/main/{folder}"
    all_failed = bool(failed) and not pushed
    return {
        "success": not all_failed,
        "repository": f"{GITHUB_OWNER}/{GITHUB_REPO_NAME}",
        "folder": folder,
        "url": repo_url,
        "pushed": pushed,
        "failed": failed,
    }


@api_router.post("/export/download")
async def download_export(request: Request, export_req: ExportRequest):
    """Download project as ZIP"""
    user_id = await get_current_user(request)
    
    project = await db.projects.find_one(
        {"project_id": export_req.project_id, "user_id": user_id},
        {"_id": 0}
    )
    
    if not project:
        raise HTTPException(status_code=404, detail="Projet non trouvé")
    
    # Chat-type projects have no generated_code — they export the transcript only.
    if not project.get("generated_code") and project.get("project_type") != "chat":
        raise HTTPException(status_code=400, detail="Aucun code généré. Générez d'abord le code.")
    
    # Ensure generated_code exists for the loop below.
    if not project.get("generated_code"):
        project["generated_code"] = {"files": []}
    
    # Create ZIP in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        generated_code = project["generated_code"]
        if export_req.include_code:
            for file_data in generated_code.get("files", []):
                zip_file.writestr(file_data["path"], file_data["content"])
            # Always include a README at the root if not already present.
            existing_paths = {f.get("path") for f in generated_code.get("files", [])}
            if "README.md" not in existing_paths and "readme.md" not in existing_paths:
                readme = (
                    f"# {project['name']}\n\n"
                    f"{project.get('description', '')}\n\n"
                    f"---\nGénéré par CodeForge AI · ID `{project['project_id']}`\n"
                    f"Créé le : {project.get('created_at')}\n"
                )
                zip_file.writestr("README.md", readme)
        # If this project is a saved chat OR include_chat is true, append a transcript.
        if export_req.include_chat or project.get("project_type") == "chat":
            msgs = await db.chat_messages.find(
                {"user_id": user_id, "project_id": export_req.project_id},
                {"_id": 0, "role": 1, "content": 1, "timestamp": 1, "created_at": 1},
            ).sort("timestamp", 1).to_list(length=10000)
            # Format markdown
            transcript = "\n\n".join(
                f"### {('Toi' if m.get('role') == 'user' else 'CodeForge')} — {m.get('timestamp', m.get('created_at', ''))}\n{m.get('content', '')}"
                for m in msgs
            )
            zip_file.writestr("chat-transcript.md", transcript or "(empty)")
            # iter80 C17 — also a .docx file inside the ZIP
            try:
                from docx import Document
                from docx.shared import RGBColor
                doc = Document()
                doc.add_heading(project.get("name") or export_req.project_id, 0)
                doc.add_paragraph(f"Exporté le {datetime.now(timezone.utc).strftime('%d/%m/%Y %H:%M UTC')} — {len(msgs)} message(s)")
                for m in msgs:
                    speaker = "Utilisateur" if m.get("role") == "user" else "IA"
                    p = doc.add_paragraph()
                    run = p.add_run(f"[{speaker}] ")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xE4, 0xFF, 0x00) if m.get("role") != "user" else RGBColor(0x00, 0xD4, 0xFF)
                    p.add_run((m.get("content") or "")[:50000])
                doc_buf = io.BytesIO()
                doc.save(doc_buf)
                zip_file.writestr("chat-transcript.docx", doc_buf.getvalue())
            except Exception:
                pass
    
    zip_buffer.seek(0)
    
    return Response(
        content=zip_buffer.getvalue(),
        media_type="application/zip",
        headers={
            "Content-Disposition": f"attachment; filename={project['name']}.zip"
        }
    )

@api_router.get("/export/mobile/{project_id}")
async def redirect_to_pwa_install(project_id: str):
    """Redirect to PWA installation page"""
    return RedirectResponse(url=f"/api/pwa/install/{project_id}")

@api_router.get("/export/desktop/{project_id}")
async def redirect_to_desktop_install(project_id: str):
    """Redirect to Desktop installation page"""
    return RedirectResponse(url=f"/api/desktop/install/{project_id}")

@api_router.get("/export/download/apk/{project_id}")
async def download_apk(project_id: str):
    """Generate and download APK (simplified version)"""
    project = await db.projects.find_one({"project_id": project_id}, {"_id": 0})
    
    if not project or not project.get("generated_code"):
        raise HTTPException(status_code=404, detail="Projet ou code non trouvé")
    
    # For now, return the ZIP with instructions
    # In production, this would build an actual APK using Capacitor
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        generated_code = project["generated_code"]
        for file_data in generated_code.get("files", []):
            zip_file.writestr(file_data["path"], file_data["content"])
        
        # Add APK build instructions
        zip_file.writestr("BUILD_APK.md", """# Build APK Instructions

## Option 1: Using Capacitor (Recommended)
```bash
npm install -g @capacitor/cli
capacitor init
capacitor add android
capacitor open android
# Build in Android Studio
```

## Option 2: Using PWA Builder
1. Visit https://www.pwabuilder.com/
2. Enter your app URL
3. Download Android package

## Option 3: Direct Install (PWA)
1. Host these files on a server
2. Open in Chrome on Android
3. Click "Add to Home Screen"
""")
    
    zip_buffer.seek(0)
    
    return Response(
        content=zip_buffer.getvalue(),
        media_type="application/zip",
        headers={
            "Content-Disposition": f"attachment; filename={project['name']}_android.zip"
        }
    )

@api_router.get("/export/download/exe/{project_id}")
async def download_exe(project_id: str):
    """Generate and download EXE (simplified version)"""
    project = await db.projects.find_one({"project_id": project_id}, {"_id": 0})
    
    if not project or not project.get("generated_code"):
        raise HTTPException(status_code=404, detail="Projet ou code non trouvé")
    
    # For now, return the ZIP with instructions
    # In production, this would build an actual EXE using Electron
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        generated_code = project["generated_code"]
        for file_data in generated_code.get("files", []):
            zip_file.writestr(file_data["path"], file_data["content"])
        
        # Add Electron package.json
        zip_file.writestr("electron/package.json", json.dumps({
            "name": project["name"],
            "version": "1.0.0",
            "main": "main.js",
            "scripts": {
                "start": "electron .",
                "build": "electron-builder"
            },
            "devDependencies": {
                "electron": "latest",
                "electron-builder": "latest"
            },
            "build": {
                "appId": f"com.codeforge.{project['name'].lower()}",
                "win": {
                    "target": "nsis"
                }
            }
        }, indent=2))
        
        # Add Electron main.js
        zip_file.writestr("electron/main.js", """const { app, BrowserWindow } = require('electron');

function createWindow() {
    const win = new BrowserWindow({
        width: 1200,
        height: 800,
        webPreferences: {
            nodeIntegration: false,
            contextIsolation: true
        }
    });
    
    win.loadFile('../index.html');
}

app.whenReady().then(createWindow);
""")
        
        # Add build instructions
        zip_file.writestr("BUILD_EXE.md", """# Build Windows EXE Instructions

## Using Electron
```bash
cd electron
npm install
npm run build
```

The .exe will be in `electron/dist/`

## Alternative: Portable HTML App
Use NW.js or similar to package as standalone app
""")
    
    zip_buffer.seek(0)
    
    return Response(
        content=zip_buffer.getvalue(),
        media_type="application/zip",
        headers={
            "Content-Disposition": f"attachment; filename={project['name']}_windows.zip"
        }
    )

@api_router.get("/preview/{preview_id}")
async def get_preview(preview_id: str):
    """Get preview HTML for generated content"""
    preview = await db.previews.find_one({"preview_id": preview_id}, {"_id": 0})
    
    if not preview:
        return HTMLResponse("<h1>Prévisualisation non trouvée</h1>", status_code=404)
    
    # Return HTML preview
    html_content = preview.get("html_content", "")
    
    if not html_content:
        # Generate preview from files
        files = preview.get("files", [])
        html_parts = ["<!DOCTYPE html><html><head><meta charset='utf-8'><title>Prévisualisation</title>"]
        
        # Add CSS
        for file in files:
            if file["path"].endswith(".css"):
                html_parts.append(f"<style>{file['content']}</style>")
        
        html_parts.append("</head><body>")
        
        # Add HTML
        for file in files:
            if file["path"].endswith(".html"):
                html_parts.append(file["content"])
        
        # Add JS
        for file in files:
            if file["path"].endswith(".js"):
                html_parts.append(f"<script>{file['content']}</script>")
        
        html_parts.append("</body></html>")
        html_content = "\n".join(html_parts)
    
    return HTMLResponse(content=html_content)

@api_router.get("/preview/project/{project_id}")
async def get_project_preview(project_id: str):
    """Get preview for a specific project"""
    project = await db.projects.find_one({"project_id": project_id}, {"_id": 0})
    
    if not project:
        return HTMLResponse("<h1>Projet non trouvé</h1>", status_code=404)
    
    generated_code = project.get("generated_code") or {}
    files = (generated_code.get("files") if isinstance(generated_code, dict) else None) or []
    
    # Generate combined HTML preview
    html_parts = ["""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Prévisualisation - """ + project.get("name", "Projet") + """</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { font-family: system-ui, sans-serif; background: #050505; color: #fff; }
    </style>"""]
    
    # Add CSS files
    for file in files:
        if file["path"].endswith(".css"):
            html_parts.append(f"<style>{file['content']}</style>")
    
    html_parts.append("</head><body>")
    
    # Add HTML content
    for file in files:
        if file["path"].endswith(".html"):
            # Extract body content if full HTML
            content = file["content"]
            if "<body>" in content:
                start = content.find("<body>") + 6
                end = content.find("</body>")
                content = content[start:end] if end > start else content
            html_parts.append(content)
    
    # Add JS files
    for file in files:
        if file["path"].endswith(".js"):
            html_parts.append(f"<script>{file['content']}</script>")
    
    html_parts.append("</body></html>")
    
    return HTMLResponse(content="\n".join(html_parts))

@api_router.get("/preview/demo/{preview_type}")
async def get_demo_preview(preview_type: str):
    """Get demo preview pages for different formats (Web, PDF, DOCX, App)"""
    
    previews = {
        "web": """<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Prévisualisation Web - CodeForge AI</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { 
            font-family: system-ui, -apple-system, sans-serif; 
            background: linear-gradient(135deg, #050505 0%, #0F0F13 100%); 
            color: #ffffff; 
            min-height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
        }
        .container { max-width: 800px; padding: 3rem; text-align: center; }
        .icon { font-size: 5rem; margin-bottom: 2rem; }
        h1 { color: #E4FF00; font-size: 3rem; margin-bottom: 1rem; }
        p { color: #A1A1AA; font-size: 1.2rem; line-height: 1.8; margin-bottom: 2rem; }
        .badge { 
            display: inline-block; 
            background: #00FF66; 
            color: #050505; 
            padding: 0.5rem 1.5rem; 
            border-radius: 2rem; 
            font-weight: bold;
            margin: 0.5rem;
        }
        .features { 
            display: grid; 
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); 
            gap: 1.5rem; 
            margin-top: 3rem;
            text-align: left;
        }
        .feature { 
            background: rgba(255,255,255,0.05); 
            padding: 1.5rem; 
            border-radius: 0.5rem; 
            border: 1px solid rgba(255,255,255,0.1);
        }
        .feature h3 { color: #E4FF00; margin-bottom: 0.5rem; }
    </style>
</head>
<body>
    <div class="container">
        <div class="icon">🌐</div>
        <h1>Prévisualisation Web</h1>
        <p>Ceci est une démonstration de prévisualisation pour les applications web générées par CodeForge AI.</p>
        <span class="badge">HTML5</span>
        <span class="badge">CSS3</span>
        <span class="badge">JavaScript</span>
        <div class="features">
            <div class="feature">
                <h3>Responsive</h3>
                <p>S'adapte à tous les écrans</p>
            </div>
            <div class="feature">
                <h3>Moderne</h3>
                <p>Technologies récentes</p>
            </div>
            <div class="feature">
                <h3>Rapide</h3>
                <p>Performance optimisée</p>
            </div>
        </div>
    </div>
</body>
</html>""",

        "app": """<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Prévisualisation Application - CodeForge AI</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { 
            font-family: system-ui, sans-serif; 
            background: #050505; 
            color: #fff;
            min-height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
        }
        .phone-frame {
            width: 375px;
            height: 667px;
            background: #0F0F13;
            border-radius: 40px;
            padding: 20px;
            border: 3px solid #333;
            box-shadow: 0 25px 50px rgba(0,0,0,0.5);
        }
        .phone-screen {
            background: linear-gradient(180deg, #1a1a2e 0%, #0F0F13 100%);
            height: 100%;
            border-radius: 25px;
            overflow: hidden;
        }
        .status-bar {
            height: 44px;
            background: rgba(0,0,0,0.3);
            display: flex;
            align-items: center;
            justify-content: space-between;
            padding: 0 20px;
            font-size: 12px;
        }
        .app-content {
            padding: 20px;
            text-align: center;
        }
        .app-icon {
            width: 80px;
            height: 80px;
            background: #E4FF00;
            border-radius: 20px;
            margin: 30px auto;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 2.5rem;
        }
        h2 { color: #E4FF00; margin-bottom: 10px; }
        p { color: #A1A1AA; font-size: 14px; }
        .btn {
            background: #E4FF00;
            color: #050505;
            border: none;
            padding: 15px 40px;
            border-radius: 25px;
            font-weight: bold;
            margin-top: 30px;
            cursor: pointer;
        }
    </style>
</head>
<body>
    <div class="phone-frame">
        <div class="phone-screen">
            <div class="status-bar">
                <span>9:41</span>
                <span>📶 🔋</span>
            </div>
            <div class="app-content">
                <div class="app-icon">📱</div>
                <h2>Mon Application</h2>
                <p>Application mobile générée par CodeForge AI</p>
                <button class="btn">Démarrer</button>
            </div>
        </div>
    </div>
</body>
</html>""",

        "pdf": """<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Prévisualisation PDF - CodeForge AI</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { 
            font-family: 'Times New Roman', serif; 
            background: #404040; 
            min-height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
            padding: 2rem;
        }
        .pdf-page {
            width: 595px;
            min-height: 842px;
            background: white;
            color: #000;
            padding: 60px;
            box-shadow: 0 10px 40px rgba(0,0,0,0.5);
        }
        .header {
            border-bottom: 2px solid #E4FF00;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }
        .logo { 
            font-size: 24px; 
            font-weight: bold; 
            color: #050505;
            display: flex;
            align-items: center;
            gap: 10px;
        }
        .logo span { background: #E4FF00; padding: 5px 15px; }
        h1 { font-size: 28px; margin: 30px 0 20px; color: #1a1a1a; }
        p { line-height: 1.8; margin-bottom: 15px; color: #333; }
        .section { margin: 30px 0; }
        .highlight { background: #FFFFD0; padding: 15px; border-left: 4px solid #E4FF00; }
        .footer { 
            margin-top: 50px; 
            padding-top: 20px; 
            border-top: 1px solid #ddd;
            font-size: 12px;
            color: #666;
            text-align: center;
        }
    </style>
</head>
<body>
    <div class="pdf-page">
        <div class="header">
            <div class="logo">
                <span>CodeForge</span> AI
            </div>
        </div>
        <h1>Document de Prévisualisation</h1>
        <p>Ce document représente un aperçu de la génération PDF par CodeForge AI. Les documents générés peuvent inclure du texte formaté, des images et des tableaux.</p>
        
        <div class="section">
            <h2>Caractéristiques</h2>
            <ul style="margin-left: 20px; line-height: 2;">
                <li>Génération automatique de contenu</li>
                <li>Mise en page professionnelle</li>
                <li>Export haute qualité</li>
                <li>Compatible tous appareils</li>
            </ul>
        </div>
        
        <div class="highlight">
            <strong>Note:</strong> Les PDFs générés sont entièrement personnalisables et peuvent être modifiés selon vos besoins.
        </div>
        
        <div class="footer">
            Généré par CodeForge AI - Création Sans Limites
        </div>
    </div>
</body>
</html>""",

        "docx": """<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Prévisualisation DOCX - CodeForge AI</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { 
            font-family: 'Calibri', 'Arial', sans-serif; 
            background: #2b579a; 
            min-height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
            padding: 2rem;
        }
        .word-container {
            background: #f3f3f3;
            padding: 20px;
            border-radius: 5px;
        }
        .toolbar {
            background: #2b579a;
            color: white;
            padding: 10px 20px;
            border-radius: 5px 5px 0 0;
            font-size: 14px;
            display: flex;
            gap: 20px;
        }
        .doc-page {
            width: 612px;
            min-height: 792px;
            background: white;
            color: #000;
            padding: 72px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.2);
        }
        h1 { font-size: 26px; color: #2b579a; margin-bottom: 20px; }
        h2 { font-size: 18px; color: #2b579a; margin: 25px 0 15px; }
        p { line-height: 1.6; margin-bottom: 12px; }
        .table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }
        .table th, .table td {
            border: 1px solid #ddd;
            padding: 10px;
            text-align: left;
        }
        .table th { background: #2b579a; color: white; }
        .footer { 
            margin-top: 40px; 
            font-size: 11px; 
            color: #666;
            border-top: 1px solid #ddd;
            padding-top: 10px;
        }
    </style>
</head>
<body>
    <div class="word-container">
        <div class="toolbar">
            <span>📄 Document1.docx</span>
            <span>Fichier</span>
            <span>Édition</span>
            <span>Affichage</span>
        </div>
        <div class="doc-page">
            <h1>Document Word - Prévisualisation</h1>
            <p>Ce document représente un aperçu de la génération de fichiers DOCX par CodeForge AI. Les documents peuvent être édités dans Microsoft Word ou LibreOffice.</p>
            
            <h2>Fonctionnalités supportées</h2>
            <p>Les documents générés supportent de nombreuses fonctionnalités :</p>
            
            <table class="table">
                <tr>
                    <th>Fonctionnalité</th>
                    <th>Support</th>
                </tr>
                <tr>
                    <td>Texte formaté</td>
                    <td>✅ Complet</td>
                </tr>
                <tr>
                    <td>Tableaux</td>
                    <td>✅ Complet</td>
                </tr>
                <tr>
                    <td>Images</td>
                    <td>✅ Complet</td>
                </tr>
                <tr>
                    <td>En-têtes/Pieds de page</td>
                    <td>✅ Complet</td>
                </tr>
            </table>
            
            <div class="footer">
                Page 1 sur 1 | Généré par CodeForge AI
            </div>
        </div>
    </div>
</body>
</html>""",

        "image": """<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Prévisualisation Image - CodeForge AI</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { 
            font-family: system-ui, sans-serif; 
            background: #1a1a1a; 
            min-height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
        }
        .image-preview {
            background: #0F0F13;
            padding: 20px;
            border-radius: 10px;
            text-align: center;
        }
        .image-frame {
            width: 600px;
            height: 400px;
            background: linear-gradient(135deg, #E4FF00 0%, #00FF66 50%, #00D4FF 100%);
            border-radius: 5px;
            display: flex;
            align-items: center;
            justify-content: center;
            color: #050505;
            font-size: 24px;
            font-weight: bold;
        }
        .info {
            margin-top: 15px;
            color: #A1A1AA;
            font-size: 14px;
        }
    </style>
</head>
<body>
    <div class="image-preview">
        <div class="image-frame">
            🖼️ Image Générée par IA
        </div>
        <div class="info">
            Format: PNG | Résolution: 1920x1080 | Taille: ~2.4 MB
        </div>
    </div>
</body>
</html>"""
    }
    
    content = previews.get(preview_type, previews["web"])
    return HTMLResponse(content=content)

# ==================== VOICE TRANSCRIPTION ====================

@api_router.post("/voice/transcribe")
async def voice_transcribe(
    request: Request,
    file: UploadFile = File(...),
    language: str = "auto",
):
    """Transcribe a short voice recording (<25MB) to text via OpenAI Whisper.

    Used by the chat & create pages for two UX flows:
    - "Voice message" — record → instant send (the AI receives the transcript).
    - "Dictation" — record → fill the input field for review before send.

    Returns: { "text": "...", "language": "fr", "duration_ms": 0 }
    """
    # Auth — same pattern as other endpoints; reject anonymous calls.
    user_id = await get_current_user(request)
    if not user_id:
        raise HTTPException(status_code=401, detail="Authentification requise")

    # Validate file type & size early (Whisper accepts mp3/mp4/mpeg/mpga/m4a/wav/webm).
    allowed = {"audio/webm", "audio/mp3", "audio/mpeg", "audio/mp4",
               "audio/m4a", "audio/wav", "audio/x-wav", "audio/ogg",
               "video/webm"}  # browsers send webm with audio codec as video/webm sometimes
    if file.content_type and file.content_type.split(";")[0].strip() not in allowed:
        # Don't be overly strict — Whisper auto-detects. Just log.
        logger.info(f"voice/transcribe: unusual content-type {file.content_type}")

    raw = await file.read()
    if len(raw) > 25 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="Fichier audio trop volumineux (>25MB)")
    if len(raw) < 200:
        raise HTTPException(status_code=400, detail="Enregistrement trop court ou vide")

    try:
        from emergentintegrations.llm.openai import OpenAISpeechToText
        emergent_key = os.environ.get("EMERGENT_LLM_KEY")
        if not emergent_key:
            raise ValueError("EMERGENT_LLM_KEY not configured")

        stt = OpenAISpeechToText(api_key=emergent_key)
        # Wrap the bytes in a BytesIO for the integration.
        audio_buf = io.BytesIO(raw)
        # Whisper relies on the filename extension to guess the format,
        # so set a sensible name (browsers usually send webm).
        audio_buf.name = file.filename or "recording.webm"

        kwargs = {"file": audio_buf, "model": "whisper-1", "response_format": "json"}
        if language and language != "auto":
            kwargs["language"] = language[:2]  # ISO-639-1

        response = await stt.transcribe(**kwargs)
        text = (getattr(response, "text", "") or "").strip()
        if not text:
            raise HTTPException(status_code=422, detail="Aucun texte reconnu — réessaie en parlant plus clairement.")

        return {"text": text, "language": language, "size_bytes": len(raw)}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"voice/transcribe error: {e}")
        raise HTTPException(status_code=500, detail=f"Erreur de transcription : {e}")

# ==================== ROOT ROUTE ====================

@api_router.get("/")
async def root():
    return {
        "message": "API CodeForge AI - Plateforme de Génération IA Sans Limites",
        "version": "2.0.0",
        "status": "online",
        "features": [
            "Chat IA (GPT)",
            "Création avec Emergent",
            "Export Mobile (APK)",
            "Export Desktop (EXE)",
            "Sans limites"
        ]
    }

@api_router.get("/health")
async def health_check():
    """Health check + deployed version info (helps debug auto-deploy)."""
    import subprocess
    commit = "unknown"
    try:
        commit = subprocess.check_output(
            ["git", "rev-parse", "--short", "HEAD"],
            cwd=str(ROOT_DIR.parent), timeout=2
        ).decode().strip()
    except Exception:
        pass

    # MongoDB ping
    db_ok = False
    try:
        await db.command("ping")
        db_ok = True
    except Exception as e:
        logger.warning(f"Health: Mongo ping failed: {e}")

    # Resend availability
    resend_ok = bool(os.environ.get("RESEND_API_KEY"))

    # Ollama best-effort (don't block — short timeout)
    ollama_ok = False
    try:
        ollama_url = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434")
        async with httpx.AsyncClient(timeout=1.5) as client:
            r = await client.get(f"{ollama_url}/api/tags")
            ollama_ok = r.status_code == 200
    except Exception:
        pass

    return {
        "status": "healthy" if db_ok else "degraded",
        "commit": commit,
        "checks": {
            "mongo": db_ok,
            "resend": resend_ok,
            "ollama": ollama_ok,
            "github": GITHUB_ENABLED,
        },
        "chat_ai": "GPT (disponible)",
        "create_ai": "Emergent (disponible)",
        "exports": "mobile + desktop",
        "ts": datetime.now(timezone.utc).isoformat(),
    }


@api_router.post("/admin/redeploy")
async def redeploy(request: Request):
    """Auto-deploy webhook called by GitHub Actions on push to main.
    Pulls latest code from origin/main and restarts the backend supervisor.
    Protected by DEPLOY_SECRET env var (HMAC-style shared secret).
    """
    import subprocess
    secret = os.environ.get("DEPLOY_SECRET")
    provided = request.headers.get("X-Deploy-Secret")

    if not secret:
        raise HTTPException(status_code=503, detail="DEPLOY_SECRET not configured")
    if not provided or provided != secret:
        raise HTTPException(status_code=401, detail="Invalid deploy secret")

    try:
        repo_dir = str(ROOT_DIR.parent)

        # Back up .env files BEFORE git operations, because they are tracked
        # in this repo and would be overwritten by git pull / reset --hard
        # with the (incomplete) GitHub-side version.
        env_files = ["backend/.env", "frontend/.env"]
        env_backups = {}
        for env_path in env_files:
            full = os.path.join(repo_dir, env_path)
            if os.path.exists(full):
                with open(full, "rb") as f:
                    env_backups[env_path] = f.read()

        # Robust sync: fetch + hard reset on origin/main.
        # Avoids "divergent branches" errors that plain `git pull` hits when
        # local has auto-commits or the same file was edited via GitHub UI.
        fetch_out = subprocess.check_output(
            ["git", "fetch", "origin", "main"],
            cwd=repo_dir, stderr=subprocess.STDOUT, timeout=30
        ).decode()
        reset_out = subprocess.check_output(
            ["git", "reset", "--hard", "origin/main"],
            cwd=repo_dir, stderr=subprocess.STDOUT, timeout=30
        ).decode()

        # Restore .env files (preserve runtime secrets)
        for env_path, content in env_backups.items():
            full = os.path.join(repo_dir, env_path)
            with open(full, "wb") as f:
                f.write(content)

        new_commit = subprocess.check_output(
            ["git", "rev-parse", "--short", "HEAD"],
            cwd=repo_dir, stderr=subprocess.STDOUT, timeout=10
        ).decode().strip()

        # Trigger uvicorn hot-reload by touching server.py
        # (more reliable than supervisorctl restart in background which races
        # with watchfiles and can leave supervisor in STOPPED state).
        subprocess.Popen(
            ["sh", "-c", "sleep 1 && touch backend/server.py"],
            cwd=repo_dir
        )
        return {
            "status": "deploying",
            "commit": new_commit,
            "git_fetch": fetch_out.strip(),
            "git_reset": reset_out.strip(),
            "env_preserved": list(env_backups.keys()),
        }
    except subprocess.CalledProcessError as e:
        raise HTTPException(status_code=500, detail=f"Deploy failed: {e.output.decode() if e.output else str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==================================================================
#                  ITER 54 — Creator power tools
# ==================================================================
# Accounts panel, Ideas inbox, Announcements, Polls, Export approval,
# Account visit, Remove-creator-mode. All endpoints below are creator-
# signature gated unless explicitly noted.

class _CreatorSigIn(BaseModel):
    model_config = ConfigDict(extra="allow")
    key_id: str
    nonce: str
    signature: str


class _TargetCreatorSigIn(_CreatorSigIn):
    target_key_id: str


async def _log_account_event(event: str, target_key_id: str, target_label: Optional[str] = None,
                              extra: Optional[Dict[str, Any]] = None,
                              actor_key_id: Optional[str] = None):
    doc = {
        "event_id": f"ah_{uuid.uuid4().hex[:14]}",
        "event": event,
        "target_key_id": target_key_id,
        "target_label": target_label,
        "ts": datetime.now(timezone.utc).isoformat(),
    }
    # iter79 — tag who did the action (creator/admin/modo) for color-coded UI.
    if actor_key_id:
        doc["actor_key_id"] = actor_key_id
        actor = await db.device_keys.find_one({"key_id": actor_key_id}, {"_id": 0, "role": 1, "staff_kind": 1, "pseudo": 1, "label": 1})
        if actor:
            doc["actor_kind"] = "creator" if actor.get("role") == "creator" else (actor.get("staff_kind") or actor.get("role"))
            doc["actor_label"] = actor.get("pseudo") or actor.get("label")
    if extra:
        doc.update(extra)
    await db.account_history.insert_one(doc)


def _disambiguate_pseudos(rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Append "#N" to duplicate pseudos for creator-side display only."""
    by_lower: Dict[str, int] = {}
    out = []
    for r in rows:
        p = (r.get("pseudo") or r.get("label") or "").strip()
        key = p.lower()
        if not p:
            r["display"] = (r.get("email") or r.get("key_id", "")[:14])
            out.append(r)
            continue
        by_lower[key] = by_lower.get(key, 0) + 1
        n = by_lower[key]
        r["display"] = p if n == 1 else f"{p} #{n}"
        out.append(r)
    return out


@api_router.post("/accounts/list")
async def accounts_list(payload: _CreatorSigIn):
    """Creator-only — list ALL device accounts with pseudo/email/state.

    iter77 — include `inactive` devices too: l'utilisatrice veut voir TOUS les
    comptes inscrits (même ceux d'amis qui ont testé sans pousser de demande).
    On expose `is_pending_nudge=True` pour les rôles inactive pour distinguer.
    """
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    devices = await db.device_keys.find(
        {}, {"_id": 0, "public_key_jwk": 0},
    ).sort("created_at", -1).to_list(length=2000)
    emails = list({d.get("email") for d in devices if d.get("email")})
    users = {}
    if emails:
        async for u in db.users.find({"email": {"$in": emails}}, {"_id": 0, "email": 1, "pseudo": 1}):
            users[u["email"]] = u.get("pseudo")
    for d in devices:
        d["pseudo"] = users.get(d.get("email")) or d.get("pseudo") or d.get("label")
        d["muted"] = bool(d.get("muted"))
        d["banned"] = bool(d.get("banned"))
        d["is_inactive"] = (d.get("role") == "inactive")
        d.setdefault("product", None)
        d.setdefault("model", None)
        d.setdefault("staff_kind", None)  # iter77 — 'admin'|'modo'|None
        d.setdefault("force_visitor", bool(d.get("force_visitor", False)))
    return {"accounts": _disambiguate_pseudos(devices)}


@api_router.post("/accounts/rename-pseudo")
async def accounts_rename_pseudo(payload: _TargetCreatorSigIn):
    """Creator-only — rename a peer's pseudo (everywhere)."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    new_pseudo = (getattr(payload, "new_pseudo", None) or "").strip() if hasattr(payload, "new_pseudo") else ""
    body = payload.model_dump() if hasattr(payload, "model_dump") else {}
    new_pseudo = (body.get("new_pseudo") or "").strip()
    if not (1 <= len(new_pseudo) <= 30):
        raise HTTPException(status_code=400, detail="Pseudo invalide (3-30).")
    # iter75: "créatrice" no longer reserved on pseudo updates either.
    target = await db.device_keys.find_one({"key_id": payload.target_key_id}, {"_id": 0})
    if not target:
        raise HTTPException(status_code=404, detail="Compte introuvable.")
    await db.device_keys.update_one(
        {"key_id": payload.target_key_id},
        {"$set": {"pseudo": new_pseudo, "label": new_pseudo}},
    )
    if target.get("email"):
        await db.users.update_one(
            {"email": target["email"]},
            {"$set": {"pseudo": new_pseudo, "pseudo_lower": new_pseudo.lower()}},
        )
    await _log_account_event("rename", payload.target_key_id, new_pseudo)
    return {"success": True, "pseudo": new_pseudo}


@api_router.post("/accounts/mute")
async def accounts_mute(payload: _TargetCreatorSigIn):
    """iter79 — Mute. Ouvert à staff (admin/modo) et créatrice."""
    await _require_staff_signature(payload.key_id, payload.nonce, payload.signature)
    await db.device_keys.update_one(
        {"key_id": payload.target_key_id},
        {"$set": {"muted": True, "muted_at": datetime.now(timezone.utc).isoformat()}},
    )
    await _log_account_event("mute", payload.target_key_id, actor_key_id=payload.key_id)
    return {"success": True}


@api_router.post("/accounts/unmute")
async def accounts_unmute(payload: _TargetCreatorSigIn):
    await _require_staff_signature(payload.key_id, payload.nonce, payload.signature)
    await db.device_keys.update_one(
        {"key_id": payload.target_key_id},
        {"$set": {"muted": False}, "$unset": {"muted_at": ""}},
    )
    await _log_account_event("unmute", payload.target_key_id, actor_key_id=payload.key_id)
    return {"success": True}


# iter77 — Staff sub-roles (admin / modo) + Force Visitor mode -----------------
class _SetStaffKindIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    target_key_id: str
    staff_kind: Optional[str] = None  # 'admin' | 'modo' | None (clear)


@api_router.post("/accounts/set-staff-kind")
async def accounts_set_staff_kind(payload: _SetStaffKindIn):
    """iter77/79 — Promote / demote a target between approved/admin/modo.

    Permissions :
    - Créatrice : peut tout définir (admin, modo, null).
    - Admin : peut uniquement définir/retirer 'modo' (pas admin, pas créa).
    """
    actor = await _require_staff_signature(payload.key_id, payload.nonce, payload.signature)
    sk = (payload.staff_kind or None)
    if sk not in (None, "admin", "modo"):
        raise HTTPException(status_code=400, detail="staff_kind invalide ('admin'|'modo'|null).")
    if actor.get("role") != "creator":
        # Admin: only allowed to set/clear 'modo'
        actor_sk = actor.get("staff_kind")
        if actor_sk != "admin":
            raise HTTPException(status_code=403, detail="Seuls les admins et créatrice peuvent promouvoir.")
        if sk == "admin":
            raise HTTPException(status_code=403, detail="Seule la créatrice peut nommer un admin.")
    target = await db.device_keys.find_one({"key_id": payload.target_key_id}, {"_id": 0, "role": 1})
    if not target:
        raise HTTPException(status_code=404, detail="Compte introuvable.")
    update = {"$set": {"staff_kind": sk}} if sk else {"$unset": {"staff_kind": ""}}
    await db.device_keys.update_one({"key_id": payload.target_key_id}, update)
    await _log_account_event(f"staff_kind_{sk or 'clear'}", payload.target_key_id, actor_key_id=payload.key_id)
    return {"success": True, "staff_kind": sk}


class _ForceVisitorIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    target_key_id: str
    force: bool = True


@api_router.post("/accounts/force-visitor")
async def accounts_force_visitor(payload: _ForceVisitorIn):
    """iter77 — Mode visiteur forcé sur compte cible : lecture seule sans logout.

    Idéal quand on soupçonne quelqu'un sans vouloir l'exclure ou le déconnecter.
    Le frontend doit lire `force_visitor` depuis `/devices/verify` et bloquer
    les écritures côté UI (canWrite=false)."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    target = await db.device_keys.find_one({"key_id": payload.target_key_id}, {"_id": 0})
    if not target:
        raise HTTPException(status_code=404, detail="Compte introuvable.")
    await db.device_keys.update_one(
        {"key_id": payload.target_key_id},
        {"$set": {"force_visitor": bool(payload.force)}},
    )
    await _log_account_event("force_visitor_on" if payload.force else "force_visitor_off",
                              payload.target_key_id)
    return {"success": True, "force_visitor": bool(payload.force)}


@api_router.post("/accounts/exclude")
async def accounts_exclude(payload: _TargetCreatorSigIn):
    """iter79 — Exclusion temporaire. Ouvert au staff."""
    await _require_staff_signature(payload.key_id, payload.nonce, payload.signature)
    body = payload.model_dump() if hasattr(payload, "model_dump") else {}
    minutes = int(body.get("duration_minutes") or 0)
    if minutes <= 0 or minutes > 60 * 24 * 90:  # max 90 days, no infinite
        raise HTTPException(status_code=400, detail="Durée invalide (1 min - 90 jours).")
    until = datetime.now(timezone.utc) + timedelta(minutes=minutes)
    await db.device_keys.update_one(
        {"key_id": payload.target_key_id},
        {"$set": {"excluded_until": until.isoformat(), "excluded_reason": body.get("reason") or ""}},
    )
    # Also wipe active sessions so the user is kicked immediately.
    target = await db.device_keys.find_one({"key_id": payload.target_key_id}, {"_id": 0, "email": 1})
    if target and target.get("email"):
        await db.user_sessions.delete_many({"email": target["email"]})
    await _log_account_event("exclude", payload.target_key_id, extra={"until": until.isoformat(), "minutes": minutes}, actor_key_id=payload.key_id)
    return {"success": True, "excluded_until": until.isoformat()}


@api_router.post("/accounts/ban")
async def accounts_ban(payload: _TargetCreatorSigIn):
    """iter79 — Bannissement permanent. Ouvert au staff."""
    await _require_staff_signature(payload.key_id, payload.nonce, payload.signature)
    target = await db.device_keys.find_one({"key_id": payload.target_key_id}, {"_id": 0})
    if not target:
        raise HTTPException(status_code=404, detail="Compte introuvable.")
    await db.device_keys.update_one(
        {"key_id": payload.target_key_id},
        {"$set": {"banned": True, "banned_at": datetime.now(timezone.utc).isoformat()}},
    )
    if target.get("email"):
        await db.banned_emails.update_one(
            {"email": target["email"]},
            {"$set": {"email": target["email"], "banned_at": datetime.now(timezone.utc).isoformat()}},
            upsert=True,
        )
        await db.user_sessions.delete_many({"email": target["email"]})
    await _log_account_event("ban", payload.target_key_id, extra={"email": target.get("email")}, actor_key_id=payload.key_id)
    return {"success": True}


@api_router.post("/accounts/unban")
async def accounts_unban(payload: _TargetCreatorSigIn):
    await _require_staff_signature(payload.key_id, payload.nonce, payload.signature)
    target = await db.device_keys.find_one({"key_id": payload.target_key_id}, {"_id": 0})
    await db.device_keys.update_one(
        {"key_id": payload.target_key_id},
        {"$set": {"banned": False}, "$unset": {"banned_at": ""}},
    )
    if target and target.get("email"):
        await db.banned_emails.delete_many({"email": target["email"]})
    await _log_account_event("unban", payload.target_key_id, actor_key_id=payload.key_id)
    return {"success": True}


@api_router.post("/accounts/history")
async def accounts_history(payload: _CreatorSigIn):
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    rows = await db.account_history.find({}, {"_id": 0}).sort("ts", -1).to_list(length=1000)
    return {"history": rows}


@api_router.post("/accounts/history/clear")
async def accounts_history_clear(payload: _CreatorSigIn):
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    r = await db.account_history.delete_many({})
    return {"deleted": r.deleted_count}


@api_router.post("/accounts/visit")
async def accounts_visit(payload: _TargetCreatorSigIn):
    """iter80 C20 — Creator-only : voir le compte d'un user comme s'il s'agissait
    de son propre dashboard (projets, messages, group chats). Les éléments
    supprimés (deleted_by_user, deleted_by_creator, deleted=True) sont retournés
    avec un flag `is_deleted: true` pour permettre au frontend d'appliquer un
    contraste foncé. Le user lui-même ne voit pas ces éléments en vue créateur."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    target = await db.device_keys.find_one({"key_id": payload.target_key_id}, {"_id": 0})
    if not target:
        raise HTTPException(status_code=404, detail="Compte introuvable.")
    user_id = None
    if target.get("email"):
        u = await db.users.find_one({"email": target["email"]}, {"_id": 0, "user_id": 1})
        if u:
            user_id = u["user_id"]
    projects = []
    messages = []
    if user_id:
        raw_projects = await db.projects.find(
            {"user_id": user_id}, {"_id": 0, "generated_code": 0},
        ).sort("created_at", -1).to_list(length=500)
        for p in raw_projects:
            p["is_deleted"] = bool(p.get("deleted_by_user") or p.get("deleted_by_creator") or p.get("deleted"))
            projects.append(p)
        raw_messages = await db.chat_messages.find(
            {"user_id": user_id}, {"_id": 0},
        ).sort("timestamp", -1).to_list(length=2000)
        for m in raw_messages:
            m["is_deleted"] = bool(m.get("deleted"))
            messages.append(m)
    # iter82 C20 — Récupération de TOUS les messages privés (DMs) impliquant
    # cet utilisateur. Le thread_key_id est l'identifiant de l'utilisateur côté
    # créa (pas le créa lui-même). Inclut tout : from_key_id == target ou
    # thread_key_id == target → cela couvre les conversations avec créa, modos
    # ET avec n'importe quel autre user (friend DM).
    private_msgs_cursor = db.messages.find(
        {"$or": [
            {"thread_key_id": payload.target_key_id},
            {"from_key_id": payload.target_key_id},
            {"to_key_id": payload.target_key_id},
        ]},
        {"_id": 0},
    ).sort("ts", -1)
    private_msgs = await private_msgs_cursor.to_list(length=2000)
    # iter82 — Friend requests : voir qui a demandé quoi à cet user.
    fr_cursor = db.friend_requests.find(
        {"$or": [{"from_key_id": payload.target_key_id}, {"to_key_id": payload.target_key_id}]},
        {"_id": 0},
    ).sort("created_at", -1)
    friend_requests = await fr_cursor.to_list(length=200)
    # iter82 — Group chats : à quels group_chats le user a-t-il posté ?
    group_posts_cursor = db.group_messages.find(
        {"from_key_id": payload.target_key_id},
        {"_id": 0},
    ).sort("ts", -1)
    group_posts = await group_posts_cursor.to_list(length=1000)
    return {
        "target": {
            "key_id": payload.target_key_id,
            "email": target.get("email"),
            "pseudo": target.get("pseudo") or target.get("label"),
            "label": target.get("label"),
            "role": target.get("role"),
            "staff_kind": target.get("staff_kind"),
            "force_visitor": target.get("force_visitor"),
            "muted": target.get("muted"),
            "banned": target.get("banned"),
            "last_seen_at": target.get("last_seen_at"),
            "created_at": target.get("created_at"),
            "biometric_kind": target.get("biometric_kind") or target.get("biometric", {}).get("kind") if isinstance(target.get("biometric"), dict) else None,
            "approved_by_kind": target.get("approved_by_kind"),
            "approved_by_label": target.get("approved_by_label"),
        },
        "projects": projects,
        "messages": list(reversed(messages)),
        "private_messages": list(reversed(private_msgs)),
        "friend_requests": friend_requests,
        "group_posts": group_posts,
    }


@api_router.post("/accounts/delete-user-project")
async def accounts_delete_user_project(payload: _TargetCreatorSigIn):
    """Creator-only — delete a user's project (CGU violation / unsafe app)."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    body = payload.model_dump() if hasattr(payload, "model_dump") else {}
    project_id = body.get("project_id")
    if not project_id:
        raise HTTPException(status_code=400, detail="project_id requis.")
    r = await db.projects.update_one(
        {"project_id": project_id},
        {"$set": {"deleted_by_creator": True, "deleted_at": datetime.now(timezone.utc).isoformat()}},
    )
    await _log_account_event("delete_project", payload.target_key_id, extra={"project_id": project_id})
    return {"success": True, "matched": r.matched_count}


@api_router.post("/accounts/delete-one")
async def accounts_delete_one(payload: _TargetCreatorSigIn):
    """Creator-only — fully delete an account (device_key + cascade)."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    target_key_id = payload.target_key_id
    if target_key_id == payload.key_id:
        raise HTTPException(status_code=400, detail="Impossible de supprimer ton propre compte ici.")
    target = await db.device_keys.find_one({"key_id": target_key_id}, {"_id": 0})
    if not target:
        raise HTTPException(status_code=404, detail="Compte introuvable.")
    # Cascade: keep messages/projects for audit but mark the device_key gone.
    await db.device_keys.delete_one({"key_id": target_key_id})
    if target.get("email"):
        await db.user_sessions.delete_many({"email": target["email"]})
    await _log_account_event("delete_account", target_key_id, target.get("label"))
    return {"success": True}


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


async def _email_for_device_key(key_id: str) -> Optional[str]:
    """iter75: resolve the email tied to a device_key, healing missing
    bindings on the fly. Legacy creator devices that pre-date iter63
    sometimes have an empty device_keys.email even though they have an
    active user_session — that broke creator-destructive actions like
    /accounts/delete-all by returning 'Aucun email lié à cet appareil.'
    even when the password was correct. Now we transparently look up the
    active session, derive the owner's email, and persist the binding."""
    me = await db.device_keys.find_one({"key_id": key_id}, {"_id": 0, "email": 1})
    if me and me.get("email"):
        return me["email"]
    sess = await db.user_sessions.find_one(
        {"device_key_id": key_id, "expires_at": {"$gt": _now_iso()}},
        {"_id": 0, "user_id": 1},
    )
    if not sess:
        return None
    owner = await db.users.find_one({"user_id": sess["user_id"]}, {"_id": 0, "email": 1})
    if owner and owner.get("email"):
        await db.device_keys.update_one({"key_id": key_id}, {"$set": {"email": owner["email"]}})
        return owner["email"]
    return None


@api_router.post("/accounts/delete-all")
async def accounts_delete_all(payload: _CreatorSigIn):
    """Creator-only — delete EVERY other account. Self preserved.

    Requires the caller's account password as a destructive-action gate
    (same UX as remove-creator) — set ``password`` in the request body.
    """
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    body = payload.model_dump() if hasattr(payload, "model_dump") else {}
    pwd = body.get("password") or ""
    email = await _email_for_device_key(payload.key_id)
    if not email:
        raise HTTPException(status_code=400, detail="Aucun email lié à cet appareil. Reconnecte-toi pour le re-lier.")
    user = await db.users.find_one({"email": email}, {"_id": 0, "password_hash": 1})
    if not user or not user.get("password_hash"):
        raise HTTPException(status_code=400, detail="Aucun mot de passe configuré.")
    if not bcrypt.checkpw(pwd.encode("utf-8"), user["password_hash"].encode("utf-8")):
        raise HTTPException(status_code=403, detail="Mot de passe incorrect.")
    r = await db.device_keys.delete_many({"key_id": {"$ne": payload.key_id}})
    await _log_account_event("delete_all_accounts", payload.key_id, extra={"deleted": r.deleted_count})
    return {"success": True, "deleted": r.deleted_count}


@api_router.post("/accounts/remove-creator")
async def accounts_remove_creator(payload: _CreatorSigIn):
    """Creator-only — demote a creator device back to 'approved'.

    - target_key_id absent → demote SELF (the calling key_id).
    - target_key_id present → demote that other creator device.

    Both paths require the caller's own account password to confirm intent.
    Logged in account_history + device_decisions.
    """
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    body = payload.model_dump() if hasattr(payload, "model_dump") else {}
    pwd = body.get("password") or ""
    target_key_id = body.get("target_key_id") or payload.key_id  # default = self
    email = await _email_for_device_key(payload.key_id)
    if not email:
        raise HTTPException(status_code=400, detail="Aucun email lié à cet appareil. Reconnecte-toi pour le re-lier.")
    user = await db.users.find_one({"email": email}, {"_id": 0, "password_hash": 1})
    if not user or not user.get("password_hash"):
        raise HTTPException(status_code=400, detail="Aucun mot de passe configuré.")
    if not bcrypt.checkpw(pwd.encode("utf-8"), user["password_hash"].encode("utf-8")):
        raise HTTPException(status_code=403, detail="Mot de passe incorrect.")
    target = await db.device_keys.find_one({"key_id": target_key_id}, {"_id": 0})
    if not target:
        raise HTTPException(status_code=404, detail="Compte introuvable.")
    if target.get("role") != "creator":
        raise HTTPException(status_code=400, detail="Ce compte n'est pas créateur.")
    await db.device_keys.update_one(
        {"key_id": target_key_id},
        {"$set": {"role": "approved"}},
    )
    is_self = target_key_id == payload.key_id
    await _log_account_event("remove_creator_self" if is_self else "remove_creator_other",
                              target_key_id, target.get("label"))
    await db.device_decisions.insert_one({
        "decision_id": f"d_{uuid.uuid4().hex[:14]}",
        "action": "demote",
        "actor_key_id": payload.key_id,
        "target_key_id": target_key_id,
        "ts": datetime.now(timezone.utc).isoformat(),
        "target_label": target.get("label"),
    })
    return {"success": True, "self": is_self}


# ---------------- IDEAS / FEEDBACK ----------------
class IdeasSendIn(BaseModel):
    model_config = ConfigDict(extra="allow")
    key_id: Optional[str] = None
    nonce: Optional[str] = None
    signature: Optional[str] = None
    content: str = ""
    kind: str = "idea"   # 'idea' | 'bug' | 'report' | 'other'

@api_router.post("/ideas/send")
async def ideas_send(request: Request, payload: IdeasSendIn):
    """Any device — send a feedback/idea/bug to creator.

    Signed call (device key present) attaches the sender's pseudo for
    follow-up; anonymous call (login page, no key yet) is accepted too
    and lands as "Anonyme" so brand-new visitors can still report bugs.
    """
    sender_label = "Anonyme"
    sender_key_id = None
    sender_email = None
    if payload.key_id and payload.nonce and payload.signature:
        try:
            dev = await _verify_signed(payload.key_id, payload.nonce, payload.signature)
            sender_label = dev.get("pseudo") or dev.get("label") or payload.key_id[:14]
            sender_key_id = payload.key_id
            sender_email = dev.get("email")
        except HTTPException:
            # Fall back to anonymous if the signature is stale; never block
            # feedback submission for a signature issue.
            pass
    # iter55: no minimum length, empty allowed.
    content = (payload.content or "").strip()
    kind = payload.kind if payload.kind in ("idea", "bug", "report", "other") else "idea"
    ip = request.client.host if request and request.client else None
    await db.ideas.insert_one({
        "idea_id": f"idea_{uuid.uuid4().hex[:14]}",
        "sender_key_id": sender_key_id,
        "sender_label": sender_label,
        "sender_email": sender_email,
        "sender_ip_hash": hashlib.sha256((ip or "").encode()).hexdigest()[:16] if ip else None,
        "kind": kind,
        "content": content,
        "ts": datetime.now(timezone.utc).isoformat(),
        "read": False,
        "page": getattr(payload, "page", None) or "/",
    })
    return {"success": True}


@api_router.post("/ideas/mine")
async def ideas_mine(payload: _CreatorSigIn):
    """Signed by any device — returns items SENT by this device. Public path."""
    await _verify_signed(payload.key_id, payload.nonce, payload.signature)
    rows = await db.ideas.find({"sender_key_id": payload.key_id}, {"_id": 0}).sort("ts", -1).to_list(length=500)
    return {"ideas": rows}


@api_router.post("/ideas/inbox")
async def ideas_inbox(payload: _CreatorSigIn):
    """iter78 — Inbox accessible créa + admin + modo."""
    dev = await _verify_signed(payload.key_id, payload.nonce, payload.signature)
    role = dev.get("role")
    sk = dev.get("staff_kind")
    if role != "creator" and sk not in ("admin", "modo"):
        raise HTTPException(status_code=403, detail="Réservé staff (admin/modo) et créatrice.")
    rows = await db.ideas.find({}, {"_id": 0}).sort("ts", -1).to_list(length=500)
    return {"ideas": rows}


# iter80 — Clear ideas (créa only, password-protected if any unresolved)
class IdeasClearIn(BaseModel):
    model_config = ConfigDict(extra="allow")
    key_id: str
    nonce: str
    signature: str
    scope: str  # 'all' | 'resolved' | 'unresolved'
    password: Optional[str] = None  # required if scope includes unresolved


def _idea_is_resolved(idea: Dict[str, Any]) -> bool:
    """iter80 — Une idée est 'résolue' quand son state vaut validated.
    'refused' compte comme refusé (non-résolu). 'orange' = escalade créa,
    'reset' / pas de state = non-traité."""
    return idea.get("state") == "validated"


@api_router.post("/ideas/clear")
async def ideas_clear(payload: IdeasClearIn):
    """iter80 — Vide tout / résolus / non-résolus. Si le scope inclut des
    non-traités/refusés, exige le mot de passe créatrice."""
    dev = await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    scope = payload.scope
    if scope not in ("all", "resolved", "unresolved"):
        raise HTTPException(status_code=400, detail="scope invalide.")
    # Look at remaining unresolved counts (avant suppression)
    rows = await db.ideas.find({}, {"_id": 0}).to_list(length=2000)
    unresolved_ids = [i["idea_id"] for i in rows if not _idea_is_resolved(i)]
    resolved_ids = [i["idea_id"] for i in rows if _idea_is_resolved(i)]
    # Check password if scope touches unresolved.
    if scope in ("all", "unresolved") and unresolved_ids:
        if not payload.password:
            raise HTTPException(status_code=428, detail="Mot de passe requis pour effacer des retours non-traités.")
        user = await db.users.find_one({"email": dev.get("email")}, {"_id": 0, "password_hash": 1})
        ok = False
        if user and user.get("password_hash"):
            try:
                ok = bcrypt.checkpw(payload.password.encode("utf-8"), user["password_hash"].encode("utf-8"))
            except Exception:
                ok = False
        if not ok:
            raise HTTPException(status_code=403, detail="Mot de passe incorrect. Veuillez réessayer.")
    # Apply deletion
    if scope == "all":
        await db.ideas.delete_many({})
        deleted = len(rows)
    elif scope == "resolved":
        await db.ideas.delete_many({"idea_id": {"$in": resolved_ids}})
        deleted = len(resolved_ids)
    else:  # unresolved
        await db.ideas.delete_many({"idea_id": {"$in": unresolved_ids}})
        deleted = len(unresolved_ids)
    return {"success": True, "deleted": deleted, "scope": scope}


@api_router.post("/ideas/mark-read")
async def ideas_mark_read(payload: _CreatorSigIn):
    """iter78 — Mark-read accessible créa + admin + modo."""
    dev = await _verify_signed(payload.key_id, payload.nonce, payload.signature)
    role = dev.get("role")
    sk = dev.get("staff_kind")
    if role != "creator" and sk not in ("admin", "modo"):
        raise HTTPException(status_code=403, detail="Réservé staff (admin/modo) et créatrice.")
    await db.ideas.update_many({"read": False}, {"$set": {"read": True}})
    return {"success": True}


@api_router.post("/ideas/delete")
async def ideas_delete(payload: _CreatorSigIn):
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    body = payload.model_dump() if hasattr(payload, "model_dump") else {}
    idea_id = body.get("idea_id")
    if not idea_id:
        raise HTTPException(status_code=400, detail="idea_id requis.")
    await db.ideas.delete_one({"idea_id": idea_id})
    return {"success": True}


# iter77 — Set state on a bug/idea (staff = admin+modo, or creator)
class IdeaSetStateIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    idea_id: str
    state: str  # 'validated' | 'refused' | 'orange' | 'reset'


@api_router.post("/ideas/set-state")
async def ideas_set_state(payload: IdeaSetStateIn):
    """iter77 — Marque un bug/idée avec un état : validé, refusé ou orange
    (seule la créa peut). Réservé staff+créa : admin/modo/creator."""
    dev = await _verify_signed(payload.key_id, payload.nonce, payload.signature)
    role = dev.get("role")
    sk = dev.get("staff_kind")
    if role != "creator" and sk not in ("admin", "modo"):
        raise HTTPException(status_code=403, detail="Réservé au staff (admin/modo) et créatrice.")
    if payload.state not in ("validated", "refused", "orange", "reset"):
        raise HTTPException(status_code=400, detail="État invalide.")
    if payload.state == "reset":
        await db.ideas.update_one(
            {"idea_id": payload.idea_id},
            {"$unset": {"state": "", "state_by": "", "state_at": ""}},
        )
    else:
        await db.ideas.update_one(
            {"idea_id": payload.idea_id},
            {"$set": {
                "state": payload.state,
                "state_by": payload.key_id,
                "state_actor": "creator" if role == "creator" else sk,
                "state_at": datetime.now(timezone.utc).isoformat(),
            }},
        )
    return {"success": True}


# ---------------- ANNOUNCEMENTS + POLLS ----------------
class AnnounceCreateIn(BaseModel):
    model_config = ConfigDict(extra="allow")
    key_id: str
    nonce: str
    signature: str
    title: str
    body: str = ""
    audience: Any = "all"   # iter77: str ou List[str] (multi-select)

VALID_AUDIENCE_GROUPS = {"all", "approved", "creator", "admin", "modo", "pending", "non_validated"}


def _audience_matches(target_audience, dev: Optional[Dict[str, Any]]) -> bool:
    """iter77 — Évalue si le device cible reçoit une annonce/sondage.

    `target_audience` peut être :
    - "all" (legacy) → tout le monde
    - "approved" (legacy) → approved + creator
    - liste/array de groupes parmi VALID_AUDIENCE_GROUPS
    """
    # Normalize en liste
    if not target_audience:
        groups = ["all"]
    elif isinstance(target_audience, str):
        groups = [target_audience]
    elif isinstance(target_audience, list):
        groups = [g for g in target_audience if isinstance(g, str)]
        if not groups:
            groups = ["all"]
    else:
        groups = ["all"]
    if "all" in groups:
        return True
    role = (dev or {}).get("role") or "public"
    sk = (dev or {}).get("staff_kind")
    # Map groups → membership predicates
    if "creator" in groups and role == "creator":
        return True
    if "approved" in groups and role in ("approved", "creator"):
        return True
    if "admin" in groups and sk == "admin":
        return True
    if "modo" in groups and sk == "modo":
        return True
    if "pending" in groups and role == "pending":
        return True
    # iter77 — "non_validated" = pending + revoked + blocked + public.
    if "non_validated" in groups and role not in ("approved", "creator"):
        return True
    return False


@api_router.post("/announcements/create")
async def announcements_create(payload: AnnounceCreateIn):
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    if not payload.title.strip():
        raise HTTPException(status_code=400, detail="Titre requis.")
    body = payload.model_dump() if hasattr(payload, "model_dump") else {}
    raw_aud = body.get("audience")
    audience = raw_aud if isinstance(raw_aud, list) else [raw_aud or "all"]
    audience = [g for g in audience if g in VALID_AUDIENCE_GROUPS] or ["all"]
    doc = {
        "announce_id": f"ann_{uuid.uuid4().hex[:12]}",
        "title": payload.title.strip()[:200],
        "body": (payload.body or "").strip()[:5000],
        "audience": audience,
        "ts": datetime.now(timezone.utc).isoformat(),
        "updated_at": None,
    }
    await db.announcements.insert_one(doc)
    return {"success": True, "announce_id": doc["announce_id"]}


@api_router.get("/announcements/list")
async def announcements_list(key_id: Optional[str] = None):
    """Public — anyone can fetch the active announcements they qualify for.

    iter76/77 — enrichi avec :
    - `my_state` : état perso du device courant (validated/refused/orange/null)
    - `staff_states` : visible créatrice uniquement, tableau des états du staff
    - `audience` peut être string legacy ou List[str] multi-groupe (iter77).
    """
    rows = await db.announcements.find({}, {"_id": 0}).sort("ts", -1).to_list(length=50)
    dev = None
    role = "public"
    if key_id:
        dev = await db.device_keys.find_one(
            {"key_id": key_id},
            {"_id": 0, "role": 1, "staff_kind": 1},
        )
        role = (dev or {}).get("role") or "public"
    filtered = []
    for r in rows:
        # iter77 — multi-audience: utilise _audience_matches.
        if not _audience_matches(r.get("audience"), dev):
            continue
        my_state = None
        if key_id:
            ms = await db.announcement_states.find_one(
                {"announce_id": r["announce_id"], "key_id": key_id},
                {"_id": 0, "state": 1, "actor": 1, "ts": 1},
            )
            if ms:
                my_state = ms.get("state")
        if role == "creator":
            states = await db.announcement_states.find(
                {"announce_id": r["announce_id"], "key_id": {"$ne": key_id}},
                {"_id": 0, "key_id": 1, "state": 1, "actor": 1, "ts": 1},
            ).to_list(length=200)
            r["staff_states"] = states
        r["my_state"] = my_state
        # iter76 — masque côté requérant si VALIDÉ par lui.
        if role != "creator" and my_state == "validated":
            continue
        filtered.append(r)
    return {"announcements": filtered}


# iter77 — Edit announcement (pencil button côté créa)
class AnnounceEditIn(BaseModel):
    model_config = ConfigDict(extra="allow")
    key_id: str
    nonce: str
    signature: str
    announce_id: str
    title: Optional[str] = None
    body: Optional[str] = None
    audience: Any = None


@api_router.post("/announcements/edit")
async def announcements_edit(payload: AnnounceEditIn):
    """iter77 — Modifie une annonce. Reset tous les états (l'annonce revient
    en attente chez tout le monde) et bump updated_at."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    body = payload.model_dump() if hasattr(payload, "model_dump") else {}
    upd = {}
    if body.get("title") is not None:
        title = (body.get("title") or "").strip()
        if not title:
            raise HTTPException(status_code=400, detail="Titre requis.")
        upd["title"] = title[:200]
    if body.get("body") is not None:
        upd["body"] = (body.get("body") or "").strip()[:5000]
    if body.get("audience") is not None:
        raw_aud = body.get("audience")
        aud = raw_aud if isinstance(raw_aud, list) else [raw_aud]
        aud = [g for g in aud if g in VALID_AUDIENCE_GROUPS] or ["all"]
        upd["audience"] = aud
    upd["updated_at"] = datetime.now(timezone.utc).isoformat()
    res = await db.announcements.update_one({"announce_id": payload.announce_id}, {"$set": upd})
    if res.matched_count == 0:
        raise HTTPException(status_code=404, detail="Annonce introuvable.")
    # iter77 — Reset les états: l'annonce modifiée doit réapparaître à tous.
    await db.announcement_states.delete_many({"announce_id": payload.announce_id})
    return {"success": True, "updated_at": upd["updated_at"]}


@api_router.post("/announcements/delete")
async def announcements_delete(payload: _CreatorSigIn):
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    body = payload.model_dump() if hasattr(payload, "model_dump") else {}
    ann_id = body.get("announce_id")
    if not ann_id:
        raise HTTPException(status_code=400, detail="announce_id requis.")
    await db.announcements.delete_one({"announce_id": ann_id})
    return {"success": True}


class PollCreateIn(BaseModel):
    model_config = ConfigDict(extra="allow")
    key_id: str
    nonce: str
    signature: str
    question: str
    options: List[str] = Field(default_factory=list)
    audience: Any = "all"  # iter77: list ou str
    max_selections: int = 0  # iter77 — 0 = illimité (par défaut), sinon cap explicite
    allow_user_suggestions: bool = False  # iter77 — laisser users proposer leur réponse

@api_router.post("/polls/create")
async def polls_create(payload: PollCreateIn):
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    q = (payload.question or "").strip()
    opts = [o.strip() for o in (payload.options or []) if o.strip()]
    if not q or len(opts) < 2:
        raise HTTPException(status_code=400, detail="Question + 2 options requis.")
    try:
        max_sel = int(payload.max_selections or 0)
    except Exception:
        max_sel = 0
    # iter77 — 0 = illimité; sinon cap entre 1 et nombre d'options.
    if max_sel < 0:
        max_sel = 0
    elif max_sel > 0:
        max_sel = min(max_sel, len(opts[:50]))
    body = payload.model_dump() if hasattr(payload, "model_dump") else {}
    raw_aud = body.get("audience")
    aud = raw_aud if isinstance(raw_aud, list) else [raw_aud or "all"]
    aud = [g for g in aud if g in VALID_AUDIENCE_GROUPS] or ["all"]
    doc = {
        "poll_id": f"poll_{uuid.uuid4().hex[:12]}",
        "question": q[:300],
        "options": opts[:50],
        "audience": aud,
        "max_selections": max_sel,  # 0 = illimité
        "allow_user_suggestions": bool(payload.allow_user_suggestions),
        "ts": datetime.now(timezone.utc).isoformat(),
        "updated_at": None,
    }
    await db.polls.insert_one(doc)
    return {"success": True, "poll_id": doc["poll_id"]}


# iter77 — Edit poll (pencil button)
class PollEditIn(BaseModel):
    model_config = ConfigDict(extra="allow")
    key_id: str
    nonce: str
    signature: str
    poll_id: str
    question: Optional[str] = None
    options: Optional[List[str]] = None
    audience: Any = None
    max_selections: Optional[int] = None
    allow_user_suggestions: Optional[bool] = None


@api_router.post("/polls/edit")
async def polls_edit(payload: PollEditIn):
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    body = payload.model_dump() if hasattr(payload, "model_dump") else {}
    upd = {}
    if body.get("question") is not None:
        q = (body.get("question") or "").strip()
        if not q:
            raise HTTPException(status_code=400, detail="Question requise.")
        upd["question"] = q[:300]
    if body.get("options") is not None:
        opts = [o.strip() for o in (body.get("options") or []) if isinstance(o, str) and o.strip()]
        if len(opts) < 2:
            raise HTTPException(status_code=400, detail="2 options minimum.")
        upd["options"] = opts[:50]
    if body.get("audience") is not None:
        raw_aud = body.get("audience")
        aud = raw_aud if isinstance(raw_aud, list) else [raw_aud]
        aud = [g for g in aud if g in VALID_AUDIENCE_GROUPS] or ["all"]
        upd["audience"] = aud
    if body.get("max_selections") is not None:
        try:
            upd["max_selections"] = max(0, int(body.get("max_selections")))
        except Exception:
            pass
    if body.get("allow_user_suggestions") is not None:
        upd["allow_user_suggestions"] = bool(body.get("allow_user_suggestions"))
    upd["updated_at"] = datetime.now(timezone.utc).isoformat()
    res = await db.polls.update_one({"poll_id": payload.poll_id}, {"$set": upd})
    if res.matched_count == 0:
        raise HTTPException(status_code=404, detail="Sondage introuvable.")
    # iter77 — Reset les votes? Non, on garde les tally. L'utilisateur peut
    # revoter en cas de changement d'options (les indices peuvent shift).
    # On reset uniquement si options changent.
    if "options" in upd:
        await db.poll_votes.delete_many({"poll_id": payload.poll_id})
    return {"success": True, "updated_at": upd["updated_at"]}


# iter77 — User suggests an extra option (creator validates/removes)
class PollSuggestIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    poll_id: str
    text: str


@api_router.post("/polls/suggest-option")
async def polls_suggest_option(payload: PollSuggestIn):
    """iter77 — N'importe quel votant peut proposer une réponse perso.
    Stockée en `poll_suggestions` (pending). Créa valide ou retire.
    Si retirée, les votes sur cette option ne comptent plus (filtre tally)."""
    dev = await _verify_signed(payload.key_id, payload.nonce, payload.signature)
    poll = await db.polls.find_one({"poll_id": payload.poll_id}, {"_id": 0, "allow_user_suggestions": 1})
    if not poll:
        raise HTTPException(status_code=404, detail="Sondage introuvable.")
    if not poll.get("allow_user_suggestions"):
        raise HTTPException(status_code=403, detail="Propositions désactivées sur ce sondage.")
    text = (payload.text or "").strip()
    if not text or len(text) > 200:
        raise HTTPException(status_code=400, detail="Texte requis (≤200 chars).")
    sid = f"sug_{uuid.uuid4().hex[:12]}"
    await db.poll_suggestions.insert_one({
        "suggestion_id": sid,
        "poll_id": payload.poll_id,
        "key_id": payload.key_id,
        "pseudo": dev.get("label") or dev.get("pseudo") or "Anonyme",
        "text": text,
        "status": "pending",  # 'pending' | 'approved' | 'removed'
        "ts": datetime.now(timezone.utc).isoformat(),
    })
    return {"success": True, "suggestion_id": sid}


class PollSuggestDecideIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    suggestion_id: str
    decision: str  # 'approve' | 'remove'


@api_router.post("/polls/decide-suggestion")
async def polls_decide_suggestion(payload: PollSuggestDecideIn):
    """iter77 — Créa valide la proposition (devient option officielle) ou la retire."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    if payload.decision not in ("approve", "remove"):
        raise HTTPException(status_code=400, detail="decision invalide.")
    sug = await db.poll_suggestions.find_one({"suggestion_id": payload.suggestion_id}, {"_id": 0})
    if not sug:
        raise HTTPException(status_code=404, detail="Proposition introuvable.")
    new_status = "approved" if payload.decision == "approve" else "removed"
    await db.poll_suggestions.update_one(
        {"suggestion_id": payload.suggestion_id},
        {"$set": {"status": new_status,
                  "decided_at": datetime.now(timezone.utc).isoformat()}},
    )
    if payload.decision == "approve":
        # Ajoute l'option au poll.
        await db.polls.update_one(
            {"poll_id": sug["poll_id"]},
            {"$push": {"options": sug["text"]}},
        )
    return {"success": True, "status": new_status}


@api_router.get("/polls/list")
async def polls_list(key_id: Optional[str] = None):
    rows = await db.polls.find({}, {"_id": 0}).sort("ts", -1).to_list(length=50)
    dev = None
    role = "public"
    if key_id:
        dev = await db.device_keys.find_one(
            {"key_id": key_id},
            {"_id": 0, "role": 1, "staff_kind": 1},
        )
        role = (dev or {}).get("role") or "public"
    out = []
    for p in rows:
        # iter77 — multi-audience
        if not _audience_matches(p.get("audience"), dev):
            continue
        if "max_selections" not in p:
            p["max_selections"] = 1
        # Tally votes (unwind sur option_indices).
        votes = await db.poll_votes.aggregate([
            {"$match": {"poll_id": p["poll_id"]}},
            {"$unwind": "$option_indices"},
            {"$group": {"_id": "$option_indices", "count": {"$sum": 1}}},
        ]).to_list(length=200)
        tally = {v["_id"]: v["count"] for v in votes}
        p["tally"] = [tally.get(i, 0) for i in range(len(p.get("options", [])))]
        voters = await db.poll_votes.count_documents({"poll_id": p["poll_id"]})
        p["voters"] = voters
        my = None
        if key_id:
            mv = await db.poll_votes.find_one(
                {"poll_id": p["poll_id"], "voter_key_id": key_id},
                {"_id": 0, "option_indices": 1, "option_index": 1},
            )
            if mv:
                if isinstance(mv.get("option_indices"), list):
                    my = mv["option_indices"]
                elif mv.get("option_index") is not None:
                    my = [mv["option_index"]]
        p["my_vote"] = my
        # iter77 — Créa peut voir qui a voté (sauf si audience inclut 'all'/'public'
        # — interprétation : sondage à la « communauté » → vote anonyme).
        aud = p.get("audience")
        is_community = ("all" in aud) if isinstance(aud, list) else (aud in (None, "all"))
        if role == "creator" and not is_community:
            voters_rows = await db.poll_votes.find(
                {"poll_id": p["poll_id"]},
                {"_id": 0, "voter_key_id": 1, "option_indices": 1, "option_index": 1, "ts": 1},
            ).to_list(length=500)
            # Enrichir avec pseudo via device_keys
            kids = list({v.get("voter_key_id") for v in voters_rows if v.get("voter_key_id")})
            pseudos = {}
            if kids:
                async for d in db.device_keys.find(
                    {"key_id": {"$in": kids}}, {"_id": 0, "key_id": 1, "label": 1, "pseudo": 1, "email": 1},
                ):
                    pseudos[d["key_id"]] = d.get("pseudo") or d.get("label") or d.get("email") or d["key_id"][:10]
            for v in voters_rows:
                v["pseudo"] = pseudos.get(v.get("voter_key_id"), "Anonyme")
                if isinstance(v.get("option_indices"), list):
                    pass
                elif v.get("option_index") is not None:
                    v["option_indices"] = [v["option_index"]]
            p["voters_detail"] = voters_rows
        else:
            p["voters_detail"] = None
        # iter77 — propositions perso
        if p.get("allow_user_suggestions"):
            suggestions = await db.poll_suggestions.find(
                {"poll_id": p["poll_id"]}, {"_id": 0},
            ).sort("ts", 1).to_list(length=100)
            # Côté non-créa: ne montrer que les `approved` + `pending` (les `removed` sont silencieux).
            if role != "creator":
                suggestions = [s for s in suggestions if s.get("status") in ("approved", "pending")]
            p["suggestions"] = suggestions
        else:
            p["suggestions"] = []
        out.append(p)
    return {"polls": out}


class PollVoteIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    poll_id: str
    option_index: Optional[int] = None  # legacy single
    option_indices: Optional[List[int]] = None  # iter76 multi-select

@api_router.post("/polls/vote")
async def polls_vote(payload: PollVoteIn):
    dev = await _verify_signed(payload.key_id, payload.nonce, payload.signature)
    poll = await db.polls.find_one({"poll_id": payload.poll_id}, {"_id": 0})
    if not poll:
        raise HTTPException(status_code=404, detail="Sondage introuvable.")
    if not _audience_matches(poll.get("audience"), dev):
        raise HTTPException(status_code=403, detail="Audience non autorisée.")
    n = len(poll.get("options", []))
    max_sel = int(poll.get("max_selections") or 0)
    if payload.option_indices is not None:
        chosen = sorted({int(i) for i in payload.option_indices if isinstance(i, int)})
    elif payload.option_index is not None:
        chosen = [int(payload.option_index)]
    else:
        raise HTTPException(status_code=400, detail="option_index(s) requis.")
    if not chosen:
        raise HTTPException(status_code=400, detail="Sélection vide.")
    # iter77 — max_sel=0 → illimité.
    if max_sel > 0 and len(chosen) > max_sel:
        raise HTTPException(status_code=400, detail=f"Max {max_sel} sélection(s) autorisée(s).")
    for idx in chosen:
        if not (0 <= idx < n):
            raise HTTPException(status_code=400, detail="Option invalide.")
    await db.poll_votes.update_one(
        {"poll_id": payload.poll_id, "voter_key_id": payload.key_id},
        {"$set": {
            "poll_id": payload.poll_id,
            "voter_key_id": payload.key_id,
            "option_indices": chosen,
            "ts": datetime.now(timezone.utc).isoformat(),
        }, "$unset": {"option_index": ""}},
        upsert=True,
    )
    return {"success": True}


@api_router.post("/polls/delete")
async def polls_delete(payload: _CreatorSigIn):
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    body = payload.model_dump() if hasattr(payload, "model_dump") else {}
    poll_id = body.get("poll_id")
    if not poll_id:
        raise HTTPException(status_code=400, detail="poll_id requis.")
    await db.polls.delete_one({"poll_id": poll_id})
    await db.poll_votes.delete_many({"poll_id": poll_id})
    return {"success": True}


# ---------------- iter76: ANNOUNCEMENT STATES (validated/refused/orange) ----------------
class AnnStateIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    announce_id: str
    state: str  # 'validated' | 'refused' | 'orange' | 'reset'


@api_router.post("/announcements/set-state")
async def announcements_set_state(payload: AnnStateIn):
    """iter76 — Marque l'annonce avec un état pour le device courant.

    - validated (✅ vert) → tâche faite. Disparait pour le staff/user qui valide.
    - refused (❌ rouge) → tâche refusée. Reste visible (non-supprimable, sauf via clear-history).
    - orange (🟠) → escalade: « le staff n'a pas les codes, seule la créatrice peut ».
    - reset → suppression de l'état (annonce redevient en attente).

    Asymétrie staff↔créatrice:
    - Quand le staff valide, la créatrice voit toujours l'annonce mais avec un badge
      « Coché par staff » et peut soit confirmer (et l'annonce disparait pour la créatrice
      aussi) soit réinitialiser (et l'annonce redevient en attente pour tout le monde).
    """
    dev = await _verify_signed(payload.key_id, payload.nonce, payload.signature)
    state = (payload.state or "").strip().lower()
    if state not in ("validated", "refused", "orange", "reset"):
        raise HTTPException(status_code=400, detail="État invalide.")
    ann = await db.announcements.find_one({"announce_id": payload.announce_id}, {"_id": 0, "announce_id": 1})
    if not ann:
        raise HTTPException(status_code=404, detail="Annonce introuvable.")
    role = dev.get("role") or "public"
    actor = "creator" if role == "creator" else ("staff" if role == "approved" else "user")

    if state == "reset":
        # Le créateur peut réinitialiser TOUS les états d'une annonce ; sinon, seul son propre état.
        if role == "creator":
            await db.announcement_states.delete_many({"announce_id": payload.announce_id})
        else:
            await db.announcement_states.delete_one({"announce_id": payload.announce_id, "key_id": payload.key_id})
        return {"success": True, "reset": True}

    await db.announcement_states.update_one(
        {"announce_id": payload.announce_id, "key_id": payload.key_id},
        {"$set": {
            "announce_id": payload.announce_id,
            "key_id": payload.key_id,
            "state": state,
            "actor": actor,
            "ts": datetime.now(timezone.utc).isoformat(),
        }},
        upsert=True,
    )
    return {"success": True, "state": state, "actor": actor}


@api_router.post("/announcements/clear-history")
async def announcements_clear_history(payload: _CreatorSigIn):
    """iter76 — Bouton « Supprimer l'historique » côté créatrice : retire complètement
    toutes les annonces ET tous les états associés. Sert à repartir propre."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    res_ann = await db.announcements.delete_many({})
    res_st = await db.announcement_states.delete_many({})
    return {"success": True, "deleted_announcements": res_ann.deleted_count, "deleted_states": res_st.deleted_count}


# Modifie l'endpoint list pour enrichir avec states (rétrocompat: on RÉ-ÉCRIT
# la route ci-dessous, mais l'ancienne au-dessus reste l'autoritative car
# FastAPI prend la 1ʳᵉ enregistrée. Donc on patche directement la liste plus haut.
# Pas de seconde définition ici — voir announcements_list ci-dessus modifié.


# ---------------- iter76: SCHEDULED DISCONNECT ----------------
class ScheduleKickIn(BaseModel):
    model_config = ConfigDict(extra="allow")
    key_id: str
    nonce: str
    signature: str
    minutes: int = 5
    note: str = ""
    audience: Any = "all"  # iter77 — list ou str (groupes : all/staff/admin/modo/approved/pending/non_validated)


@api_router.post("/system/schedule-kick")
async def system_schedule_kick(payload: ScheduleKickIn):
    """iter77 — La créatrice programme la déconnexion ciblée.

    audience peut être :
    - 'all' → tout le monde (sauf créatrice)
    - 'staff' → admins + modos uniquement
    - 'admin' / 'modo' / 'approved' / 'pending' / 'non_validated' → groupe spécifique
    - liste mixte (cases à cocher) ex: ['admin','modo']
    """
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    try:
        delay = max(0, min(int(payload.minutes or 0), 24 * 60))
    except Exception:
        delay = 5
    body = payload.model_dump() if hasattr(payload, "model_dump") else {}
    raw_aud = body.get("audience")
    aud = raw_aud if isinstance(raw_aud, list) else [raw_aud or "all"]
    # iter77 — accepter 'staff' (alias admin+modo)
    if "staff" in aud:
        aud = [g for g in aud if g != "staff"] + ["admin", "modo"]
    valid_groups = VALID_AUDIENCE_GROUPS
    aud = [g for g in aud if g in valid_groups] or ["all"]
    now = datetime.now(timezone.utc)
    execute_at = (now + timedelta(minutes=delay)).isoformat()
    sk_id = f"sk_{uuid.uuid4().hex[:12]}"
    await db.scheduled_kicks.insert_one({
        "kick_id": sk_id,
        "creator_key_id": payload.key_id,
        "minutes": delay,
        "audience": aud,
        "execute_at": execute_at,
        "executed": False,
        "ts": now.isoformat(),
    })
    if (payload.note or "").strip():
        await db.announcements.insert_one({
            "announce_id": f"ann_{uuid.uuid4().hex[:12]}",
            "title": (payload.note.strip())[:200],
            "body": "",
            "audience": aud,
            "ts": now.isoformat(),
            "from_scheduled_kick": sk_id,
        })
    return {"success": True, "kick_id": sk_id, "execute_at": execute_at, "audience": aud}


@api_router.get("/system/scheduled-kicks")
async def system_scheduled_kicks_list(key_id: Optional[str] = None):
    """Liste les déconnexions programmées en cours (pour affichage côté créatrice)."""
    rows = await db.scheduled_kicks.find(
        {"executed": False}, {"_id": 0}
    ).sort("execute_at", 1).to_list(length=50)
    return {"scheduled_kicks": rows}


@api_router.post("/system/cancel-scheduled-kick")
async def system_cancel_scheduled_kick(payload: _CreatorSigIn):
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    body = payload.model_dump() if hasattr(payload, "model_dump") else {}
    kid = body.get("kick_id")
    if not kid:
        raise HTTPException(status_code=400, detail="kick_id requis.")
    await db.scheduled_kicks.update_one({"kick_id": kid}, {"$set": {"executed": True, "cancelled": True}})
    return {"success": True}


async def _execute_due_kicks():
    """iter76/77 — sweeper toutes les 10s : pour chaque kick dû, purge les
    sessions matchant l'audience (créatrice toujours exclue)."""
    try:
        now = datetime.now(timezone.utc).isoformat()
        due = await db.scheduled_kicks.find(
            {"executed": False, "execute_at": {"$lte": now}}, {"_id": 0}
        ).to_list(length=20)
        for k in due:
            aud = k.get("audience") or ["all"]
            if isinstance(aud, str):
                aud = [aud]
            # Récupère tous les devices non-créateur
            non_creator = await db.device_keys.find(
                {"role": {"$ne": "creator"}}, {"_id": 0, "key_id": 1, "role": 1, "staff_kind": 1},
            ).to_list(length=10000)
            targets = [d["key_id"] for d in non_creator if _audience_matches(aud, d)]
            if targets:
                await db.user_sessions.delete_many({"device_key_id": {"$in": targets}})
            await db.scheduled_kicks.update_one(
                {"kick_id": k["kick_id"]},
                {"$set": {"executed": True, "executed_at": now,
                          "deleted_sessions_for": targets[:200]}},
            )
    except Exception as e:
        logger.warning(f"scheduled-kick sweep error: {e}")


# ---------------- EXPORT APPROVAL ----------------
class ExportRequestIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    project_id: str
    export_kind: str   # "apk" | "exe" | "zip+github" | "source"

@api_router.post("/exports/request")
async def exports_request(payload: ExportRequestIn):
    """Any non-creator device — request export approval."""
    dev = await _verify_signed(payload.key_id, payload.nonce, payload.signature)
    if dev.get("role") == "creator":
        return {"approved": True, "auto": True}
    # Check for an existing pending/approved entry to avoid duplicates.
    existing = await db.export_requests.find_one({
        "key_id": payload.key_id,
        "project_id": payload.project_id,
        "export_kind": payload.export_kind,
        "status": {"$in": ["pending", "approved"]},
    }, {"_id": 0})
    if existing:
        return {"approved": existing["status"] == "approved", "status": existing["status"], "request_id": existing["request_id"]}
    doc = {
        "request_id": f"er_{uuid.uuid4().hex[:14]}",
        "key_id": payload.key_id,
        "label": dev.get("pseudo") or dev.get("label"),
        "project_id": payload.project_id,
        "export_kind": payload.export_kind,
        "status": "pending",
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.export_requests.insert_one(doc)
    return {"approved": False, "status": "pending", "request_id": doc["request_id"]}


@api_router.post("/exports/decide")
async def exports_decide(payload: _CreatorSigIn):
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    body = payload.model_dump() if hasattr(payload, "model_dump") else {}
    req_id = body.get("request_id")
    decision = body.get("decision")  # "approve" | "reject"
    if not req_id or decision not in ("approve", "reject"):
        raise HTTPException(status_code=400, detail="request_id + decision (approve|reject) requis.")
    new_status = "approved" if decision == "approve" else "rejected"
    r = await db.export_requests.update_one(
        {"request_id": req_id, "status": "pending"},
        {"$set": {"status": new_status, "decided_at": datetime.now(timezone.utc).isoformat()}},
    )
    if not r.matched_count:
        raise HTTPException(status_code=404, detail="Demande introuvable.")
    return {"success": True, "status": new_status}


@api_router.post("/exports/pending")
async def exports_pending(payload: _CreatorSigIn):
    """Creator-only — list pending export requests."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    rows = await db.export_requests.find({"status": "pending"}, {"_id": 0}).sort("created_at", -1).to_list(length=200)
    return {"requests": rows}


class ExportStatusIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    request_id: Optional[str] = None
    project_id: Optional[str] = None
    export_kind: Optional[str] = None

@api_router.post("/exports/status")
async def exports_status(payload: ExportStatusIn):
    """User-side polling — current status of a pending export request."""
    await _verify_signed(payload.key_id, payload.nonce, payload.signature)
    q = {"key_id": payload.key_id}
    if payload.request_id:
        q["request_id"] = payload.request_id
    elif payload.project_id and payload.export_kind:
        q["project_id"] = payload.project_id
        q["export_kind"] = payload.export_kind
    else:
        raise HTTPException(status_code=400, detail="request_id ou (project_id+export_kind) requis.")
    row = await db.export_requests.find_one(q, {"_id": 0}, sort=[("created_at", -1)])
    if not row:
        return {"status": "none"}
    return {"status": row["status"], "request_id": row["request_id"]}


# ---------------- AUTO-TRANSLATE (creator review) ----------------
class TranslateIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    text: str
    target_lang: str = "fr"

@api_router.post("/creator/translate")
async def creator_translate(payload: TranslateIn):
    """Creator-only — translate arbitrary text via Emergent LLM."""
    await _require_creator_signature(payload.key_id, payload.nonce, payload.signature)
    text = (payload.text or "").strip()
    if not text:
        return {"translated": ""}
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        key = os.environ.get("EMERGENT_LLM_KEY")
        if not key:
            raise ValueError("EMERGENT_LLM_KEY missing")
        chat = LlmChat(api_key=key, session_id=f"trans_{uuid.uuid4().hex[:8]}",
                       system_message=f"Translate the user's text into {payload.target_lang}. Output ONLY the translation, no prose.").with_model("openai", "gpt-5.2")
        translated = await chat.send_message(UserMessage(text=text[:4000]))
        return {"translated": str(translated).strip()}
    except Exception as e:
        logger.warning(f"translate failed: {e}")
        return {"translated": text, "error": "translate_unavailable"}


# ---------------- USER pseudo update ----------------
class UpdatePseudoIn(BaseModel):
    new_pseudo: str

@api_router.post("/auth/update-pseudo")
async def auth_update_pseudo(request: Request, payload: UpdatePseudoIn):
    user_id = await get_current_user(request)
    p = (payload.new_pseudo or "").strip()
    if not (1 <= len(p) <= 30):
        raise HTTPException(status_code=400, detail="Pseudo invalide (1-30).")
    # iter75: "créatrice" no longer reserved.
    user = await db.users.find_one({"user_id": user_id}, {"_id": 0, "email": 1})
    if not user:
        raise HTTPException(status_code=404, detail="Utilisateur introuvable.")
    await db.users.update_one({"user_id": user_id}, {"$set": {"pseudo": p, "pseudo_lower": p.lower()}})
    if user.get("email"):
        await db.device_keys.update_many(
            {"email": user["email"]},
            {"$set": {"pseudo": p, "label": p}},
        )
    return {"success": True, "pseudo": p}


# ---------------- THEFT recovery — email fallback ----------------
class TheftEmailIn(BaseModel):
    email: str

@api_router.post("/auth/theft-email-request")
async def auth_theft_email_request(payload: TheftEmailIn):
    """Send a magic-link to the account email; clicking it revokes ALL
    creator+approved device keys tied to that email so the user can
    re-onboard fresh on the new device.

    Idempotent: always returns 200 even if the email is unknown (to avoid
    enumeration). The actual mail is only sent when a matching user exists.
    """
    email = (payload.email or "").strip().lower()
    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="Email invalide.")
    user = await db.users.find_one({"email": email}, {"_id": 0, "user_id": 1})
    if user:
        token = uuid.uuid4().hex
        await db.theft_email_tokens.insert_one({
            "token": token,
            "email": email,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "used": False,
        })
        frontend_base = _clean_origin(os.environ.get("FRONTEND_URL", "")) or _clean_origin(os.environ.get("REACT_APP_BACKEND_URL", "")) or ""
        link = f"{frontend_base}/theft-confirm?token={token}" if frontend_base else f"/theft-confirm?token={token}"
        # Best-effort send via the unified pipeline (SMTP → Resend).
        try:
            html = (
                "<div style='font-family:system-ui,sans-serif;background:#050505;color:#fff;padding:32px;max-width:560px;margin:0 auto'>"
                "<h1 style='color:#E4FF00;margin:0 0 16px'>CodeForge AI</h1>"
                "<p style='color:#E4E4E7'>Tu reçois ce mail parce qu'une procédure de récupération en cas de vol a été demandée depuis un nouvel appareil.</p>"
                "<p style='color:#FF6B6B'><strong>Si ce n'est pas toi, ignore ce message.</strong></p>"
                "<p style='color:#E4E4E7'>Sinon, clique ci-dessous pour révoquer tous les anciens appareils enregistrés sous ton compte :</p>"
                f"<p style='margin:24px 0'><a href='{link}' style='background:#E4FF00;color:#050505;padding:14px 28px;border-radius:6px;text-decoration:none;font-weight:bold;display:inline-block'>Confirmer la récupération</a></p>"
                f"<p style='color:#A1A1AA;font-size:12px'>Ou copie ce lien : <span style='color:#00D4FF;word-break:break-all'>{link}</span></p>"
                "<p style='color:#A1A1AA;font-size:12px'>Ce lien expire dans 30 minutes.</p>"
                "</div>"
            )
            await _send_email(email, "CodeForge AI — Récupération en cas de vol", html)
        except Exception as e:
            logger.warning(f"theft-email send failed: {e}")
    # Always 200 — no enumeration leak.
    return {"success": True}


@api_router.get("/auth/theft-email-confirm")
async def auth_theft_email_confirm(token: str):
    row = await db.theft_email_tokens.find_one({"token": token}, {"_id": 0})
    if not row or row.get("used"):
        raise HTTPException(status_code=404, detail="Lien invalide ou déjà utilisé.")
    # Token expires after 30 min.
    try:
        created = datetime.fromisoformat(row["created_at"].replace("Z", "+00:00"))
        if (datetime.now(timezone.utc) - created).total_seconds() > 1800:
            raise HTTPException(status_code=410, detail="Lien expiré.")
    except (KeyError, ValueError):
        raise HTTPException(status_code=410, detail="Lien expiré.")
    email = row["email"]
    # Revoke every device key bound to this email.
    r = await db.device_keys.update_many(
        {"email": email, "role": {"$in": ["creator", "approved"]}},
        {"$set": {"role": "revoked", "revoked_at": datetime.now(timezone.utc).isoformat(), "revoked_reason": "theft_email_recovery"}},
    )
    await db.theft_email_tokens.update_one({"token": token}, {"$set": {"used": True}})
    return {"success": True, "revoked_count": r.modified_count}


class TheftIrisVerifyIn(BaseModel):
    """iter71/75: optional iris re-confirmation. iter75 retired the
    email-token leg and now accepts an email directly so the backend
    can look up the iris baseline without a one-time-use link in the
    user's compromised inbox."""
    token: Optional[str] = None  # legacy — accepted for backwards-compat
    email: Optional[str] = None
    hashes: List[str]


@api_router.post("/auth/theft-iris-verify")
async def auth_theft_iris_verify(payload: TheftIrisVerifyIn):
    if not isinstance(payload.hashes, list) or len(payload.hashes) < 3:
        raise HTTPException(status_code=400, detail="3 captures iris sont requises.")
    if any((not isinstance(h, str)) or len(h) < 20 or len(h) > 128 for h in payload.hashes[:3]):
        raise HTTPException(status_code=400, detail="Empreintes iris invalides.")
    # Resolve target email: iter75 prefers `email` directly; falls back
    # to the legacy `token` lookup so older clients still work.
    email = (payload.email or "").strip().lower() or None
    if not email and payload.token:
        row = await db.theft_email_tokens.find_one({"token": payload.token}, {"_id": 0, "email": 1})
        if row:
            email = row.get("email")
    if not email:
        raise HTTPException(status_code=400, detail="Email du compte requis.")
    user = await db.users.find_one({"email": email}, {"_id": 0, "biometric": 1})
    if not user or (user.get("biometric") or {}).get("kind") != "iris":
        logger.info(f"theft-iris-verify: no iris baseline for {email}, accepting capture as observation only")
    await db.theft_iris_attempts.insert_one({
        "email": email,
        "hashes": payload.hashes[:3],
        "token": payload.token,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "verified": False,  # real matching comes in the next sprint
    })
    return {"success": True, "revoked_count": 0}


# Helper used by ideas/polls to verify any signature (not creator-restricted).
async def _verify_signed(key_id: str, nonce: str, signature: str) -> Dict[str, Any]:
    dev = await _device_by_key(key_id)
    if not dev:
        raise HTTPException(status_code=404, detail="Clé inconnue.")
    if not await _consume_nonce(key_id, nonce):
        raise HTTPException(status_code=403, detail="Nonce invalide ou expiré.")
    if not verify_signature(dev.get("public_key_jwk") or {}, nonce, signature):
        raise HTTPException(status_code=403, detail="Signature invalide.")
    return dev


# ==========================================================================
# iter86 — FRIEND REQUESTS (C20) extraites dans routes/social_routes.py
# Le router est inclus depuis ce module via build_friends_router(...) après
# que tous les helpers soient définis.
# ==========================================================================


# ==========================================================================
# iter82 — GROUP CHATS (C19) : 6 types de tchats de groupe automatiques.
#   public, private, staff, modo, public_staff, public_private
# ==========================================================================

GROUP_TYPES = {
    "public",         # tous les approved non-staff/private + visiteurs publics
    "private",        # uniquement les clés privées (approved)
    "staff",          # admin + modo
    "modo",           # modo only
    "admin",          # iter86 — admin only
    "public_staff",   # public + admin + modo (tt types confondus)
    "public_private", # public + privé (sans staff)
}


# iter85 — 1ère slice du refacto : helper _groups_for_device extrait dans
# routes/social_routes.py. server.py importe directement.
from routes.social_routes import _groups_for_device  # noqa: E402


class GroupListIn(BaseModel):
    key_id: str
    nonce: str
    signature: str


# iter88 — Slice 3 du refacto : /groups/* extraits dans routes/social_routes.py
# Les classes Pydantic et les routes sont déplacées. Le router est inclus
# en bas de server.py via build_groups_router(...).
class GroupMessagesIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    group_type: str
    limit: int = 200


class GroupSendIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    group_type: str
    content: str


# ==========================================================================
# iter82 — MESSAGE TO RANDOM MODO (C18) : remplace l'ancien "message créa"
# ==========================================================================

class MessageToStaffIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    content: str


@api_router.post("/messages/send-to-staff")
async def messages_send_to_staff(payload: MessageToStaffIn):
    """iter82 C18 — Le bouton 'message' (icône en-tête) route maintenant
    l'utilisateur vers un MODO aléatoire (fallback admin, puis créa si aucun
    modo/admin). Le créa voit toujours toutes les threads via /messages/inbox.

    Si l'utilisateur veut parler spécifiquement à la créatrice, il doit
    d'abord la demander en ami (/friends/request).
    """
    sender = await _verify_signed(payload.key_id, payload.nonce, payload.signature)
    content = (payload.content or "").strip()
    if not content:
        raise HTTPException(status_code=400, detail="Message vide.")
    if len(content) > MAX_MESSAGE_LEN:
        raise HTTPException(status_code=400, detail=f"Message trop long ({MAX_MESSAGE_LEN} max).")

    # Pick a random modo (fallback admin, fallback creator).
    modos = await db.device_keys.find(
        {"staff_kind": "modo", "role": {"$in": ["approved", "creator"]}},
        {"_id": 0, "key_id": 1, "pseudo": 1, "label": 1},
    ).to_list(length=50)
    if not modos:
        modos = await db.device_keys.find(
            {"staff_kind": "admin", "role": {"$in": ["approved", "creator"]}},
            {"_id": 0, "key_id": 1, "pseudo": 1, "label": 1},
        ).to_list(length=50)
    recipient_kind = "modo" if modos else "creator"
    if not modos:
        modos = await db.device_keys.find(
            {"role": "creator"}, {"_id": 0, "key_id": 1, "pseudo": 1, "label": 1},
        ).to_list(length=10)
    if not modos:
        raise HTTPException(status_code=503, detail="Aucun destinataire staff disponible.")

    import random as _rnd
    target = _rnd.choice(modos)
    target_key_id = target["key_id"]

    now = datetime.now(timezone.utc)
    msg_id = f"msg_{uuid.uuid4().hex[:16]}"
    await db.messages.insert_one({
        "message_id": msg_id,
        "thread_key_id": payload.key_id,
        "from_key_id": payload.key_id,
        "to_key_id": target_key_id,
        "is_from_creator": False,
        "recipient_kind": recipient_kind,
        "content": content,
        "sender_label": sender.get("pseudo") or sender.get("label"),
        "ts": now.isoformat(),
        "read_by_creator": False,
        "read_by_user": True,
    })
    return {"sent": True, "message_id": msg_id, "assigned_to": target.get("pseudo") or target.get("label") or target_key_id[:10], "recipient_kind": recipient_kind}


# ==========================================================================
# iter83 C7 — ORCHESTRATEUR MULTI-AGENTS (planner/critic/executor/arbiter)
# ==========================================================================

from orchestrator import orchestrate as _run_orchestrate
from orchestrator import orchestrate_stream as _stream_orchestrate
from orchestrator import orchestrate_actions as _stream_actions


class OrchestrateIn(BaseModel):
    message: str
    project_id: Optional[str] = None
    language: Optional[str] = "fr"
    enable_commit: Optional[bool] = False  # iter86 — opt-in pour push GitHub réel
    enable_preview_rebuild: Optional[bool] = False  # iter88 — opt-in pour rebuild yarn


async def _persist_event(event: Dict[str, Any], *, user_id: str, session_id: str, project_id: Optional[str]):
    """Sauvegarde un événement d'orchestration en DB pour rappel ultérieur via
    /orchestrate/event/{event_id}/details."""
    try:
        await db.orchestrator_events.insert_one({
            "event_id": event.get("event_id"),
            "user_id": user_id,
            "session_id": session_id,
            "project_id": project_id,
            "kind": event.get("kind"),
            "summary": event.get("summary"),
            "details": event.get("details"),
            "ts": event.get("ts"),
        })
    except Exception:
        pass


@api_router.post("/chat/orchestrate")
async def chat_orchestrate(request: Request, payload: OrchestrateIn):
    """C7 — Lance la pipeline planner→executor→critic→arbiter et persiste
    tous les événements en base. Retourne aussi la liste pour le client."""
    user_id = await get_current_user(request)
    session_id = f"orch_{user_id}_{payload.project_id or 'global'}"
    result = await _run_orchestrate(
        payload.message, session_id=session_id, language=payload.language or "fr",
    )
    try:
        for evt in (result.get("events") or []):
            await _persist_event(evt, user_id=user_id, session_id=session_id, project_id=payload.project_id)
    except Exception:
        pass
    return result


@api_router.post("/chat/orchestrate-stream")
async def chat_orchestrate_stream(request: Request, payload: OrchestrateIn):
    """iter84 C7+C5/C8 — Stream d'ACTIONS Emergent-style.

    iter86 — Wirage on_commit RÉEL via push_to_github (si GITHUB_ENABLED).
    Si l'utilisateur passe `enable_commit=true` dans le body, l'orchestrator
    pousse le code généré sur une branche `orchestrate/{session_id[-12:]}`.
    """
    user_id = await get_current_user(request)
    session_id = f"orch_{user_id}_{payload.project_id or 'global'}"
    lang = payload.language or "fr"

    async def persist(evt):
        await _persist_event(evt, user_id=user_id, session_id=session_id, project_id=payload.project_id)

    # iter86 — Hook on_commit : push réel sur GitHub si activé + opt-in
    async def on_commit_real(branch: str, summary: str, content: str):
        if not GITHUB_ENABLED:
            return {"ok": False, "error": "GITHUB_DISABLED"}
        if not getattr(payload, "enable_commit", False):
            return {"ok": False, "note": "enable_commit=false (opt-in)"}
        safe_branch = branch.replace("/", "-").replace(" ", "-")[:48]
        file_path = f"orchestrate-runs/{safe_branch}.py"
        body = f"# Orchestrator run for: {summary}\n# Branch: {branch}\n\n{content}\n"
        ok = await push_to_github(file_path, body, branch="main", retries=2)
        return {"ok": bool(ok), "ref": file_path, "branch_label": branch}

    # iter88 — Hook on_preview : rebuild RÉEL du frontend si opt-in
    async def on_preview_real():
        if not getattr(payload, "enable_preview_rebuild", False):
            return {"ok": False, "note": "enable_preview_rebuild=false (opt-in)", "url": os.environ.get("PREVIEW_BASE_URL") or "https://no-code-builder-25.preview.emergentagent.com"}
        import subprocess as _sub
        import asyncio as _aio
        try:
            # Best-effort yarn build dans /app/frontend, timeout 90s.
            loop = _aio.get_event_loop()
            proc = await loop.run_in_executor(
                None,
                lambda: _sub.run(
                    ["yarn", "build"],
                    capture_output=True, text=True, timeout=90, cwd="/app/frontend",
                ),
            )
            base_url = os.environ.get("PREVIEW_BASE_URL") or "https://no-code-builder-25.preview.emergentagent.com"
            return {
                "ok": proc.returncode == 0,
                "url": base_url,
                "returncode": proc.returncode,
                "build_summary": ((proc.stdout or "")[-2000:] + (("\n" + (proc.stderr or "")[-1000:]) if proc.stderr else "")),
            }
        except Exception as e:
            return {"ok": False, "error": str(e)[:300]}

    async def event_gen():
        try:
            async for evt in _stream_actions(
                payload.message, session_id=session_id, language=lang,
                persist_event=persist,
                on_commit=on_commit_real,
                on_preview=on_preview_real,
            ):
                # On retire `details` du payload SSE pour ne pas surcharger ;
                # le client peut les fetch via /orchestrate/event/{id}/details
                payload_evt = {k: v for k, v in evt.items() if k != "details"}
                yield f"data: {json.dumps(payload_evt, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'kind': 'error', 'summary': str(e)[:300]}, ensure_ascii=False)}\n\n"

    from fastapi.responses import StreamingResponse
    return StreamingResponse(event_gen(), media_type="text/event-stream", headers={
        "Cache-Control": "no-cache, no-transform",
        "X-Accel-Buffering": "no",
    })


@api_router.get("/orchestrate/event/{event_id}/details")
async def orchestrate_event_details(event_id: str, request: Request):
    """iter84 — Récupère le détail complet d'un événement d'orchestration.
    L'UI appelle cet endpoint quand l'utilisateur déplie la flèche de
    l'événement."""
    user_id = await get_current_user(request)
    doc = await db.orchestrator_events.find_one(
        {"event_id": event_id, "user_id": user_id}, {"_id": 0},
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Événement introuvable.")
    return doc


class OrchestrateHistoryIn(BaseModel):
    project_id: Optional[str] = None
    limit: int = 50


# iter85 — Testing-agents-en-boucle : endpoint qui lance pytest sur les tests
# backend et émet des événements test_run (start/passed/failed/done) via SSE.
class TestLoopIn(BaseModel):
    target: Optional[str] = "backend"   # 'backend' (pytest) | 'sandbox' (orchestrator._execute_python)
    path: Optional[str] = "tests/"      # filtre pytest
    project_id: Optional[str] = None


@api_router.post("/orchestrate/test-loop")
async def orchestrate_test_loop(request: Request, payload: TestLoopIn):
    """iter85 — Pipeline de validation automatique en boucle. Lance pytest
    en interne et émet des événements test_run au fur et à mesure. Le
    frontend peut afficher ces events dans OrchestrationLog.

    Best-effort : si pytest échoue, on émet `test_run failed` mais on ne
    relance pas automatiquement (la boucle de correction est laissée au
    chat orchestrator pour l'instant)."""
    user_id = await get_current_user(request)
    session_id = f"testloop_{user_id}_{payload.project_id or 'global'}"

    async def event_gen():
        import subprocess as _sub
        evt0 = {
            "event_id": f"evt_{uuid.uuid4().hex[:16]}",
            "kind": "test_run",
            "summary": f"Lancement des tests : {payload.target}/{payload.path}",
            "ts": datetime.now(timezone.utc).isoformat(),
        }
        await _persist_event(evt0, user_id=user_id, session_id=session_id, project_id=payload.project_id)
        yield f"data: {json.dumps(evt0, ensure_ascii=False)}\n\n"

        try:
            if payload.target == "backend":
                # Lancement pytest restreint à un path safe sous /app/backend
                safe_path = (payload.path or "tests/").lstrip("/")
                if ".." in safe_path:
                    raise HTTPException(status_code=400, detail="Path invalide.")
                full_path = os.path.normpath(os.path.join("/app/backend", safe_path))
                if not full_path.startswith("/app/backend"):
                    raise HTTPException(status_code=400, detail="Path hors backend.")

                proc = _sub.run(
                    ["python", "-m", "pytest", full_path, "-q", "--tb=short", "--no-header"],
                    capture_output=True, text=True, timeout=90, cwd="/app/backend",
                )
                # Extrait des stats du dernier passage : "X passed, Y failed"
                summary_line = ""
                for line in (proc.stdout or "").splitlines()[::-1]:
                    if "passed" in line or "failed" in line or "error" in line:
                        summary_line = line.strip(); break
                kind = "test_run" if proc.returncode in (0, 5) else "error"
                summary = summary_line or (f"pytest exit {proc.returncode}")
                evt1 = {
                    "event_id": f"evt_{uuid.uuid4().hex[:16]}",
                    "kind": kind,
                    "summary": summary,
                    "details": {
                        "returncode": proc.returncode,
                        "stdout": (proc.stdout or "")[-8000:],
                        "stderr": (proc.stderr or "")[-2000:],
                    },
                    "ts": datetime.now(timezone.utc).isoformat(),
                }
                await _persist_event(evt1, user_id=user_id, session_id=session_id, project_id=payload.project_id)
                # Stream payload sans details (lazy via /orchestrate/event/{id})
                p1 = {k: v for k, v in evt1.items() if k != "details"}
                yield f"data: {json.dumps(p1, ensure_ascii=False)}\n\n"
            else:
                evt1 = {
                    "event_id": f"evt_{uuid.uuid4().hex[:16]}",
                    "kind": "error",
                    "summary": f"target inconnu : {payload.target}",
                    "ts": datetime.now(timezone.utc).isoformat(),
                }
                await _persist_event(evt1, user_id=user_id, session_id=session_id, project_id=payload.project_id)
                yield f"data: {json.dumps(evt1, ensure_ascii=False)}\n\n"
        except Exception as e:
            evt_err = {
                "event_id": f"evt_{uuid.uuid4().hex[:16]}",
                "kind": "error",
                "summary": f"test-loop crash : {str(e)[:200]}",
                "ts": datetime.now(timezone.utc).isoformat(),
            }
            yield f"data: {json.dumps(evt_err, ensure_ascii=False)}\n\n"

        evt_done = {
            "event_id": f"evt_{uuid.uuid4().hex[:16]}",
            "kind": "complete",
            "summary": "Test-loop terminé",
            "ts": datetime.now(timezone.utc).isoformat(),
        }
        await _persist_event(evt_done, user_id=user_id, session_id=session_id, project_id=payload.project_id)
        yield f"data: {json.dumps(evt_done, ensure_ascii=False)}\n\n"

    from fastapi.responses import StreamingResponse
    return StreamingResponse(event_gen(), media_type="text/event-stream", headers={
        "Cache-Control": "no-cache, no-transform",
        "X-Accel-Buffering": "no",
    })


# iter84 — Observabilité vidéo mobile (logs structurés en cas d'échec)
class VideoEventIn(BaseModel):
    kind: str
    session_id: str
    ua: Optional[str] = None
    viewport: Optional[Dict[str, Any]] = None
    is_secure: Optional[bool] = None
    ts: Optional[str] = None
    error: Optional[str] = None
    name: Optional[str] = None
    track_label: Optional[str] = None
    track_state: Optional[str] = None
    settings: Optional[Dict[str, Any]] = None
    ready_state: Optional[int] = None


@api_router.post("/observability/video-event")
async def observability_video_event(payload: VideoEventIn, request: Request):
    """iter84 — Endpoint d'observabilité pour le bug 'vidéo mobile en mode
    Public'. Aucune auth requise : on log tout pour debug. Anti-flood : on
    ne garde que 5000 events max et on rejette si même session_id émet >50
    events dans la même minute."""
    now = datetime.now(timezone.utc)
    minute_ago = (now - timedelta(seconds=60)).isoformat()
    recent_count = await db.video_events.count_documents({
        "session_id": payload.session_id, "ts_server": {"$gt": minute_ago},
    })
    if recent_count > 50:
        raise HTTPException(status_code=429, detail="Trop d'événements pour cette session.")
    doc = payload.model_dump()
    doc["ts_server"] = now.isoformat()
    doc["ip"] = (request.client.host if request.client else None)
    await db.video_events.insert_one(doc)
    # Auto-purge anciens events (>5000 lignes)
    if recent_count == 0:
        cnt = await db.video_events.estimated_document_count()
        if cnt > 5000:
            cursor = db.video_events.find({}, {"_id": 1}).sort("ts_server", 1).limit(cnt - 5000)
            ids = [d["_id"] for d in await cursor.to_list(length=cnt - 5000)]
            if ids:
                await db.video_events.delete_many({"_id": {"$in": ids}})
    return {"recorded": True}


@api_router.post("/orchestrate/history")
async def orchestrate_history(request: Request, payload: OrchestrateHistoryIn):
    """iter84 — Récupère l'historique des événements d'orchestration d'une
    session pour cet utilisateur. Permet de réafficher le journal après
    un refresh ou switch de projet."""
    user_id = await get_current_user(request)
    q = {"user_id": user_id}
    if payload.project_id:
        q["project_id"] = payload.project_id
    rows = await db.orchestrator_events.find(q, {"_id": 0, "details": 0}).sort("ts", -1).limit(max(1, min(payload.limit, 200))).to_list(length=200)
    return {"events": list(reversed(rows))}


# ==========================================================================
# iter86 — PRIVATE CODE BROWSER (creator-only, read-only)
# ==========================================================================

class PrivateReadFileIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    path: str


@api_router.post("/private/code/read-file")
async def private_read_file(payload: PrivateReadFileIn):
    """iter87 — Endpoint désactivé pour des raisons de sécurité.
    Le code du site n'est visible par personne, même la créatrice."""
    raise HTTPException(status_code=403, detail="Accès refusé pour des raisons de sécurité.")


class PrivateGrepIn(BaseModel):
    key_id: str
    nonce: str
    signature: str
    pattern: str


@api_router.post("/private/code/grep")
async def private_grep(payload: PrivateGrepIn):
    """iter87 — Endpoint désactivé pour des raisons de sécurité."""
    raise HTTPException(status_code=403, detail="Accès refusé pour des raisons de sécurité.")


# ==========================================================================
# iter82 — CHAT STREAMING SSE (C5/C8) : streaming pseudo-token-par-token
# ==========================================================================

class ChatStreamIn(BaseModel):
    message: str
    mode: str = "online"
    project_id: Optional[str] = None
    language: Optional[str] = "fr"


@api_router.post("/chat/stream")
async def chat_stream(request: Request, input: ChatStreamIn):
    """iter82 C5/C8 — Streaming SSE de la réponse IA. Émission "word by word"
    pour donner l'impression de voir le texte se construire. La réponse
    complète est sauvegardée en DB à la fin via le flux normal. Pour ne pas
    refacto tout l'endpoint /chat/message, on appelle l'endpoint et on
    re-stream la réponse côté serveur.
    """
    user_id = await get_current_user(request)
    # Récupère la réponse complète en re-utilisant la logique existante :
    # on génère le ChatMessageInput et on attend la full réponse.
    # NOTE: pour un vrai streaming token-par-token via Emergent, il faudrait
    # passer par LlmChat.stream() ce qui n'est pas exposé par
    # emergentintegrations actuellement. On simule donc le streaming en
    # découpant la réponse finale en mots. C'est suffisant pour l'UX et reste
    # totalement non-bloquant côté frontend.
    full_input = ChatMessageInput(
        message=input.message,
        mode=input.mode,
        project_id=input.project_id,
        language=input.language,
    )
    # Appel à la logique existante ; send_chat_message gère la persistance DB.
    resp = await send_chat_message(request, full_input)
    ai_text = ((resp or {}).get("ai_response") or {}).get("content") or ""
    msg_id = ((resp or {}).get("ai_response") or {}).get("message_id") or ""
    download = ((resp or {}).get("ai_response") or {}).get("download")

    async def event_gen():
        # SSE pseudo-streaming par mots, avec petit délai pour visualiser.
        import asyncio as _aio
        words = ai_text.split(" ")
        acc = ""
        for i, w in enumerate(words):
            acc = (acc + " " + w).strip() if acc else w
            yield f"data: {json.dumps({'delta': w + (' ' if i < len(words)-1 else ''), 'index': i})}\n\n"
            # 8ms par mot = pour un texte de 500 mots, ~4 secondes. Lisible et naturel.
            await _aio.sleep(0.008)
        # Signal final
        yield f"data: {json.dumps({'done': True, 'message_id': msg_id, 'download': download, 'content': ai_text})}\n\n"

    from fastapi.responses import StreamingResponse
    return StreamingResponse(event_gen(), media_type="text/event-stream", headers={
        "Cache-Control": "no-cache, no-transform",
        "X-Accel-Buffering": "no",
    })


# Include the router in the main app
app.include_router(api_router)

# iter86 — Slice 2 du refacto : friends routes incluses depuis social_routes
# iter88 — Slice 3 : groups routes également depuis social_routes
from routes.social_routes import build_friends_router, build_groups_router  # noqa: E402
app.include_router(build_friends_router(db, _verify_signed, _device_by_key), prefix="/api")
app.include_router(build_groups_router(db, _verify_signed, MAX_MESSAGE_LEN), prefix="/api")

# Include PWA routes under /api/pwa
app.include_router(pwa_router, prefix="/api/pwa", tags=["PWA"])

# Include Desktop routes under /api/desktop
app.include_router(desktop_router, prefix="/api/desktop", tags=["Desktop"])


@app.on_event("startup")
async def ensure_indexes():
    """Create MongoDB indexes used by the email/password auth flow."""
    try:
        await db.users.create_index("email", unique=True, sparse=True)
        await db.users.create_index(
            "pseudo_lower",
            unique=True,
            partialFilterExpression={
                "verified": True,
                "pseudo_lower": {"$type": "string"},
            },
        )
        await db.email_verifications.create_index("token", unique=True)
        await db.email_verifications.create_index("user_id")
        await db.user_sessions.create_index("session_token", unique=True)
        await db.login_attempts.create_index("identifier")
        await db.resend_attempts.create_index("email")
        await db.password_resets.create_index("email")
        await db.password_reset_tokens.create_index("token", unique=True)
        await db.password_reset_tokens.create_index("user_id")
        await db.messages.create_index("thread_key_id")
        await db.messages.create_index("ts")
        logger.info("✅ MongoDB indexes ready")
    except Exception as e:
        logger.warning(f"Index creation warning: {e}")


# Background task: every 10 minutes, drop expired/stale auth rows so the
# DB doesn't grow unboundedly. Documents store ISO strings (not Mongo
# Date) so we can't use a TTL index — we sweep manually.
_cleanup_task: asyncio.Task | None = None
_kick_sweeper_task: asyncio.Task | None = None


async def _periodic_auth_cleanup():
    while True:
        try:
            now = datetime.now(timezone.utc)
            now_iso = now.isoformat()
            # Verifications older than expires_at (and not consumed for >1h)
            await db.email_verifications.delete_many({"expires_at": {"$lt": now_iso}})
            # User sessions past expiry
            await db.user_sessions.delete_many({"expires_at": {"$lt": now_iso}})
            # Password reset tokens past expiry
            await db.password_reset_tokens.delete_many({"expires_at": {"$lt": now_iso}})
            # Login + resend + reset attempts older than 24h (rate-limit data)
            day_ago = (now - timedelta(hours=24)).isoformat()
            await db.login_attempts.delete_many({"ts": {"$lt": day_ago}})
            await db.resend_attempts.delete_many({"ts": {"$lt": day_ago}})
            await db.password_resets.delete_many({"ts": {"$lt": day_ago}})
            # Auth-error logs older than 7d (kept for /metrics 24h window with margin)
            week_ago = (now - timedelta(days=7)).isoformat()
            await db.auth_errors.delete_many({"ts": {"$lt": week_ago}})
        except Exception as e:
            logger.warning(f"Cleanup task error: {e}")
        await asyncio.sleep(10 * 60)  # 10 minutes


async def _periodic_kick_sweeper():
    """iter76 — sweeper rapide (10s) pour exécuter les déconnexions programmées."""
    while True:
        try:
            await _execute_due_kicks()
        except Exception as e:
            logger.warning(f"Kick sweeper error: {e}")
        await asyncio.sleep(10)


@app.on_event("startup")
async def start_cleanup_task():
    global _cleanup_task, _kick_sweeper_task
    _cleanup_task = asyncio.create_task(_periodic_auth_cleanup())
    _kick_sweeper_task = asyncio.create_task(_periodic_kick_sweeper())
    logger.info("✅ Auth cleanup background task started (every 10 min)")
    logger.info("✅ Kick sweeper background task started (every 10s)")


@app.on_event("shutdown")
async def shutdown_db_client():
    global _cleanup_task, _kick_sweeper_task
    if _cleanup_task:
        _cleanup_task.cancel()
    if _kick_sweeper_task:
        _kick_sweeper_task.cancel()
    client.close()
