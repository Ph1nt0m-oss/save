"""iter124 — Routes lourdes /chat/* extraites de server.py (4 endpoints).

  - POST /chat/analyze-attachment    : upload + extraction PDF/DOCX/XLSX/PPTX/SQLite/IMG
  - GET  /chat/models                : catalogue des modèles IA disponibles (UI selector)
  - GET  /chat/export-ipynb/{pid}    : conversation → Jupyter notebook
  - GET  /chat/export-docx/{pid}     : conversation → Word .docx

Helpers injectés : db, logger, get_current_user, sanitize_filename,
analyze_pdf, analyze_docx, analyze_xlsx, analyze_pptx, analyze_sqlite,
analyze_image_with_vision.
"""
import base64
import io as _io
import json
import re
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, File, HTTPException, Request, UploadFile
from fastapi.responses import Response
from pydantic import BaseModel


def build_chat_exports_router(
    db,
    *,
    get_current_user,
    logger,
    sanitize_filename,
    analyze_pdf,
    analyze_docx,
    analyze_xlsx,
    analyze_pptx,
    analyze_sqlite,
    analyze_image_with_vision,
):
    router = APIRouter()

    @router.post("/chat/analyze-attachment")
    async def chat_analyze_attachment(request: Request, file: UploadFile = File(...)):
        """Extract the usable content of an uploaded file for the chat.

        Returns a JSON object with a `kind` ('text' or 'image') and the content the
        frontend should embed in the next chat message.
        """
        _ = await get_current_user(request)  # auth gate
        data = await file.read()
        if len(data) > 20 * 1024 * 1024:  # 20 MB cap
            raise HTTPException(status_code=413, detail="Fichier trop lourd (max 20 Mo)")

        name = file.filename or "attachment"
        mime = (file.content_type or "").lower()
        lower = name.lower()

        if mime.startswith("image/") or lower.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif")):
            description = await analyze_image_with_vision(data, mime)
            return {
                "kind": "image",
                "filename": name,
                "mime_type": mime or "image/png",
                "content": description,
                "data_base64": base64.b64encode(data).decode("utf-8"),
            }

        text = ""
        if mime == "application/pdf" or lower.endswith(".pdf"):
            text = await analyze_pdf(data)
        elif lower.endswith(".docx") or mime.endswith("wordprocessingml.document"):
            text = await analyze_docx(data)
        elif lower.endswith(".xlsx") or mime.endswith("spreadsheetml.sheet"):
            text = await analyze_xlsx(data)
        elif lower.endswith(".pptx") or mime.endswith("presentationml.presentation"):
            text = await analyze_pptx(data)
        elif lower.endswith((".sqlite", ".db", ".sqlite3")):
            text = await analyze_sqlite(data)
        elif lower.endswith((".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".log",
                            ".py", ".js", ".jsx", ".ts", ".tsx", ".html", ".htm", ".css", ".scss",
                            ".xml", ".yaml", ".yml", ".ini", ".env", ".cfg", ".conf", ".toml",
                            ".sql", ".sh", ".ps1", ".bat", ".cmd", ".rb", ".go", ".rs", ".java",
                            ".c", ".cpp", ".h", ".hpp", ".cs", ".php", ".kt", ".swift")):
            try:
                text = data.decode("utf-8", errors="ignore")[:30000]
            except Exception:
                text = ""
        else:
            try:
                text = data.decode("utf-8", errors="ignore")[:20000]
            except Exception:
                text = ""

        if not text.strip():
            raise HTTPException(status_code=415, detail="Impossible d'extraire le contenu du fichier.")

        return {"kind": "text", "filename": name, "mime_type": mime or "text/plain", "content": text}

    @router.get("/chat/models")
    async def list_chat_models(request: Request):
        """Liste les modèles disponibles pour le sélecteur, avec descriptions
        contextuelles (chat vs création) pour aider l'utilisateur à choisir.

        Le frontend passe `?context=chat` ou `?context=create` pour adapter le wording.
        """
        await get_current_user(request)
        context = (request.query_params.get("context") or "chat").lower()
        is_create = context == "create"

        # Descriptions adaptées au contexte d'utilisation
        # iter87 — Liste mise à jour avec les meilleures versions de chaque famille
        # Emergent. Les anciens IDs (gpt-5.2/claude-opus/claude-sonnet/gemini-3-pro/
        # gemini-3-flash) restent supportés via MODEL_ROUTES — ces entrées sont
        # purement UI (le mapping provider/model_id côté backend doit suivre).
        online = [
            {
                "id": "emergent-collab", "name": "Emergent (Multi-IA)", "provider": "Emergent", "badge": "Collaboration", "color": "fuchsia",
                "description": (
                    "MODE EMERGENT : orchestre Claude 5 Fable, GPT 5.5 et Gemini 3.1 Pro en parallèle. Chaque IA répond, puis un arbitre fusionne les meilleures parties. Sans choisir."
                    if not is_create else
                    "Combine les forces de 3 IA frontières en arrière-plan. Pour les projets critiques nécessitant le plus haut niveau de fiabilité."
                ),
                "good_for": (["Décisions importantes", "Auto-fact-checking", "Réponse maximale"]
                    if not is_create else ["Projets critiques", "Production", "Sécurité élevée"]),
            },
            {
                "id": "vexub-video", "name": "Vexub Vidéo", "provider": "Vexub", "badge": "Vidéo", "color": "pink",
                "description": (
                    "Génère des vidéos courtes (TikTok/Shorts/Reels) à partir d'un texte, audio ou vidéo longue. Voix off, sous-titres, montage automatique. Édit dans le navigateur."
                    if not is_create else
                    "Pour projets vidéo : génère shorts/reels automatiquement. Idéal landings avec démo vidéo intégrée."
                ),
                "good_for": (["Vidéo TikTok/Shorts", "Texte→vidéo", "Voix off", "Montage auto"]
                    if not is_create else ["Site avec vidéo intégrée", "Démo produit", "Landing vidéo"]),
            },
            {
                "id": "claude-5-fable", "name": "Claude 5 Fable", "provider": "Anthropic", "badge": "Le plus capable", "color": "violet",
                "description": (
                    "Le modèle le plus puissant et le plus sécurisé d'Anthropic. Recommandé pour les décisions critiques, audits, raisonnements complexes."
                    if not is_create else
                    "Génère le code le plus robuste et le plus sûr. Idéal pour apps sensibles (auth, paiements, médical, juridique)."
                ),
                "good_for": (["Décisions critiques", "Audit sécurité", "Analyse longue", "Raisonnement multi-étape"]
                    if not is_create else ["Apps sensibles", "Production-ready", "Sécurité élevée"]),
            },
            {
                "id": "gpt-5.5", "name": "GPT 5.5", "provider": "OpenAI", "badge": "Défaut", "color": "yellow",
                "description": (
                    "Le dernier modèle d'OpenAI. Conversation généraliste rapide et fluide, mémoire complète, raisonnement solide."
                    if not is_create else
                    "Génère le code complet du projet (FastAPI + React + DB). Hypothèses intelligentes, README clair, prêt à exécuter."
                ),
                "good_for": (["Brainstormer", "Écrire un email/texte", "Analyser un fichier", "Conversation longue"]
                    if not is_create else ["App complète", "Projet équilibré", "Site web standard"]),
            },
            {
                "id": "claude-4.8-opus", "name": "Claude 4.8 Opus", "provider": "Anthropic", "badge": "Thinking", "color": "amber",
                "description": (
                    "Performance frontière d'Anthropic. Raisonne avant de répondre. Idéal pour problèmes complexes, dilemmes, longues analyses."
                    if not is_create else
                    "Architecte avant de coder. Meilleure pour projets multi-fichiers, logique métier complexe, sécurité."
                ),
                "good_for": (["Problèmes complexes", "Dilemmes", "Code review profond", "Recherche approfondie"]
                    if not is_create else ["Architecture complexe", "Backend avec règles métier", "Apps multi-modules"]),
            },
            {
                "id": "claude-4.7-opus-1m", "name": "Claude 4.7 Opus (1M)", "provider": "Anthropic", "badge": "Contexte long", "color": "indigo",
                "description": (
                    "Contexte d'1 million de tokens. Idéal pour analyser un repo entier, plusieurs PDFs, ou des conversations très longues."
                    if not is_create else
                    "Pour porter un projet existant volumineux : repo legacy, refactor complet, lecture de gros datasets."
                ),
                "good_for": (["Analyser un repo", "Lire 100+ pages", "Conversation infinie"]
                    if not is_create else ["Refactor legacy", "Migration de framework", "Audit codebase"]),
            },
            {
                "id": "claude-4.6-sonnet", "name": "Claude 4.6 Sonnet", "provider": "Anthropic", "badge": "Code", "color": "orange",
                "description": (
                    "Excellente pour ÉCRIRE DU CODE dans le chat — clique sur ▶ Exécuter pour le lancer dans le sandbox Python."
                    if not is_create else
                    "Le PLUS RAPIDE pour générer un projet complet propre, exécutable, prêt à pousser sur GitHub. Recommandée par défaut."
                ),
                "good_for": (["Snippets de code", "Refactor", "Debug ligne par ligne", "Réécrire un texte"]
                    if not is_create else ["App standard", "Site marketing", "Outil CRUD", "Recommandé par défaut"]),
            },
            {
                "id": "gpt-5.3-codex", "name": "GPT 5.3 Codex", "provider": "OpenAI", "badge": "Code", "color": "emerald",
                "description": (
                    "Modèle flagship OpenAI spécialisé code. Idéal pour write/debug/refactor complexe avec exécution."
                    if not is_create else
                    "Optimisé pour les patterns Python/JS/TS modernes. Code propre, idiomatique, testé."
                ),
                "good_for": (["Debug complexe", "Tests unitaires", "Architecture code"]
                    if not is_create else ["Backend Python", "Frontend TS", "API REST/GraphQL"]),
            },
            {
                "id": "gemini-3.1-pro", "name": "Gemini 3.1 Pro", "provider": "Google", "badge": "Multimodal", "color": "blue",
                "description": (
                    "Le meilleur de Google. Idéal quand tu joins une IMAGE — décrit, analyse, OCR, explique des schémas."
                    if not is_create else
                    "Plus créative visuellement. Idéale pour UI originales, design audacieux, identité visuelle marquée."
                ),
                "good_for": (["Analyser une image", "OCR", "Lire un schéma", "Décrire une photo"]
                    if not is_create else ["UI design original", "Landing page", "Portfolio créatif"]),
            },
            {
                "id": "gpt-5.4-1m", "name": "GPT 5.4 (1M)", "provider": "OpenAI", "badge": "Contexte long", "color": "cyan",
                "description": (
                    "Variante 1M tokens de GPT 5.4. Pour ingérer de grandes quantités de docs en une seule passe."
                    if not is_create else
                    "Pour projets nécessitant un contexte massif (specs longues, multiples APIs externes)."
                ),
                "good_for": (["Analyse docs volumineux", "Multi-PDF", "Conversation sans coupure"]
                    if not is_create else ["Projet enterprise", "Specs complexes"]),
            },
            {
                "id": "grok-4.3", "name": "Grok 4.3", "provider": "xAI", "badge": "Temps réel", "color": "rose",
                "description": (
                    "Modèle xAI d'Elon Musk avec accès temps réel à X (Twitter). Idéal pour analyser des tendances, actualités, ou des sujets controversés sans filtre."
                    if not is_create else
                    "Génère du code avec un ton sarcastique optionnel. Idéal pour outils d'analyse sociale ou intégrations X/Twitter."
                ),
                "good_for": (["Actualités", "Tendances X", "Analyse sociale", "Sujets sans tabou"]
                    if not is_create else ["Bot X/Twitter", "Outil de veille", "Dashboard tendances"]),
            },
            {
                "id": "grok-4.20-reasoning", "name": "Grok 4.20 Reasoning", "provider": "xAI", "badge": "Thinking", "color": "rose",
                "description": (
                    "Variante raisonnement de Grok. Étape de réflexion explicite avant la réponse. Idéal pour résoudre des problèmes complexes."
                    if not is_create else
                    "Architecte avec un mode raisonnement profond. Pour projets nécessitant des décisions structurées."
                ),
                "good_for": (["Raisonnement complexe", "Maths", "Logique avancée"]
                    if not is_create else ["Architecture complexe", "Algorithmes pointus"]),
            },
            {
                "id": "lindy-flow", "name": "Lindy Flow", "provider": "Lindy", "badge": "Workflow", "color": "teal",
                "description": (
                    "Plateforme no-code de workflows IA. Crée des agents automatisés (email triage, calendrier, CRM) avec des étapes visuelles. Mode no-code par essence."
                    if not is_create else
                    "Génère des workflows d'agents IA pré-câblés (Slack/email/Calendar). Idéal pour automatiser un business sans coder."
                ),
                "good_for": (["Automatisation business", "Email triage", "Agent personnel"]
                    if not is_create else ["Outil interne", "Automatisation CRM", "Workflow no-code"]),
            },
        ]
        offline = [
            {
                "id": "deepseek", "name": "DeepSeek R1 (7B)", "provider": "Ollama", "badge": "Code", "color": "sky",
                "description": "Excellent en code et raisonnement step-by-step, totalement hors-ligne." if not is_create else "Génère du code propre hors-ligne. Lent mais privé.",
                "good_for": ["Code", "Maths", "Logique"] if not is_create else ["App standard hors-ligne"],
            },
            {
                "id": "gemma", "name": "Gemma 3 (12B)", "provider": "Ollama", "badge": "Équilibré", "color": "indigo",
                "description": "Modèle Google open-source. Polyvalent, multilingue, conversation naturelle, sans rien envoyer en ligne.",
                "good_for": ["Tout-terrain hors-ligne", "Multilingue", "Discussion privée"] if not is_create else ["Petites apps hors-ligne", "Privacy first"],
            },
            {
                "id": "gemma-27b", "name": "Gemma 3 (27B)", "provider": "Ollama", "badge": "Puissant", "color": "purple",
                "description": "Le plus capable des Gemma — exigeant en RAM (+24 Go). Qualité GPT-3.5 hors-ligne.",
                "good_for": ["Tâches lourdes hors-ligne", "Analyse profonde"] if not is_create else ["App complète hors-ligne haut de gamme"],
            },
            {
                "id": "gemma-4b", "name": "Gemma 3 (4B)", "provider": "Ollama", "badge": "Léger", "color": "violet",
                "description": "Très léger (~3 Go RAM), idéal pour machines modestes et démarrage rapide.",
                "good_for": ["Vieux laptop", "Tâches simples", "Test rapide"] if not is_create else ["Page statique", "MVP minimal"],
            },
            {
                "id": "llama", "name": "Llama 3.2", "provider": "Ollama", "badge": "Généraliste", "color": "emerald",
                "description": "Meta Llama — alternative équilibrée, communauté très active.",
                "good_for": ["Discussion générale"] if not is_create else ["App standard"],
            },
            {
                "id": "qwen", "name": "Qwen 2.5 (7B)", "provider": "Ollama", "badge": "Multilingue", "color": "teal",
                "description": "Très bon support chinois + anglais + français. Top pour i18n.",
                "good_for": ["Traduction", "Multilangue"] if not is_create else ["Site multilingue"],
            },
            {
                "id": "mistral", "name": "Mistral 7B", "provider": "Ollama", "badge": "Européen", "color": "rose",
                "description": "Modèle français performant, léger, respectueux du RGPD si tu héberges toi-même.",
                "good_for": ["Français natif", "RGPD"] if not is_create else ["App RGPD-friendly"],
            },
            {
                "id": "phi", "name": "Phi-3 Medium", "provider": "Ollama", "badge": "Compact", "color": "fuchsia",
                "description": "Microsoft Phi — petit mais étonnamment fort en maths et code.",
                "good_for": ["Maths", "Code"] if not is_create else ["Outils techniques"],
            },
        ]
        return {"online": online, "offline": offline, "context": context}

    @router.get("/chat/export-ipynb/{project_id}")
    async def export_chat_as_ipynb(project_id: str, request: Request):
        """Exporte une conversation chat en notebook Jupyter (.ipynb).

        Chaque message utilisateur devient une cellule markdown. Chaque bloc
        ```python``` trouvé dans les réponses AI devient une cellule code (avec le
        stdout déjà capturé en output si le bloc `▶️ Exécution Python (sandbox)` suit).
        """
        user_id = await get_current_user(request)
        # Validate ownership
        proj = await db.projects.find_one({"project_id": project_id, "user_id": user_id}, {"_id": 0})
        if not proj:
            raise HTTPException(status_code=404, detail="Projet introuvable")
        # Collect messages in order
        msgs = await db.chat_messages.find(
            {"user_id": user_id, "project_id": project_id},
            {"_id": 0}
        ).sort("timestamp", 1).to_list(length=None)
        if not msgs:
            raise HTTPException(status_code=400, detail="Aucun message à exporter")

        cells: List[Dict[str, Any]] = []
        # Title cell
        cells.append({
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                f"# {proj.get('name', 'Conversation')}\n",
                f"\n*Exporté depuis CodeForge AI — {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}*\n",
                "\n---\n",
            ],
        })

        code_block_re = re.compile(r"```(?:python|py)\s*\n(.*?)```", re.DOTALL)
        stdout_after_re = re.compile(
            r"\*\*▶️ Exécution Python \(sandbox\)[^\n]*\n+```\n(.*?)```",
            re.DOTALL,
        )

        for m in msgs:
            role = m.get("role", "assistant")
            content = m.get("content", "") or ""
            if role == "user":
                cells.append({
                    "cell_type": "markdown",
                    "metadata": {},
                    "source": ["**👤 Utilisateur**\n\n", content],
                })
            else:
                # Try to split AI response into: pre-code markdown, code cells, post-code markdown.
                pos = 0
                code_blocks = list(code_block_re.finditer(content))
                if not code_blocks:
                    cells.append({
                        "cell_type": "markdown",
                        "metadata": {},
                        "source": ["**🤖 CodeForge**\n\n", content],
                    })
                    continue
                # Preamble markdown
                prefix = content[: code_blocks[0].start()].rstrip()
                if prefix:
                    cells.append({
                        "cell_type": "markdown",
                        "metadata": {},
                        "source": ["**🤖 CodeForge**\n\n", prefix],
                    })
                for i, cb in enumerate(code_blocks):
                    code_src = cb.group(1).rstrip()
                    # Look for an inline sandbox stdout block right after this cell.
                    tail = content[cb.end():]
                    stdout_m = stdout_after_re.search(tail[:4000])
                    outputs = []
                    if stdout_m:
                        outputs.append({
                            "name": "stdout",
                            "output_type": "stream",
                            "text": stdout_m.group(1).splitlines(keepends=True),
                        })
                    cells.append({
                        "cell_type": "code",
                        "execution_count": i + 1,
                        "metadata": {},
                        "outputs": outputs,
                        "source": code_src.splitlines(keepends=True),
                    })
                    pos = cb.end()
                # Trailing markdown
                trailing = content[pos:]
                # Strip the "Exécution Python (sandbox)" blocks from trailing since we moved them.
                trailing_clean = re.sub(stdout_after_re, "", trailing).strip()
                if trailing_clean:
                    cells.append({
                        "cell_type": "markdown",
                        "metadata": {},
                        "source": [trailing_clean],
                    })

        notebook = {
            "cells": cells,
            "metadata": {
                "kernelspec": {"display_name": "Python 3", "language": "python", "name": "python3"},
                "language_info": {"name": "python", "version": "3.11"},
                "codeforge_export": {
                    "project_id": project_id,
                    "project_name": proj.get("name"),
                    "exported_at": datetime.now(timezone.utc).isoformat(),
                },
            },
            "nbformat": 4,
            "nbformat_minor": 5,
        }
        raw = json.dumps(notebook, ensure_ascii=False, indent=1).encode("utf-8")
        safe_name = sanitize_filename(proj.get("name") or project_id) + ".ipynb"
        return Response(
            content=raw,
            media_type="application/x-ipynb+json",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
        )


    @router.get("/chat/export-docx/{project_id}")
    async def export_chat_as_docx(project_id: str, request: Request):
        """iter79 — Exporte une conversation chat en .docx (Microsoft Word)."""
        from docx import Document
        from docx.shared import Pt, RGBColor
        import io as _io

        user_id = await get_current_user(request)
        proj = await db.projects.find_one({"project_id": project_id, "user_id": user_id}, {"_id": 0})
        if not proj:
            raise HTTPException(status_code=404, detail="Projet introuvable.")
        msgs = await db.chat_messages.find(
            {"project_id": project_id, "user_id": user_id}, {"_id": 0},
        ).sort("created_at", 1).to_list(length=10000)

        doc = Document()
        doc.add_heading(proj.get("name") or project_id, 0)
        doc.add_paragraph(f"Exporté le {datetime.now(timezone.utc).strftime('%d/%m/%Y %H:%M UTC')}")
        doc.add_paragraph(f"{len(msgs)} message(s)")
        doc.add_paragraph()
        for m in msgs:
            speaker = "Utilisateur" if m.get("role") == "user" else "IA"
            p = doc.add_paragraph()
            run = p.add_run(f"[{speaker}] ")
            run.bold = True
            run.font.color.rgb = RGBColor(0xE4, 0xFF, 0x00) if m.get("role") != "user" else RGBColor(0x00, 0xD4, 0xFF)
            p.add_run((m.get("content") or "")[:50000])
        buf = _io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        safe_name = sanitize_filename(proj.get("name") or project_id) + ".docx"
        return Response(
            content=buf.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
        )


    return router
